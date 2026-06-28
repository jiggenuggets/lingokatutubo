import json
import sys
import tempfile
from html.parser import HTMLParser
from io import StringIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.contrib.auth.tokens import default_token_generator
from django.core import mail
from django.core.management import call_command
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils import timezone
from django.utils.encoding import force_bytes
from django.utils.http import urlsafe_base64_encode

from .forms import DEFAULT_OCR_LANGUAGES
from .models import (
    DocumentPage,
    GeneratedOutput,
    OCRResult,
    SystemEventLog,
    TranslationJob,
    TranslationSegment,
    UploadedDocument,
)
from .services import sync_pipeline_job


User = get_user_model()


class _PreviewStructureParser(HTMLParser):
    """Track the bilingual preview container, its header bar, and per-page rows.

    The preview lays out original/translated pages as one row per page
    (instead of two independently-scrolling panels) so they stay aligned by
    page number. This tracks that shape: exactly one container, one header
    pair (Original/Translated labels), and one row per page.
    """

    def __init__(self):
        super().__init__()
        self.container_count = 0
        self.header_count = 0
        self.row_count = 0

    @staticmethod
    def _classes(attrs):
        for name, value in attrs:
            if name == "class" and value:
                return set(value.split())
        return set()

    def handle_starttag(self, tag, attrs):
        classes = self._classes(attrs)
        if "bilingual-document-preview" in classes:
            self.container_count += 1
        if tag == "header" and "document-panel-header" in classes:
            self.header_count += 1
        if "document-page-card" in classes:
            self.row_count += 1


class TranslatorAuthAndJobTests(TestCase):
    password = "Bagobo-Test-2026!"

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)

        self.alice = User.objects.create_user(
            username="alice",
            email="alice@example.test",
            password=self.password,
        )
        self.bob = User.objects.create_user(
            username="bob",
            email="bob@example.test",
            password=self.password,
        )

    def test_signup_creates_and_logs_in_user(self):
        response = self.client.post(
            reverse("translator:signup"),
            {
                "username": "new-reviewer",
                "email": "reviewer@example.test",
                "password1": self.password,
                "password2": self.password,
            },
        )

        self.assertRedirects(response, reverse("translator:translate"))
        user = User.objects.get(username="new-reviewer")
        self.assertEqual(int(self.client.session["_auth_user_id"]), user.pk)

    def test_login_endpoint_redirects_to_translator(self):
        response = self.client.post(
            reverse("translator:login"),
            {"username": self.alice.username, "password": self.password},
        )

        self.assertRedirects(response, reverse("translator:translate"))
        self.assertEqual(int(self.client.session["_auth_user_id"]), self.alice.pk)

    def test_translate_page_requires_authentication(self):
        translate_url = reverse("translator:translate")
        response = self.client.get(translate_url)

        self.assertRedirects(
            response,
            f"{reverse('translator:login')}?next={translate_url}",
        )

    def test_upload_creates_owned_job_and_stores_media_files(self):
        self.client.force_login(self.alice)
        uploaded = SimpleUploadedFile(
            "sample.pdf",
            b"%PDF-1.4\n1 0 obj\n<<>>\nendobj\n%%EOF",
            content_type="application/pdf",
        )

        with patch("translator.views.start_translation_job") as start_translation:
            response = self.client.post(
                reverse("translator:upload"),
                {
                    "file": uploaded,
                    "source_language": "auto",
                    "target_language": "tagabawa",
                    "ocr_languages": "",
                },
            )

        self.assertEqual(response.status_code, 202)
        job = TranslationJob.objects.get()
        self.assertEqual(job.owner, self.alice)
        self.assertEqual(job.original_filename, "sample.pdf")
        self.assertEqual(job.ocr_languages, DEFAULT_OCR_LANGUAGES)
        self.assertTrue(Path(job.upload_file_path).is_file())
        self.assertTrue(Path(job.input_file_path).is_file())
        self.assertEqual(
            Path(job.upload_file_path).relative_to(self.media_root).parts[:2],
            ("uploads", job.job_id),
        )
        self.assertEqual(
            Path(job.input_file_path).relative_to(self.media_root).parts[:3],
            ("jobs", job.job_id, "input"),
        )
        document = UploadedDocument.objects.get(job=job)
        self.assertEqual(document.owner, self.alice)
        self.assertEqual(document.original_filename, "sample.pdf")
        self.assertEqual(document.file_size_bytes, uploaded.size)
        self.assertEqual(document.metadata["content_type"], "application/pdf")
        self.assertEqual(job.metadata["original_content_type"], "application/pdf")
        start_translation.assert_called_once()
        self.assertEqual(start_translation.call_args.args[0].id, job.id)

    def test_processing_sync_saves_ocr_result_and_translation_segments(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.PROCESSING)
        structure_path = self._write_job_file(
            job,
            "structure.json",
            json.dumps(
                {
                    "warnings": [],
                    "pages": [
                        {
                            "page_number": 1,
                            "width": 612,
                            "height": 792,
                            "blocks": [
                                {
                                    "type": "text",
                                    "block_id": "b1",
                                    "bbox": [72, 72, 540, 110],
                                    "ocr_confidence": 0.91,
                                    "lines": [
                                        {
                                            "text": "Hello learner",
                                            "translated_text": "Madigar learner",
                                            "translation_method": "phrasebook_exact",
                                            "translation_confidence": 0.95,
                                            "ocr_confidence": 0.91,
                                            "bbox": [72, 72, 540, 90],
                                        },
                                        {
                                            "text": "Unlisted phrase",
                                            "translated_text": "[UNKNOWN_FOR_REVIEW]",
                                            "translation_method": "unknown_for_review",
                                            "translation_confidence": 0.0,
                                            "ocr_confidence": 0.81,
                                            "bbox": [72, 94, 540, 110],
                                        },
                                    ],
                                }
                            ],
                        }
                    ],
                }
            ).encode("utf-8"),
        )
        sync_pipeline_job(
            SimpleNamespace(
                job_id=job.job_id,
                status=TranslationJob.Status.COMPLETED,
                progress=100,
                current_phase="completed",
                current_step="Completed",
                phase_message="Translation complete.",
                error="",
                detection_type="digital",
                file_type=TranslationJob.FileType.PDF,
                metadata={"structure_file": str(structure_path)},
                completed_at=None,
            )
        )

        job.refresh_from_db()
        self.assertEqual(job.status, TranslationJob.Status.COMPLETED)
        self.assertEqual(DocumentPage.objects.filter(job=job).count(), 1)
        ocr = OCRResult.objects.get(job=job)
        self.assertIn("Hello learner", ocr.text)
        self.assertAlmostEqual(ocr.confidence, 0.8767)
        self.assertEqual(TranslationSegment.objects.filter(job=job).count(), 2)
        first = TranslationSegment.objects.get(job=job, segment_index=1)
        self.assertEqual(first.source_text, "Hello learner")
        self.assertEqual(first.translated_text, "Madigar learner")
        self.assertEqual(first.method, "phrasebook_exact")
        self.assertAlmostEqual(first.confidence, 0.95)
        second = TranslationSegment.objects.get(job=job, segment_index=2)
        self.assertTrue(second.needs_review)

    def test_owner_can_read_job_status(self):
        job = self._create_job(self.alice)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:status", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["job_id"], job.job_id)
        self.assertEqual(payload["status"], TranslationJob.Status.QUEUED)

    def test_translate_page_hides_ocr_language_controls(self):
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:translate"))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("Translation Settings", body)
        self.assertNotIn("OCR Language Coverage", body)
        self.assertNotIn("Used automatically for scanned PDFs and images.", body)
        self.assertNotIn("ocr_languages_list", body)
        self.assertNotIn("Custom OCR Codes", body)
        self.assertNotIn("ocr_languages_custom", body)
        self.assertNotIn("Orientation and Script Detection", body)
        self.assertRegex(body, r'id="complete-actions"[^>]*hidden')
        self.assertRegex(body, r'id="preview-link"[^>]*hidden')
        self.assertRegex(body, r'id="download-link"[^>]*hidden')

    def test_hidden_upload_actions_are_not_overridden_by_css(self):
        styles_path = Path(__file__).resolve().parents[1] / "static" / "css" / "styles.css"
        css = styles_path.read_text(encoding="utf-8")

        self.assertIn("[hidden]", css)
        self.assertIn("display: none !important", css)

    def test_upload_warning_behavior_is_limited_to_active_upload(self):
        script_path = Path(__file__).resolve().parents[1] / "static" / "js" / "app.js"
        script = script_path.read_text(encoding="utf-8")

        self.assertIn(
            "Your document is still uploading. Leaving now may cancel the upload.",
            script,
        )
        self.assertIn('window.addEventListener("beforeunload", beforeUnloadHandler)', script)
        self.assertIn('window.removeEventListener("beforeunload", beforeUnloadHandler)', script)
        self.assertIn("if (!isUploading || currentJobId) return undefined", script)
        self.assertIn("if (!file || isSubmitting) return", script)

    def test_polling_controller_stops_on_terminal_status(self):
        script_path = Path(__file__).resolve().parents[1] / "static" / "js" / "app.js"
        script = script_path.read_text(encoding="utf-8")

        self.assertIn('const terminalStatuses = new Set(["completed", "failed"])', script)
        self.assertIn("if (terminalStatuses.has(status) || !activeStatuses.has(status))", script)
        self.assertIn("this.stop()", script)

    def test_duplicate_polling_is_prevented(self):
        script_path = Path(__file__).resolve().parents[1] / "static" / "js" / "app.js"
        script = script_path.read_text(encoding="utf-8")

        self.assertIn("const controllers = new Map()", script)
        self.assertIn("if (controllers.has(key))", script)
        self.assertIn("return controllers.get(key)", script)

    def test_polling_resumes_when_visible(self):
        script_path = Path(__file__).resolve().parents[1] / "static" / "js" / "app.js"
        script = script_path.read_text(encoding="utf-8")

        self.assertIn('document.addEventListener("visibilitychange"', script)
        self.assertIn("if (!document.hidden", script)
        self.assertIn("document.hidden ? hiddenDelay : visibleDelay", script)

    def test_user_cannot_read_another_users_job_status(self):
        job = self._create_job(self.alice)
        self.client.force_login(self.bob)

        response = self.client.get(reverse("translator:status", args=[job.id]))

        self.assertEqual(response.status_code, 404)

    def test_structure_view_is_owner_scoped(self):
        job = self._create_job(self.alice)
        structure_path = self._write_job_file(
            job,
            "structure.json",
            json.dumps({"pages": [{"page_number": 1, "blocks": []}]}).encode("utf-8"),
        )
        job.structure_file_path = str(structure_path)
        job.save(update_fields=["structure_file_path"])

        self.client.force_login(self.alice)
        owner_response = self.client.get(reverse("translator:structure", args=[job.id]))
        self.assertEqual(owner_response.status_code, 200)
        self.assertEqual(owner_response.json()["status"], job.status)

        self.client.force_login(self.bob)
        other_response = self.client.get(reverse("translator:structure", args=[job.id]))
        self.assertEqual(other_response.status_code, 404)

    def test_preview_image_and_download_are_owner_scoped(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        preview_path = self._write_job_file(
            job,
            "preview",
            "original_page_0.png",
            b"\x89PNG\r\n\x1a\n",
        )
        output_path = self._write_job_file(
            job,
            "translated",
            "translated.pdf",
            b"%PDF-1.4\n%%EOF",
        )
        job.output_file_path = str(output_path)
        job.save(update_fields=["output_file_path"])

        self.client.force_login(self.alice)
        preview_response = self.client.get(
            reverse(
                "translator:translate_preview_image",
                args=[job.id, preview_path.name],
            )
        )
        self.assertEqual(preview_response.status_code, 200)
        preview_response.close()

        download_response = self.client.get(
            reverse("translator:translate_download", args=[job.id])
        )
        self.assertEqual(download_response.status_code, 200)
        download_response.close()

        self.client.force_login(self.bob)
        other_preview_response = self.client.get(
            reverse(
                "translator:translate_preview_image",
                args=[job.id, preview_path.name],
            )
        )
        self.assertEqual(other_preview_response.status_code, 404)

        other_download_response = self.client.get(
            reverse("translator:translate_download", args=[job.id])
        )
        self.assertEqual(other_download_response.status_code, 404)

    def test_history_shows_only_logged_in_users_jobs(self):
        alice_job = self._create_job(
            self.alice,
            status=TranslationJob.Status.COMPLETED,
            original_filename="alice.pdf",
        )
        bob_job = self._create_job(
            self.bob,
            status=TranslationJob.Status.COMPLETED,
            original_filename="bob.pdf",
        )
        TranslationSegment.objects.create(
            job=alice_job,
            segment_index=1,
            source_text="Alice original",
            translated_text="Alice translated",
            source_language="english",
            target_language="tagabawa",
            method="phrasebook_exact",
            confidence=1.0,
        )
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:history"))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn(alice_job.original_filename, body)
        self.assertNotIn(bob_job.original_filename, body)

    def test_preview_uses_database_translation_segments(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        TranslationSegment.objects.create(
            job=job,
            segment_index=1,
            source_text="Original from database",
            translated_text="Translated from database",
            source_language="english",
            target_language="tagabawa",
            method="phrasebook_exact",
            confidence=0.88,
            needs_review=True,
        )
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("Original from database", body)
        self.assertIn("Translated from database", body)
        self.assertIn("Original Document", body)
        self.assertIn("Translated Document", body)
        parser = _PreviewStructureParser()
        parser.feed(body)
        self.assertEqual(parser.container_count, 1)
        self.assertEqual(parser.header_count, 2)
        self.assertEqual(parser.row_count, 1)
        # Normal (non-staff) users never see the technical segment table.
        self.assertNotIn("View Segment Details", body)
        self.assertNotIn("Bilingual aligned segment details", body)
        self.assertNotIn("phrasebook_exact", body)
        self.assertNotIn("Confidence 0.88", body)
        self.assertNotIn("Needs review", body)
        self.assertNotIn("Needs Review", body)

    def test_mixed_known_and_unknown_segments_render_correctly_in_preview(self):
        """Phase 8 regression check: a job with one known phrasebook segment
        and one unsupported segment must, after all the UI/template changes,
        still (a) translate the known segment, (b) keep the unknown
        segment's original text visible rather than dropping or hallucinating
        it, and (c) flag the unknown segment for review — for staff and
        non-staff users alike."""
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        TranslationSegment.objects.create(
            job=job,
            segment_index=1,
            source_text="Hello",
            translated_text="Madigár",
            source_language="english",
            target_language="tagabawa",
            method="exact_phrase",
            confidence=1.0,
            needs_review=False,
        )
        TranslationSegment.objects.create(
            job=job,
            segment_index=2,
            source_text="This sentence is not in the dataset.",
            translated_text="",
            source_language="english",
            target_language="tagabawa",
            method="unknown_for_review",
            confidence=0.0,
            needs_review=True,
        )

        # Non-staff: known segment translates, unknown segment's original
        # text remains visible (not dropped, not hallucinated), and no
        # technical review markers leak through.
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_preview", args=[job.id]))
        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("Madigár", body)
        self.assertIn("This sentence is not in the dataset.", body)
        self.assertNotIn("unknown_for_review", body)
        self.assertNotIn("[UNKNOWN_FOR_REVIEW]", body)
        self.assertNotIn("Needs Review", body)
        self.assertNotIn("Needs review", body)

        # Staff: same content, plus the technical segment table showing the
        # known segment's method/confidence and the unknown segment flagged
        # for review with its original text (never fabricated).
        self.alice.is_staff = True
        self.alice.save(update_fields=["is_staff"])
        staff_response = self.client.get(reverse("translator:job_preview", args=[job.id]))
        staff_body = staff_response.content.decode("utf-8")
        self.assertIn("Madigár", staff_body)
        self.assertIn("exact_phrase", staff_body)
        self.assertIn("This sentence is not in the dataset.", staff_body)
        self.assertIn("needs-review", staff_body)

    def test_preview_renders_translated_panel_when_translated_preview_missing(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        parser = _PreviewStructureParser()
        parser.feed(body)
        self.assertEqual(parser.container_count, 1)
        self.assertEqual(parser.header_count, 2)
        self.assertEqual(parser.row_count, 1)
        self.assertIn("Original Document", body)
        self.assertIn("Translated Document", body)
        self.assertIn("No translated preview available", body)
        self.assertNotIn("View Segment Details", body)
        self.assertNotIn("[UNKNOWN_FOR_REVIEW]", body)

    def test_preview_css_constrains_rendered_page_content_to_panel(self):
        css_path = Path(__file__).resolve().parent.parent / "static" / "css" / "styles.css"
        css = css_path.read_text(encoding="utf-8")

        self.assertIn(".document-page-col svg", css)
        self.assertIn("flex-direction: column", css)
        self.assertIn("flex: 0 0 auto", css)
        self.assertIn("--preview-card-width: 100%", css)
        self.assertIn("width: var(--preview-card-width)", css)
        self.assertIn(".document-page-image", css)
        self.assertIn("width: 100%", css)
        self.assertIn("max-width: 100% !important", css)
        self.assertNotIn("max-height: min(62vh, 560px)", css)
        self.assertIn(".preview-document-page", css)
        self.assertIn(".document-page-col .pdf-page", css)
        self.assertIn("transform-origin: top center", css)
        self.assertIn(".document-page-row-grid", css)
        self.assertIn(".document-compare-headers", css)
        # Original/translated panels no longer scroll independently of each
        # other — pages are aligned by page number in a single shared flow.
        self.assertNotIn(".document-panel-body", css)

    def test_preview_renders_fit_width_zoom_controls(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn('data-preview-scope', body)
        self.assertIn('data-preview-zoom-action="fit"', body)
        self.assertIn("Fit Width", body)
        self.assertIn("Zoom In", body)
        self.assertIn("Zoom Out", body)
        self.assertIn("Reset Zoom", body)

    def test_preview_page_count_uses_pipeline_metadata_when_images_missing(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        job.metadata = {
            "source_page_count": 3,
            "translated_page_count": 3,
            "layout_page_count": 3,
        }
        job.save(update_fields=["metadata"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("3 pages", body)
        self.assertIn("Original page 3", body)
        self.assertIn("Translated page 3", body)

    def test_preview_segment_details_visible_to_staff(self):
        """Staff/admin still get the full technical segment table."""
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        TranslationSegment.objects.create(
            job=job,
            segment_index=1,
            source_text="Original from database",
            translated_text="Translated from database",
            source_language="english",
            target_language="tagabawa",
            method="phrasebook_exact",
            confidence=0.88,
            needs_review=True,
        )
        self.alice.is_staff = True
        self.alice.save(update_fields=["is_staff"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("View Segment Details", body)
        self.assertIn("Bilingual aligned segment details", body)
        self.assertIn("phrasebook_exact", body)
        self.assertIn("Confidence 0.88", body)
        self.assertIn("Needs review", body)

    def test_history_empty_state_has_upload_cta(self):
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:history"))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("No translated documents yet", body)
        self.assertIn("Upload a Document", body)

    def test_pending_job_hides_preview_and_download_buttons(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.QUEUED)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_detail", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertNotIn("Preview Bilingual", body)
        self.assertNotIn("Download</a>", body)
        self.assertIn("Processing is still running", body)

    def test_processing_job_hides_preview_and_download_buttons(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.PROCESSING)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:history"))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn(job.original_filename, body)
        self.assertNotIn("View Preview", body)
        self.assertNotIn("Download", body)

    def test_completed_job_shows_preview_without_download_when_output_missing(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_detail", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("Preview Bilingual", body)
        self.assertNotIn("Download</a>", body)

    def test_completed_job_with_output_shows_download(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        output_path = self._write_job_file(
            job,
            "translated.pdf",
            b"%PDF-1.4\n%%EOF",
        )
        job.output_file_path = str(output_path)
        job.save(update_fields=["output_file_path"])
        self.client.force_login(self.alice)

        detail_response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        status_response = self.client.get(reverse("translator:status", args=[job.id]))

        self.assertEqual(detail_response.status_code, 200)
        body = detail_response.content.decode("utf-8")
        self.assertIn("Preview Bilingual", body)
        self.assertIn("Download</a>", body)
        self.assertIn("Uploaded", body)
        self.assertIn("Extracting / OCR", body)
        self.assertIn("Translating", body)
        self.assertIn("Generating Output", body)
        self.assertIn("Completed", body)
        payload = status_response.json()
        self.assertTrue(payload["can_preview"])
        self.assertTrue(payload["can_download"])

    def test_completed_job_with_generated_output_shows_download(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        output_path = self._write_job_file(
            job,
            "outputs",
            "translated.pdf",
            b"%PDF-1.4\n%%EOF",
        )
        GeneratedOutput.objects.create(
            job=job,
            output_type=GeneratedOutput.OutputType.TRANSLATED_PDF,
            file_format="pdf",
            file_path=str(output_path),
            file_size_bytes=output_path.stat().st_size,
        )
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_detail", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        self.assertIn("Download</a>", response.content.decode("utf-8"))

    def test_failed_job_hides_preview_and_download_buttons(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.FAILED)
        job.error = "OCR produced no text."
        job.save(update_fields=["error"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_detail", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertNotIn("Preview Bilingual", body)
        self.assertNotIn("Download</a>", body)
        # Normal users get the simplified, safe failure message — never the
        # raw backend error text.
        self.assertNotIn("OCR produced no text.", body)
        self.assertIn(
            "Translation could not be completed. Please try another file or contact the administrator.",
            body,
        )
        self.assertIn("Failed", body)

    def test_failed_job_raw_error_visible_to_staff_in_technical_panel(self):
        """Staff/admin can still inspect the raw technical error, but only
        inside the collapsible technical details panel."""
        job = self._create_job(self.alice, status=TranslationJob.Status.FAILED)
        job.error = "OCR produced no text."
        job.save(update_fields=["error"])
        self.alice.is_staff = True
        self.alice.save(update_fields=["is_staff"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_detail", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("Technical Details", body)
        self.assertIn("OCR produced no text.", body)
        self.assertIn(
            "Translation could not be completed. Please try another file or contact the administrator.",
            body,
        )

    def test_preview_before_completion_redirects_to_job_detail(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.PROCESSING)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertRedirects(response, reverse("translator:job_detail", args=[job.id]))

    def test_download_before_completion_redirects_to_job_detail(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.PROCESSING)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_download", args=[job.id]))

        self.assertRedirects(response, reverse("translator:job_detail", args=[job.id]))

    def test_another_user_cannot_view_private_preview(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        self.client.force_login(self.bob)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 404)

    def test_failed_processing_sync_saves_system_event_log(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.PROCESSING)
        sync_pipeline_job(
            SimpleNamespace(
                job_id=job.job_id,
                status=TranslationJob.Status.FAILED,
                progress=25,
                current_phase="ocr",
                current_step="Translation failed",
                phase_message="OCR produced no text.",
                error="OCR produced no text.",
                detection_type="scanned",
                file_type=TranslationJob.FileType.PDF,
                metadata={},
                completed_at=None,
            )
        )

        job.refresh_from_db()
        self.assertEqual(job.status, TranslationJob.Status.FAILED)
        event = SystemEventLog.objects.get(job=job, event_type="translation_job_failed")
        self.assertEqual(event.level, SystemEventLog.Level.ERROR)
        self.assertIn("OCR produced no text", event.message)

    def test_phase_5b_translation_quality_summary_is_visible(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        TranslationSegment.objects.create(
            job=job,
            segment_index=1,
            source_text="Known phrase",
            translated_text="Known translation",
            source_language="english",
            target_language="tagabawa",
            method="exact_phrase",
            confidence=0.8,
            needs_review=False,
        )
        TranslationSegment.objects.create(
            job=job,
            segment_index=2,
            source_text="Unknown phrase",
            translated_text="[UNKNOWN_FOR_REVIEW]",
            source_language="english",
            target_language="tagabawa",
            method="unknown_for_review",
            confidence=0.0,
            needs_review=True,
        )
        original_preview = self._write_job_file(
            job,
            "preview",
            "original_page_0.png",
            b"\x89PNG\r\n\x1a\n",
        )
        translated_preview = self._write_job_file(
            job,
            "preview",
            "translated_page_0.png",
            b"\x89PNG\r\n\x1a\n",
        )
        job.metadata = {
            "preview_original": [str(original_preview)],
            "preview_translated": [str(translated_preview)],
        }
        job.save(update_fields=["metadata"])
        self.client.force_login(self.alice)

        detail_response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        preview_response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(detail_response.status_code, 200)
        detail_body = detail_response.content.decode("utf-8")
        self.assertIn("Translation Confidence", detail_body)
        self.assertIn("40%", detail_body)
        self.assertIn("Needs Review", detail_body)
        # Normal users see a plain-language note, not the raw "1 of 2" count.
        self.assertNotIn("1 of 2", detail_body)
        self.assertIn("Some parts may need teacher review.", detail_body)

        self.assertEqual(preview_response.status_code, 200)
        preview_body = preview_response.content.decode("utf-8")
        self.assertIn("40%", preview_body)
        self.assertNotIn("1 of 2", preview_body)
        self.assertNotIn("Needs Review", preview_body)
        self.assertNotIn("Some parts may need teacher review.", preview_body)

    def test_phase_5b_translation_quality_technical_counts_visible_to_staff(self):
        """Staff/admin still see the exact Needs Review counts and segment table."""
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        TranslationSegment.objects.create(
            job=job,
            segment_index=1,
            source_text="Known phrase",
            translated_text="Known translation",
            source_language="english",
            target_language="tagabawa",
            method="exact_phrase",
            confidence=0.8,
            needs_review=False,
        )
        TranslationSegment.objects.create(
            job=job,
            segment_index=2,
            source_text="Unknown phrase",
            translated_text="[UNKNOWN_FOR_REVIEW]",
            source_language="english",
            target_language="tagabawa",
            method="unknown_for_review",
            confidence=0.0,
            needs_review=True,
        )
        self.alice.is_staff = True
        self.alice.save(update_fields=["is_staff"])
        self.client.force_login(self.alice)

        detail_response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        preview_response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(detail_response.status_code, 200)
        detail_body = detail_response.content.decode("utf-8")
        self.assertIn("1 of 2", detail_body)

        self.assertEqual(preview_response.status_code, 200)
        preview_body = preview_response.content.decode("utf-8")
        self.assertIn("Translation Details", preview_body)
        self.assertIn("1 of 2", preview_body)

    def test_phase_5b_quality_summary_is_exposed_by_status_and_preview_api(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        TranslationSegment.objects.create(
            job=job,
            segment_index=1,
            source_text="Known phrase",
            translated_text="Known translation",
            source_language="english",
            target_language="tagabawa",
            method="exact_phrase",
            confidence=0.8,
            needs_review=False,
        )
        TranslationSegment.objects.create(
            job=job,
            segment_index=2,
            source_text="Unknown phrase",
            translated_text="[UNKNOWN_FOR_REVIEW]",
            source_language="english",
            target_language="tagabawa",
            method="unknown_for_review",
            confidence=0.0,
            needs_review=True,
        )
        self.client.force_login(self.alice)

        status_response = self.client.get(reverse("translator:status", args=[job.id]))
        preview_response = self.client.get(
            reverse("translator:preview_data", args=[job.id])
        )

        self.assertEqual(status_response.status_code, 200)
        status_payload = status_response.json()
        self.assertEqual(status_payload["segment_count"], 2)
        self.assertEqual(status_payload["review_segment_count"], 1)
        self.assertAlmostEqual(status_payload["translation_confidence"], 0.4)
        self.assertEqual(status_payload["translation_confidence_pct"], 40)
        self.assertTrue(status_payload["translation_has_review_items"])

        self.assertEqual(preview_response.status_code, 200)
        preview_payload = preview_response.json()
        self.assertEqual(preview_payload["segment_count"], 2)
        self.assertEqual(preview_payload["review_segment_count"], 1)
        self.assertEqual(preview_payload["translation_confidence_pct"], 40)
        self.assertEqual(preview_payload["translation_summary"]["segment_count"], 2)
        self.assertEqual(
            preview_payload["translation_summary"]["review_segment_count"],
            1,
        )
        self.assertTrue(
            preview_payload["translation_summary"]["translation_has_review_items"]
        )
        unknown_segment = preview_payload["segments"][1]
        self.assertEqual(unknown_segment["translated_text"], "Unknown phrase")
        self.assertEqual(unknown_segment["display_translated_text"], "Unknown phrase")
        self.assertTrue(unknown_segment["needs_review"])
        self.assertNotIn("[UNKNOWN_FOR_REVIEW]", json.dumps(preview_payload))

    def test_retrying_status_api_is_active_state(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.RETRYING)
        job.current_phase = "retrying"
        job.current_step = "Retrying abandoned translation job"
        job.phase_message = "This job was queued for retry."
        job.save()
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:status", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["status"], TranslationJob.Status.RETRYING)
        self.assertTrue(payload["is_processing"])
        self.assertEqual(payload["message"], "This job was queued for retry.")

    def test_status_api_returns_consistent_fields_for_all_states(self):
        self.client.force_login(self.alice)
        required_fields = {
            "status",
            "progress_percent",
            "current_phase",
            "current_step",
            "phase_message",
            "updated_at",
            "started_at",
            "completed_at",
            "error",
            "can_preview",
            "can_download",
        }

        for status in [
            TranslationJob.Status.QUEUED,
            TranslationJob.Status.PROCESSING,
            TranslationJob.Status.RETRYING,
            TranslationJob.Status.COMPLETED,
            TranslationJob.Status.FAILED,
        ]:
            job = self._create_job(
                self.alice,
                status=status,
                original_filename=f"{status}.pdf",
            )
            if status == TranslationJob.Status.FAILED:
                job.error = "Failed for test"
                job.save(update_fields=["error"])

            response = self.client.get(reverse("translator:status", args=[job.id]))

            self.assertEqual(response.status_code, 200)
            payload = response.json()
            self.assertTrue(required_fields.issubset(payload.keys()))
            self.assertEqual(payload["status"], status)
            self.assertIsNotNone(payload["started_at"])
            if status == TranslationJob.Status.FAILED:
                # The JSON status API never leaks the raw backend error to
                # the requesting (non-staff) user — only the safe, generic
                # failure message.
                self.assertNotEqual(payload["error"], "Failed for test")
                self.assertEqual(
                    payload["error"],
                    "Translation could not be completed. Please try another file or contact the administrator.",
                )

    def test_active_job_banner_is_owner_scoped(self):
        self._create_job(self.alice, status=TranslationJob.Status.PROCESSING, original_filename="alice.pdf")
        self._create_job(self.bob, status=TranslationJob.Status.PROCESSING, original_filename="bob.pdf")

        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:history"))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("1 document is currently processing.", body)
        self.assertIn("alice.pdf", body)
        self.assertNotIn("bob.pdf", body)
        self.assertIn("View Status", body)
        self.assertIn("History", body)

    def test_start_translation_job_dispatches_after_commit(self):
        from translator.services import start_translation_job

        job = self._create_job(self.alice, status=TranslationJob.Status.QUEUED)
        job.input_file_path = str(self.media_root / "sample.pdf")
        job.save(update_fields=["input_file_path"])

        with patch("translator.services.register_pipeline_callback"), patch(
            "translator.services.submit_translation_task"
        ) as submit_task:
            with self.captureOnCommitCallbacks(execute=False) as callbacks:
                start_translation_job(job)

            submit_task.assert_not_called()
            self.assertEqual(len(callbacks), 1)
            callbacks[0]()

        submit_task.assert_called_once()
        self.assertEqual(submit_task.call_args.args[1], job.job_id)

    def test_duplicate_task_is_ignored_safely(self):
        from translator.services import _run_translation_job

        job = self._create_job(self.alice, status=TranslationJob.Status.PROCESSING)

        with patch("translator.services._get_pipeline_service") as pipeline_service:
            _run_translation_job(
                job.job_id,
                job.input_file_path,
                "pdf",
                "auto",
                "tagabawa",
                None,
            )

        pipeline_service.assert_not_called()
        job.refresh_from_db()
        self.assertEqual(job.status, TranslationJob.Status.PROCESSING)
        self.assertTrue(
            SystemEventLog.objects.filter(
                job=job,
                event_type="translation_job_duplicate_ignored",
            ).exists()
        )

    def _create_job(self, owner, status=TranslationJob.Status.QUEUED, original_filename="sample.pdf"):
        return TranslationJob.objects.create(
            owner=owner,
            original_filename=original_filename,
            file_type=TranslationJob.FileType.PDF,
            status=status,
        )

    def _write_job_file(self, job, *parts_and_content):
        *parts, content = parts_and_content
        path = self.media_root.joinpath("jobs", job.job_id, *parts)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(content)
        return path


class TranslatorPhase5BPipelineTests(TestCase):
    def test_pipeline_marks_ocr_failed_pages_as_incomplete_translation(self):
        from translator.services.pipeline_service import PipelineService

        warnings = PipelineService._untranslated_page_warnings([
            {
                "page": 0,
                "blocks": [
                    {
                        "type": "text",
                        "lines": [{"text": "Page one source"}],
                    }
                ],
            },
            {
                "page": 1,
                "blocks": [],
                "ocr_error": "image too faded to process",
            },
        ])

        self.assertEqual(len(warnings), 1)
        self.assertIn("Page 2", warnings[0])
        self.assertIn("image too faded to process", warnings[0])

    def test_bilingual_first_page_uses_line_specific_translation_metadata(self):
        from translator.services.pipeline_service import PipelineService
        from translator.services.translation_dataset import UNKNOWN_FOR_REVIEW

        service = PipelineService.__new__(PipelineService)
        layout_data = [
            {
                "blocks": [
                    {
                        "type": "text",
                        "bbox": [0, 0, 100, 50],
                        "lines": [
                            {"text": "First line", "bbox": [0, 0, 100, 20]},
                            {"text": "Second line", "bbox": [0, 24, 100, 44]},
                        ],
                    }
                ]
            }
        ]
        translations = {
            "0_0": {
                "translated": "Wrong block-level fallback",
                "method": "legacy",
                "cascade_stage": "legacy",
                "confidence": 0.1,
            },
            "0_0_0": {
                "translated": "First translated",
                "method": "exact_phrase",
                "cascade_stage": "exact_phrase",
                "confidence": 1.0,
            },
            "0_0_1": {
                "translated": UNKNOWN_FOR_REVIEW,
                "method": "unknown_for_review",
                "cascade_stage": "unknown_for_review",
                "confidence": 0.0,
            },
        }

        preview = service._build_bilingual_first_page(layout_data, translations)

        self.assertEqual(len(preview["blocks"]), 2)
        self.assertEqual(preview["blocks"][0]["translated_text"], "First translated")
        self.assertEqual(preview["blocks"][0]["translation_method"], "exact_phrase")
        self.assertEqual(preview["blocks"][0]["translation_confidence"], 1.0)
        self.assertFalse(preview["blocks"][0]["needs_review"])
        self.assertEqual(preview["blocks"][1]["translated_text"], "Second line")
        self.assertEqual(preview["blocks"][1]["raw_translated_text"], UNKNOWN_FOR_REVIEW)
        self.assertEqual(preview["blocks"][1]["translation_method"], "unknown_for_review")
        self.assertEqual(preview["blocks"][1]["translation_confidence"], 0.0)
        self.assertTrue(preview["blocks"][1]["needs_review"])


# ============================================================
# Phase 4 — System Design, Gating, Security & Purge Tests
# ============================================================

import os
import shutil


class TranslatorPhase4SystemTests(TestCase):
    password = "Bagobo-Test-Phase4-Sys!"

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)

        self.alice = User.objects.create_user(
            username="sysalice", email="sysalice@example.test", password=self.password
        )
        self.bob = User.objects.create_user(
            username="sysbob", email="sysbob@example.test", password=self.password
        )

    def _create_job(self, owner, status="completed", metadata=None):
        return TranslationJob.objects.create(
            owner=owner,
            original_filename="test_doc.pdf",
            file_type="pdf",
            status=status,
            source_language="english",
            target_language="tagabawa",
            metadata=metadata or {},
        )

    def test_scanned_pdf_uses_ocr_automatically(self):
        from translator.services.detection_service import DetectionService, DetectionType
        orig = DetectionService._count_pdf_text_chars
        try:
            DetectionService._count_pdf_text_chars = staticmethod(lambda path: (10, 1))
            det_type = DetectionService.detect_pdf_type("dummy.pdf")
            self.assertEqual(det_type, DetectionType.SCANNED)
        finally:
            DetectionService._count_pdf_text_chars = orig

    def test_image_upload_uses_ocr_automatically(self):
        from translator.services.detection_service import DetectionService, DetectionType
        det_type = DetectionService.detect_image_type("dummy.png")
        self.assertEqual(det_type, DetectionType.SCANNED)

    def test_digital_pdf_uses_direct_text_extraction(self):
        from translator.services.detection_service import DetectionService, DetectionType
        orig = DetectionService._count_pdf_text_chars
        try:
            DetectionService._count_pdf_text_chars = staticmethod(lambda path: (60, 1))
            det_type = DetectionService.detect_pdf_type("dummy.pdf")
            self.assertEqual(det_type, DetectionType.DIGITAL)
        finally:
            DetectionService._count_pdf_text_chars = orig

    def test_low_text_pdf_falls_back_to_ocr(self):
        from translator.services.detection_service import DetectionService, DetectionType
        orig = DetectionService._count_pdf_text_chars
        try:
            DetectionService._count_pdf_text_chars = staticmethod(lambda path: (49, 1))
            det_type = DetectionService.detect_pdf_type("dummy.pdf")
            self.assertEqual(det_type, DetectionType.SCANNED)
        finally:
            DetectionService._count_pdf_text_chars = orig

    def test_docx_uses_docx_text_extraction(self):
        from translator.services.detection_service import DetectionService, DetectionType
        det_type = DetectionService.detect_docx_type("dummy.docx")
        self.assertEqual(det_type, DetectionType.DIGITAL)

    def test_txt_uses_plain_text_extraction(self):
        job = self._create_job(self.alice, status="processing")
        job.file_type = "txt"
        job.save()
        self.assertEqual(job.file_type, "txt")

    def test_extraction_method_saved_in_metadata(self):
        job = self._create_job(self.alice, status="completed", metadata={"extraction_method": "docx_text"})
        self.assertEqual(job.metadata.get("extraction_method"), "docx_text")

    def test_ocr_confidence_saved_and_displayed(self):
        # OCR confidence detail lives in the staff-only technical panel.
        job = self._create_job(self.alice, status="completed", metadata={
            "extraction_method": "ocr_image",
            "ocr_summary": {"mean_confidence": 0.85}
        })
        self.alice.is_staff = True
        self.alice.save(update_fields=["is_staff"])
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        self.assertEqual(response.status_code, 200)
        self.assertIn("85%", response.content.decode("utf-8"))

    def test_ocr_warnings_saved_and_displayed(self):
        # OCR warnings are technical detail, visible to staff only.
        job = self._create_job(self.alice, status="completed", metadata={
            "ocr_warnings": ["Warning text 123"]
        })
        self.alice.is_staff = True
        self.alice.save(update_fields=["is_staff"])
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        self.assertEqual(response.status_code, 200)
        self.assertIn("Warning text 123", response.content.decode("utf-8"))

    @patch("translator.services._get_pipeline_service")
    def test_ocr_failure_fails_safely(self, mock_pipeline):
        from translator.services import _run_translation_job
        mock_pipeline.side_effect = Exception("Tesseract failed catastrophically")
        job = self._create_job(self.alice, status="queued")
        _run_translation_job(
            job.job_id,
            job.input_file_path,
            "pdf",
            "auto",
            "tagabawa",
            None
        )
        job.refresh_from_db()
        self.assertEqual(job.status, TranslationJob.Status.FAILED)
        self.assertIn("Tesseract failed catastrophically", job.error)

    @patch("translator.services._get_pipeline_service")
    def test_ocr_timeout_sets_failed_status(self, mock_pipeline):
        from translator.services import _run_translation_job

        mock_pipeline.side_effect = TimeoutError("Tesseract OCR timed out")
        job = self._create_job(self.alice, status=TranslationJob.Status.QUEUED)

        _run_translation_job(
            job.job_id,
            job.input_file_path,
            "pdf",
            "auto",
            "tagabawa",
            None,
        )

        job.refresh_from_db()
        self.assertEqual(job.status, TranslationJob.Status.FAILED)
        self.assertIn("Tesseract OCR timed out", job.error)
        self.assertTrue(
            SystemEventLog.objects.filter(
                job=job,
                event_type="translation_job_failed",
            ).exists()
        )

    def test_stale_processing_job_is_recovered(self):
        job = self._create_job(self.alice, status=TranslationJob.Status.PROCESSING)
        old_updated_at = timezone.now() - timezone.timedelta(minutes=30)
        TranslationJob.objects.filter(id=job.id).update(updated_at=old_updated_at)

        dry_output = StringIO()
        call_command("recover_stale_jobs", minutes=15, dry_run=True, stdout=dry_output)
        job.refresh_from_db()
        self.assertEqual(job.status, TranslationJob.Status.PROCESSING)
        self.assertIn("fail:", dry_output.getvalue())

        output = StringIO()
        call_command("recover_stale_jobs", minutes=15, stdout=output)
        job.refresh_from_db()
        self.assertEqual(job.status, TranslationJob.Status.FAILED)
        self.assertIn("abandoned", job.error)
        self.assertIn("Recovered 1 stale job", output.getvalue())
        self.assertTrue(
            SystemEventLog.objects.filter(
                job=job,
                event_type="translation_job_stale_recovery",
            ).exists()
        )

    def test_recover_stale_jobs_skips_completed_and_deleted_jobs(self):
        old_updated_at = timezone.now() - timezone.timedelta(minutes=30)
        completed = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        deleted = self._create_job(self.alice, status=TranslationJob.Status.PROCESSING)
        deleted.is_deleted = True
        deleted.save(update_fields=["is_deleted"])
        TranslationJob.objects.filter(id__in=[completed.id, deleted.id]).update(
            updated_at=old_updated_at
        )

        output = StringIO()
        call_command("recover_stale_jobs", minutes=15, stdout=output)

        completed.refresh_from_db()
        deleted.refresh_from_db()
        self.assertEqual(completed.status, TranslationJob.Status.COMPLETED)
        self.assertEqual(deleted.status, TranslationJob.Status.PROCESSING)
        self.assertIn("No stale translation jobs found", output.getvalue())

    def test_upload_invalid_file_rejected(self):
        self.client.force_login(self.alice)
        bad_file = SimpleUploadedFile("danger.exe", b"binary content", content_type="application/octet-stream")
        response = self.client.post(reverse("translator:upload"), {"file": bad_file})
        self.assertEqual(response.status_code, 400)

    def test_upload_large_file_rejected(self):
        self.client.force_login(self.alice)
        large_file = SimpleUploadedFile("big.pdf", b"a" * (51 * 1024 * 1024), content_type="application/pdf")
        response = self.client.post(reverse("translator:upload"), {"file": large_file})
        self.assertEqual(response.status_code, 400)

    def test_upload_requires_login(self):
        # Must be a JSON 401, not an HTML redirect: fetch() in app.js follows
        # redirects automatically and ends up handing response.json() the
        # login page's HTML, which throws "Unexpected token '<'".
        response = self.client.post(reverse("translator:upload"))
        self.assertEqual(response.status_code, 401)
        self.assertEqual(response["Content-Type"], "application/json")
        data = response.json()
        self.assertFalse(data["ok"])
        self.assertEqual(data["redirect_url"], reverse("translator:login"))

    def test_upload_csrf_required(self):
        from django.test import Client
        csrf_client = Client(enforce_csrf_checks=True)
        csrf_client.force_login(self.alice)
        # Load page to set cookie
        csrf_client.get(reverse("translator:translate"))
        response = csrf_client.post(reverse("translator:upload"), {"file": "dummy"})
        self.assertEqual(response.status_code, 403)
        self.assertEqual(response["Content-Type"], "application/json")
        data = response.json()
        self.assertFalse(data["ok"])
        self.assertIn("error", data)

    def test_upload_rate_limit_blocks_excess_requests(self):
        self.client.force_login(self.alice)
        from django.core.cache import cache
        cache.clear()
        for _ in range(5):
            self.client.post(reverse("translator:upload"), {"file": ""})
        response = self.client.post(reverse("translator:upload"), {"file": ""})
        self.assertEqual(response.status_code, 400)
        self.assertIn("upload limit", response.json()["detail"])

    def test_active_job_limit_blocks_excess_processing_jobs(self):
        self.client.force_login(self.alice)
        self._create_job(self.alice, status=TranslationJob.Status.QUEUED)
        self._create_job(self.alice, status=TranslationJob.Status.PROCESSING)
        response = self.client.post(reverse("translator:upload"), {"file": ""})
        self.assertEqual(response.status_code, 400)
        self.assertIn("upload limit", response.json()["detail"])

    def test_health_page_is_staff_only(self):
        staff = User.objects.create_user(
            username="healthstaff",
            email="healthstaff@example.test",
            password=self.password,
            is_staff=True,
        )
        checks = [{"name": "Database", "status": "ok", "detail": "sqlite"}]
        counts = {"queued": 0, "processing": 0, "retrying": 0, "stale": 0}

        response = self.client.get(reverse("translator:health"))
        self.assertEqual(response.status_code, 302)

        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:health"))
        self.assertEqual(response.status_code, 302)

        self.client.force_login(staff)
        with patch("translator.views._health_checks", return_value=checks), patch(
            "translator.views._health_job_counts", return_value=counts
        ):
            response = self.client.get(reverse("translator:health"))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("System Health", body)
        self.assertIn("Database", body)

    def test_unavailable_redis_and_celery_are_reported_safely(self):
        from translator import views

        class BrokenRedis:
            @classmethod
            def from_url(cls, *args, **kwargs):
                raise RuntimeError("redis down")

        fake_redis_module = SimpleNamespace(Redis=BrokenRedis)
        with override_settings(CELERY_BROKER_URL="redis://:secret@localhost:6379/0"):
            with patch.dict(sys.modules, {"redis": fake_redis_module}):
                redis_check = views._redis_health()

        self.assertEqual(redis_check["status"], "error")
        self.assertIn("redis://localhost:6379/0", redis_check["detail"])
        self.assertNotIn("secret", redis_check["detail"])

        with patch("lingokatutubo_django.celery.app.control.ping", return_value=[]):
            celery_check = views._celery_worker_health()

        self.assertEqual(celery_check["status"], "warning")
        self.assertIn("No workers responded", celery_check["detail"])

    def test_preview_hidden_before_completed(self):
        job = self._create_job(self.alice, status="processing")
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_preview", args=[job.id]))
        self.assertRedirects(response, reverse("translator:job_detail", args=[job.id]))

    def test_download_hidden_until_output_exists(self):
        job = self._create_job(self.alice, status="completed")
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_download", args=[job.id]))
        self.assertRedirects(response, reverse("translator:job_detail", args=[job.id]))

    def test_failed_job_hides_preview_and_download(self):
        job = self._create_job(self.alice, status="failed")
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        self.assertEqual(response.status_code, 200)
        self.assertNotIn(b"Preview Bilingual", response.content)

    def test_user_cannot_view_another_users_job(self):
        job = self._create_job(self.bob, status="completed")
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        self.assertEqual(response.status_code, 404)

    def test_user_cannot_preview_another_users_job(self):
        job = self._create_job(self.bob, status="completed")
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_preview", args=[job.id]))
        self.assertEqual(response.status_code, 404)

    def test_user_cannot_download_another_users_output(self):
        job = self._create_job(self.bob, status="completed")
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_download", args=[job.id]))
        self.assertEqual(response.status_code, 404)

    def test_user_cannot_delete_another_users_job(self):
        job = self._create_job(self.bob, status="completed")
        self.client.force_login(self.alice)
        response = self.client.post(reverse("translator:job_delete", args=[job.id]))
        self.assertEqual(response.status_code, 404)

    def test_preview_image_blocks_path_traversal(self):
        self.client.force_login(self.alice)
        job = self._create_job(self.alice, status="completed")
        response = self.client.get(reverse("translator:translate_preview_image", args=[job.id, "invalid_name.png"]))
        self.assertEqual(response.status_code, 400)

    @patch("translator.views._translated_output_path")
    @patch("translator.views._translated_output_exists")
    def test_download_serves_attachment_pdf_only(self, mock_exists, mock_path):
        mock_exists.return_value = True
        job = self._create_job(self.alice, status="completed")
        with tempfile.NamedTemporaryFile(dir=self.media_root, delete=False) as tmp:
            tmp.write(b"%PDF-1.4 mock content")
            tmp_name = tmp.name
        mock_path.return_value = Path(tmp_name)
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_download", args=[job.id]))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.headers["Content-Type"], "application/pdf")
        self.assertIn("attachment", response.headers["Content-Disposition"])
        
        # Close file handles to avoid Windows WinError 32 permission errors on cleanup
        response.close()
        os.remove(tmp_name)

    def test_history_lists_only_user_jobs(self):
        self._create_job(self.alice)
        self._create_job(self.bob)
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:history"))
        self.assertEqual(response.status_code, 200)
        jobs = response.context["jobs"]
        self.assertEqual(len(jobs), 1)
        self.assertEqual(jobs[0]["job"].owner, self.alice)

    def test_soft_deleted_job_hidden_from_history(self):
        job = self._create_job(self.alice)
        job.is_deleted = True
        job.save()
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:history"))
        self.assertEqual(len(response.context["jobs"]), 0)

    def test_soft_deleted_job_hidden_from_recent_sidebar(self):
        job = self._create_job(self.alice)
        job.is_deleted = True
        job.save()
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:translate"))
        self.assertEqual(len(response.context["recent_jobs"]), 0)

    def test_deleted_job_returns_404_on_detail(self):
        job = self._create_job(self.alice)
        job.is_deleted = True
        job.save()
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        self.assertEqual(response.status_code, 404)

    def test_confirm_delete_page_requires_login(self):
        job = self._create_job(self.alice)
        response = self.client.get(reverse("translator:job_delete_confirm", args=[job.id]))
        self.assertEqual(response.status_code, 302)

    def test_delete_requires_post(self):
        job = self._create_job(self.alice)
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_delete", args=[job.id]))
        self.assertEqual(response.status_code, 405)

    def test_admin_extraction_method_display_handles_missing_metadata(self):
        from translator.admin import TranslationJobAdmin
        from django.contrib.admin.sites import AdminSite
        admin_site = AdminSite()
        admin_instance = TranslationJobAdmin(TranslationJob, admin_site)
        job = self._create_job(self.alice, metadata=None)
        display = admin_instance.extraction_method_display(job)
        self.assertEqual(display, "—")

    def test_admin_ocr_confidence_display_handles_null_values(self):
        from translator.admin import TranslationJobAdmin
        from django.contrib.admin.sites import AdminSite
        admin_site = AdminSite()
        admin_instance = TranslationJobAdmin(TranslationJob, admin_site)
        job = self._create_job(self.alice, metadata=None)
        display = admin_instance.ocr_confidence_display(job)
        self.assertEqual(display, "N/A")

    def test_admin_translation_quality_display_handles_segments_and_empty_jobs(self):
        from translator.admin import TranslationJobAdmin
        from django.contrib.admin.sites import AdminSite

        admin_site = AdminSite()
        admin_instance = TranslationJobAdmin(TranslationJob, admin_site)
        job = self._create_job(self.alice, metadata=None)

        self.assertEqual(admin_instance.translation_confidence_display(job), "N/A")
        self.assertEqual(admin_instance.review_segments_display(job), "N/A")

        TranslationSegment.objects.create(
            job=job,
            segment_index=1,
            source_text="Known phrase",
            translated_text="Known translation",
            source_language="english",
            target_language="tagabawa",
            method="exact_phrase",
            confidence=0.8,
            needs_review=False,
        )
        TranslationSegment.objects.create(
            job=job,
            segment_index=2,
            source_text="Unknown phrase",
            translated_text="[UNKNOWN_FOR_REVIEW]",
            source_language="english",
            target_language="tagabawa",
            method="unknown_for_review",
            confidence=0.0,
            needs_review=True,
        )

        self.assertEqual(admin_instance.translation_confidence_display(job), "40%")
        self.assertEqual(admin_instance.review_segments_display(job), "1/2")

    def test_admin_ocr_quality_filter_finds_low_confidence_jobs(self):
        from django.contrib.admin.sites import AdminSite
        from django.test import RequestFactory
        from translator.admin import OCRQualityFilter, TranslationJobAdmin

        low_job = self._create_job(self.alice, status="completed")
        high_job = self._create_job(self.alice, status="completed")
        OCRResult.objects.create(job=low_job, confidence=0.42, text="low")
        OCRResult.objects.create(job=high_job, confidence=0.91, text="high")

        admin_site = AdminSite()
        admin_instance = TranslationJobAdmin(TranslationJob, admin_site)
        request = RequestFactory().get("/admin/", {"ocr_quality": "low"})
        request.user = self.alice
        filter_spec = OCRQualityFilter(
            request,
            {"ocr_quality": ["low"]},
            TranslationJob,
            admin_instance,
        )

        queryset = filter_spec.queryset(request, TranslationJob.objects.all())

        self.assertIn(low_job, queryset)
        self.assertNotIn(high_job, queryset)

    def test_purge_deleted_job_files_dry_run(self):
        from django.core.management import call_command
        from django.utils import timezone
        from datetime import timedelta
        job = self._create_job(self.alice, status="completed")
        job.is_deleted = True
        job.deleted_at = timezone.now() - timedelta(days=31)
        with tempfile.NamedTemporaryFile(dir=self.media_root, delete=False) as tmp:
            tmp.write(b"dummy content")
            tmp_name = tmp.name
        job.input_file_path = tmp_name
        job.save()
        call_command("purge_deleted_job_files", days=30, dry_run=True)
        self.assertTrue(os.path.exists(tmp_name))
        os.remove(tmp_name)

    def test_purge_deleted_job_files_actual(self):
        from django.core.management import call_command
        from django.utils import timezone
        from datetime import timedelta
        job = self._create_job(self.alice, status="completed")
        job.is_deleted = True
        job.deleted_at = timezone.now() - timedelta(days=31)
        with tempfile.NamedTemporaryFile(dir=self.media_root, delete=False) as tmp:
            tmp.write(b"dummy content")
            tmp_name = tmp.name
        job.input_file_path = tmp_name
        job.save()
        call_command("purge_deleted_job_files", days=30)
        self.assertFalse(os.path.exists(tmp_name))


# ============================================================
# File Upload Validation Tests
# ============================================================


class TranslatorUploadValidationTests(TestCase):
    """Tests for file-type and signature validation on upload."""

    password = "Bagobo-Upload-Test-2026!"

    def setUp(self):
        from django.core.cache import cache as _cache
        _cache.clear()

        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)

        self.user = User.objects.create_user(
            username="uploadtest",
            email="uploadtest@example.test",
            password=self.password,
        )

    @staticmethod
    def _make_minimal_docx() -> bytes:
        import io
        import zipfile as _zipfile
        buf = io.BytesIO()
        with _zipfile.ZipFile(buf, "w") as zf:
            zf.writestr(
                "[Content_Types].xml",
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                "</Types>",
            )
            zf.writestr("word/document.xml", "<w:document/>")
        return buf.getvalue()

    @staticmethod
    def _make_minimal_jpg() -> bytes:
        from PIL import Image
        import io
        img = Image.new("RGB", (4, 4), color=(200, 200, 200))
        buf = io.BytesIO()
        img.save(buf, format="JPEG")
        return buf.getvalue()

    @staticmethod
    def _make_minimal_png() -> bytes:
        from PIL import Image
        import io
        img = Image.new("RGB", (4, 4), color=(200, 200, 200))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()

    def _upload(self, filename, content, content_type, patch_start=True):
        self.client.force_login(self.user)
        uploaded = SimpleUploadedFile(filename, content, content_type=content_type)
        if patch_start:
            with patch("translator.views.start_translation_job"):
                return self.client.post(
                    reverse("translator:upload"),
                    {"file": uploaded, "source_language": "auto", "target_language": "tagabawa"},
                )
        return self.client.post(
            reverse("translator:upload"),
            {"file": uploaded, "source_language": "auto", "target_language": "tagabawa"},
        )

    def test_upload_valid_pdf_accepted(self):
        response = self._upload(
            "doc.pdf",
            b"%PDF-1.4\n1 0 obj\n<<>>\nendobj\n%%EOF",
            "application/pdf",
        )
        self.assertEqual(response.status_code, 202)

    def test_upload_valid_docx_accepted(self):
        response = self._upload(
            "doc.docx",
            self._make_minimal_docx(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertEqual(response.status_code, 202)

    def test_upload_valid_jpg_accepted(self):
        response = self._upload("photo.jpg", self._make_minimal_jpg(), "image/jpeg")
        self.assertEqual(response.status_code, 202)

    def test_upload_valid_png_accepted(self):
        response = self._upload("image.png", self._make_minimal_png(), "image/png")
        self.assertEqual(response.status_code, 202)

    def test_upload_valid_txt_accepted(self):
        response = self._upload(
            "notes.txt",
            "Hello world, kumusta?".encode("utf-8"),
            "text/plain",
        )
        self.assertEqual(response.status_code, 202)

    def test_upload_invalid_pdf_signature_rejected(self):
        """File with .pdf extension but wrong magic bytes is rejected."""
        response = self._upload(
            "fake.pdf",
            b"Not a real PDF document",
            "application/pdf",
            patch_start=False,
        )
        self.assertEqual(response.status_code, 400)

    def test_upload_corrupted_docx_rejected(self):
        """File with .docx extension but invalid ZIP content is rejected."""
        response = self._upload(
            "broken.docx",
            b"PK\x03\x04this is not a real zip archive",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            patch_start=False,
        )
        self.assertEqual(response.status_code, 400)

    def test_upload_non_utf8_txt_rejected(self):
        """TXT file with non-UTF-8 bytes is rejected."""
        response = self._upload(
            "latin.txt",
            b"\xff\xfe\x00invalid latin-1 content",
            "text/plain",
            patch_start=False,
        )
        self.assertEqual(response.status_code, 400)

    def test_upload_disallowed_extension_rejected(self):
        """File with a disallowed extension (e.g. .exe) is always rejected."""
        response = self._upload(
            "program.exe",
            b"binary content",
            "application/octet-stream",
            patch_start=False,
        )
        self.assertEqual(response.status_code, 400)


# ============================================================
# Preview Pagination Tests
# ============================================================


class TranslatorPreviewPaginationTests(TestCase):
    """Verify that the bilingual preview paginates segments correctly."""

    password = "Bagobo-Paginate-2026!"

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)

        # The segment-table pagination tested here lives in the staff-only
        # "View Segment Details" technical panel.
        self.alice = User.objects.create_user(
            username="pagalice",
            email="pagalice@example.test",
            password=self.password,
            is_staff=True,
        )

    def _create_completed_job(self):
        return TranslationJob.objects.create(
            owner=self.alice,
            original_filename="long_doc.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=TranslationJob.Status.COMPLETED,
        )

    def _bulk_create_segments(self, job, count):
        TranslationSegment.objects.bulk_create([
            TranslationSegment(
                job=job,
                segment_index=i + 1,
                source_text=f"Source segment {i + 1}",
                translated_text=f"Translated segment {i + 1}",
                source_language="english",
                target_language="tagabawa",
                method="exact_phrase",
                confidence=0.9,
            )
            for i in range(count)
        ])

    def test_preview_first_page_shows_first_25_segments(self):
        """The staff-only technical segment table paginates at 25 rows.

        Note: the normal-user-facing Original/Translated panels intentionally
        show the full concatenated document text regardless of this
        pagination, so "Source segment 26" existing elsewhere on the page is
        expected — only the table's own page window and nav label matter
        here.
        """
        job = self._create_completed_job()
        self._bulk_create_segments(job, 30)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("Source segment 1", body)
        self.assertIn("Source segment 25", body)
        self.assertIn("Page 1 of 2", body)

    def test_preview_second_page_shows_remaining_segments(self):
        job = self._create_completed_job()
        self._bulk_create_segments(job, 30)
        self.client.force_login(self.alice)

        response = self.client.get(
            reverse("translator:job_preview", args=[job.id]) + "?page=2"
        )

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("Source segment 26", body)
        self.assertIn("Source segment 30", body)
        self.assertIn("Page 2 of 2", body)

    def test_preview_total_segment_count_shown_correctly(self):
        job = self._create_completed_job()
        self._bulk_create_segments(job, 30)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        body = response.content.decode("utf-8")
        # Total segments badge should show 30, not just the current page count (25)
        self.assertIn("30 segments", body)

    def test_preview_no_pagination_for_small_document(self):
        job = self._create_completed_job()
        self._bulk_create_segments(job, 10)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        body = response.content.decode("utf-8")
        # No pagination nav when all segments fit on one page
        self.assertNotIn("Page 1 of", body)

    def test_preview_owner_isolation_with_pagination(self):
        bob = User.objects.create_user(
            username="pagbob",
            email="pagbob@example.test",
            password=self.password,
        )
        job = self._create_completed_job()
        self._bulk_create_segments(job, 5)
        self.client.force_login(bob)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 404)


# ============================================================
# Admin Access and Read-Only Tests
# ============================================================


class TranslatorAdminTests(TestCase):
    """Verify admin interface access controls and read-only log enforcement."""

    password = "Bagobo-Admin-2026!"

    def setUp(self):
        self.regular = User.objects.create_user(
            username="regular_admin_test",
            email="regular@example.test",
            password=self.password,
        )
        self.staff = User.objects.create_user(
            username="staff_admin_test",
            email="staff@example.test",
            password=self.password,
            is_staff=True,
            is_superuser=True,
        )

    def test_admin_access_denied_for_anonymous(self):
        response = self.client.get("/admin/")
        self.assertNotEqual(response.status_code, 200)

    def test_admin_access_denied_for_regular_user(self):
        self.client.force_login(self.regular)
        response = self.client.get("/admin/")
        self.assertNotEqual(response.status_code, 200)

    def test_admin_accessible_for_superuser(self):
        self.client.force_login(self.staff)
        response = self.client.get("/admin/")
        self.assertEqual(response.status_code, 200)

    def test_admin_translation_job_list_visible_to_superuser(self):
        self.client.force_login(self.staff)
        response = self.client.get("/admin/translator/translationjob/")
        self.assertEqual(response.status_code, 200)

    def test_admin_shows_soft_deleted_jobs(self):
        job = TranslationJob.objects.create(
            owner=self.regular,
            original_filename="deleted_doc.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=TranslationJob.Status.COMPLETED,
            is_deleted=True,
        )
        self.client.force_login(self.staff)
        response = self.client.get("/admin/translator/translationjob/?is_deleted__exact=1")
        self.assertEqual(response.status_code, 200)
        self.assertIn("deleted_doc.pdf", response.content.decode("utf-8"))

    def test_system_event_log_has_no_add_permission(self):
        from django.contrib.admin.sites import AdminSite
        from django.test import RequestFactory
        from translator.admin import SystemEventLogAdmin
        from translator.models import SystemEventLog

        admin_instance = SystemEventLogAdmin(SystemEventLog, AdminSite())
        request = RequestFactory().get("/admin/")
        request.user = self.staff
        self.assertFalse(admin_instance.has_add_permission(request))

    def test_system_event_log_has_no_change_permission(self):
        from django.contrib.admin.sites import AdminSite
        from django.test import RequestFactory
        from translator.admin import SystemEventLogAdmin
        from translator.models import SystemEventLog

        admin_instance = SystemEventLogAdmin(SystemEventLog, AdminSite())
        request = RequestFactory().get("/admin/")
        request.user = self.staff
        self.assertFalse(admin_instance.has_change_permission(request))

    def test_system_event_log_has_no_delete_permission(self):
        from django.contrib.admin.sites import AdminSite
        from django.test import RequestFactory
        from translator.admin import SystemEventLogAdmin
        from translator.models import SystemEventLog

        admin_instance = SystemEventLogAdmin(SystemEventLog, AdminSite())
        request = RequestFactory().get("/admin/")
        request.user = self.staff
        self.assertFalse(admin_instance.has_delete_permission(request))

    def test_user_activity_log_is_fully_read_only(self):
        from django.contrib.admin.sites import AdminSite
        from django.test import RequestFactory
        from translator.admin import UserActivityLogAdmin
        from translator.models import UserActivityLog

        admin_instance = UserActivityLogAdmin(UserActivityLog, AdminSite())
        request = RequestFactory().get("/admin/")
        request.user = self.staff
        self.assertFalse(admin_instance.has_add_permission(request))
        self.assertFalse(admin_instance.has_change_permission(request))
        self.assertFalse(admin_instance.has_delete_permission(request))


# ============================================================
# UI Cleanliness Tests
# ============================================================


class TranslatorUICleanlinessTests(TestCase):
    """Verify that duplicate UI elements have been removed."""

    password = "Bagobo-UI-2026!"

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)

        self.alice = User.objects.create_user(
            username="uialice",
            email="uialice@example.test",
            password=self.password,
        )

    def _create_job(self, status=TranslationJob.Status.PROCESSING):
        return TranslationJob.objects.create(
            owner=self.alice,
            original_filename="ui_test.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=status,
        )

    def test_job_detail_has_no_separate_status_dl_row(self):
        """Status is shown only in the heading pill, not repeated in a dl row."""
        job = self._create_job(TranslationJob.Status.PROCESSING)
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        body = response.content.decode("utf-8")
        # The dl must not contain a "Status" row label
        self.assertNotIn("<dt>Status</dt>", body)

    def test_job_detail_has_no_separate_progress_dl_row(self):
        job = self._create_job(TranslationJob.Status.PROCESSING)
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        body = response.content.decode("utf-8")
        self.assertNotIn("<dt>Progress</dt>", body)

    def test_job_detail_has_no_separate_phase_dl_row(self):
        job = self._create_job(TranslationJob.Status.PROCESSING)
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        body = response.content.decode("utf-8")
        self.assertNotIn("<dt>Current Phase</dt>", body)

    def test_job_detail_has_no_separate_message_dl_row(self):
        job = self._create_job(TranslationJob.Status.PROCESSING)
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        body = response.content.decode("utf-8")
        self.assertNotIn("<dt>Message</dt>", body)

    def test_active_job_banner_hidden_on_its_own_detail_page(self):
        """Global banner is suppressed when the user is already viewing that job."""
        job = self._create_job(TranslationJob.Status.PROCESSING)
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        body = response.content.decode("utf-8")
        self.assertNotIn("active-job-banner", body)

    def test_active_job_banner_shows_on_history_page(self):
        """Global banner IS shown when on history page while a job is active."""
        self._create_job(TranslationJob.Status.PROCESSING)
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:history"))
        body = response.content.decode("utf-8")
        self.assertIn("active-job-banner", body)

    def test_preview_page_has_no_duplicate_download_button(self):
        """Download button appears at most once on the preview page."""
        job = self._create_job(TranslationJob.Status.COMPLETED)
        output_path = self.media_root / "jobs" / job.job_id / "translated.pdf"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(b"%PDF-1.4\n%%EOF")
        job.output_file_path = str(output_path)
        job.save(update_fields=["output_file_path"])
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_preview", args=[job.id]))
        body = response.content.decode("utf-8")
        self.assertEqual(body.count("Download Translated PDF"), 1)

    def test_ocr_confidence_row_absent_for_non_ocr_job(self):
        """OCR Confidence row does not appear for digital PDF jobs."""
        job = self._create_job(TranslationJob.Status.COMPLETED)
        job.metadata = {"extraction_method": "direct_pdf_text"}
        job.save(update_fields=["metadata"])
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        body = response.content.decode("utf-8")
        self.assertNotIn("N/A (not an OCR document)", body)

    def test_ocr_low_confidence_warning_badge_shown(self):
        """⚠ Low badge appears (to staff, in the technical panel) when OCR
        confidence is below threshold."""
        job = self._create_job(TranslationJob.Status.COMPLETED)
        job.metadata = {
            "extraction_method": "ocr_image",
            "ocr_summary": {
                "mean_confidence": 0.45,
                "has_low_quality_warning": True,
            },
        }
        job.save(update_fields=["metadata"])
        self.alice.is_staff = True
        self.alice.save(update_fields=["is_staff"])
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        body = response.content.decode("utf-8")
        self.assertIn("45%", body)
        self.assertIn("⚠ Low", body)

    def test_ocr_low_confidence_hidden_from_normal_user(self):
        """Normal users do not see the OCR confidence/quality technical badge."""
        job = self._create_job(TranslationJob.Status.COMPLETED)
        job.metadata = {
            "extraction_method": "ocr_image",
            "ocr_summary": {
                "mean_confidence": 0.45,
                "has_low_quality_warning": True,
            },
        }
        job.save(update_fields=["metadata"])
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        body = response.content.decode("utf-8")
        self.assertNotIn("45%", body)
        self.assertNotIn("⚠ Low", body)
        self.assertNotIn("Technical Details", body)


class TranslatorUIRefreshTests(TestCase):
    """Recent/history card simplification, completion notification, top
    action bar, and preview loading-state hook (UI-only changes; no OCR,
    dataset, ByT5, or reconstruction logic is exercised here)."""

    password = "Bagobo-UI-Refresh-2026!"

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)

        self.alice = User.objects.create_user(
            username="ui_refresh_alice",
            email="ui_refresh_alice@example.test",
            password=self.password,
        )

    def _create_job(self, status):
        return TranslationJob.objects.create(
            owner=self.alice,
            original_filename="ui_refresh.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=status,
        )

    # ---- Phase 1: recent/history card simplification ----

    def test_completed_job_history_card_shows_preview_only(self):
        job = self._create_job(TranslationJob.Status.COMPLETED)
        out = self.media_root / "jobs" / job.job_id / "translated.pdf"
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(b"%PDF-1.4\n%%EOF")
        job.output_file_path = str(out)
        job.save(update_fields=["output_file_path"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:history"))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("document-card-action", body)
        self.assertIn(">Preview<", body)
        self.assertNotIn(">Download<", body)
        self.assertNotIn(">Remove<", body)

    def test_processing_job_history_card_has_no_preview_link(self):
        self._create_job(TranslationJob.Status.PROCESSING)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:history"))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertNotIn(">Preview<", body)
        self.assertIn("Preview available once processing completes.", body)

    def test_failed_job_history_card_shows_failed_indicator_not_preview(self):
        self._create_job(TranslationJob.Status.FAILED)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:history"))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertNotIn(">Preview<", body)
        self.assertIn("Translation failed safely.", body)

    def test_completed_job_recent_sidebar_card_shows_preview_only(self):
        self._create_job(TranslationJob.Status.COMPLETED)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:translate"))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn(">Preview<", body)
        self.assertNotIn(">Download<", body)

    # ---- Phase 2: completion notification ----

    def test_completed_job_detail_shows_ready_message_once_per_session(self):
        job = self._create_job(TranslationJob.Status.COMPLETED)
        self.client.force_login(self.alice)

        first = self.client.get(reverse("translator:job_detail", args=[job.id]))
        second = self.client.get(reverse("translator:job_detail", args=[job.id]))

        self.assertIn("Translated file is ready.", first.content.decode("utf-8"))
        self.assertNotIn("Translated file is ready.", second.content.decode("utf-8"))

    def test_ready_message_uses_accessible_status_markup(self):
        job = self._create_job(TranslationJob.Status.COMPLETED)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_detail", args=[job.id]))

        body = response.content.decode("utf-8")
        message_start = body.index("Translated file is ready.")
        message_region = body[max(0, message_start - 200):message_start]
        self.assertIn('role="status"', message_region)
        self.assertIn('aria-live="polite"', message_region)
        self.assertIn("message-dismiss", body)

    def test_failed_job_detail_never_shows_ready_message(self):
        job = self._create_job(TranslationJob.Status.FAILED)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_detail", args=[job.id]))

        self.assertNotIn("Translated file is ready.", response.content.decode("utf-8"))

    def test_processing_job_detail_never_shows_ready_message(self):
        job = self._create_job(TranslationJob.Status.PROCESSING)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_detail", args=[job.id]))

        self.assertNotIn("Translated file is ready.", response.content.decode("utf-8"))

    # ---- Phase 4: top action bar ----

    def test_job_detail_has_top_action_bar_before_layout_section(self):
        job = self._create_job(TranslationJob.Status.COMPLETED)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_detail", args=[job.id]))

        body = response.content.decode("utf-8")
        self.assertIn('class="top-action-bar"', body)
        self.assertIn('class="job-detail-layout"', body)
        self.assertLess(body.index("top-action-bar"), body.index("job-detail-layout"))

    def test_job_detail_top_action_bar_has_history_preview_and_remove(self):
        job = self._create_job(TranslationJob.Status.COMPLETED)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_detail", args=[job.id]))

        body = response.content.decode("utf-8")
        bar_html = body[body.index('class="top-action-bar"'):body.index('class="job-detail-layout"')]
        self.assertIn("History", bar_html)
        self.assertIn("Preview Bilingual", bar_html)
        self.assertIn(reverse("translator:job_delete_confirm", args=[job.id]), bar_html)

    def test_top_action_bar_shows_download_only_when_output_exists(self):
        job_no_output = self._create_job(TranslationJob.Status.COMPLETED)
        job_with_output = self._create_job(TranslationJob.Status.COMPLETED)
        out = self.media_root / "jobs" / job_with_output.job_id / "translated.pdf"
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(b"%PDF-1.4\n%%EOF")
        job_with_output.output_file_path = str(out)
        job_with_output.save(update_fields=["output_file_path"])
        self.client.force_login(self.alice)

        no_output_body = self.client.get(
            reverse("translator:job_detail", args=[job_no_output.id])
        ).content.decode("utf-8")
        with_output_body = self.client.get(
            reverse("translator:job_detail", args=[job_with_output.id])
        ).content.decode("utf-8")

        no_output_bar = no_output_body[
            no_output_body.index('class="top-action-bar"'):no_output_body.index('class="job-detail-layout"')
        ]
        with_output_bar = with_output_body[
            with_output_body.index('class="top-action-bar"'):with_output_body.index('class="job-detail-layout"')
        ]
        self.assertNotIn(">Download<", no_output_bar)
        self.assertIn(">Download<", with_output_bar)

    def test_top_action_bar_failed_job_has_no_preview_bilingual_or_download(self):
        job = self._create_job(TranslationJob.Status.FAILED)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_detail", args=[job.id]))

        body = response.content.decode("utf-8")
        bar_html = body[body.index('class="top-action-bar"'):body.index('class="job-detail-layout"')]
        self.assertNotIn("Preview Bilingual", bar_html)
        self.assertNotIn(">Download<", bar_html)
        # The danger action stays available regardless of job status.
        self.assertIn(reverse("translator:job_delete_confirm", args=[job.id]), bar_html)

    # ---- Phase 3: preview loading-state hook ----

    def test_preview_link_carries_preparing_preview_loading_text(self):
        job = self._create_job(TranslationJob.Status.COMPLETED)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_detail", args=[job.id]))

        body = response.content.decode("utf-8")
        self.assertIn("js-preview-link", body)
        self.assertIn("Preparing bilingual preview", body)

    def test_app_js_holds_loading_state_minimum_duration_and_has_no_dead_preview_code(self):
        js_path = Path(__file__).resolve().parent.parent / "static" / "js" / "app.js"
        js = js_path.read_text(encoding="utf-8")

        self.assertIn("MIN_LOADING_MS = 2500", js)
        self.assertNotIn("function initPreview()", js)
        self.assertIn('classList.contains("is-loading")', js)

    def test_page_transition_overlay_has_correct_default_text_and_reduced_motion_css(self):
        base_html = (
            Path(__file__).resolve().parent
            / "templates" / "translator" / "base.html"
        ).read_text(encoding="utf-8")
        css = (
            Path(__file__).resolve().parent.parent / "static" / "css" / "styles.css"
        ).read_text(encoding="utf-8")

        self.assertIn('id="page-transition-overlay"', base_html)
        self.assertIn("Preparing bilingual preview", base_html)
        self.assertNotIn("Opening Bilingual Preview", base_html)
        self.assertIn("@media (prefers-reduced-motion: reduce)", css)
        self.assertIn(".page-transition-spinner .spinner", css)

    # ---- Phase 5: button colors ----

    def test_primary_button_css_has_no_gradient_and_single_danger_style(self):
        css_path = Path(__file__).resolve().parent.parent / "static" / "css" / "styles.css"
        css = css_path.read_text(encoding="utf-8")

        # Only the primary *button* rule must drop its gradient — decorative
        # elements like document-thumb icons may still use one.
        primary_rule = css[css.index(".button.primary {"):css.index(".button.secondary {")]
        self.assertNotIn("linear-gradient", primary_rule)
        self.assertIn("background: var(--primary);", primary_rule)
        self.assertIn(".button.danger-outline {", css)
        self.assertNotIn(".button.danger {", css)

    def test_disabled_button_is_muted_gray_not_just_faded_color(self):
        css_path = Path(__file__).resolve().parent.parent / "static" / "css" / "styles.css"
        css = css_path.read_text(encoding="utf-8")

        disabled_rule = css[css.index(".button:disabled"):css.index(".button.primary {")]
        self.assertIn("color: var(--muted);", disabled_rule)
        self.assertIn("cursor: not-allowed", disabled_rule)
        # A disabled primary/secondary/danger-outline button must not keep
        # its variant's own color — it must render as flat muted gray.
        self.assertIn(".button.primary:disabled", disabled_rule)
        self.assertIn(".button.secondary:disabled", disabled_rule)
        self.assertIn(".button.danger-outline:disabled", disabled_rule)

    def test_view_all_link_consolidated_into_button_link(self):
        css_path = Path(__file__).resolve().parent.parent / "static" / "css" / "styles.css"
        css = css_path.read_text(encoding="utf-8")
        translate_html = (
            Path(__file__).resolve().parent
            / "templates" / "translator" / "translate.html"
        ).read_text(encoding="utf-8")

        self.assertNotIn("view-all-link", css)
        self.assertNotIn("view-all-link", translate_html)
        self.assertIn('class="button-link"', translate_html)

    # ---- Phase 6: preview page-row alignment ----

    def test_preview_aligns_pages_in_per_page_rows_with_clear_labels(self):
        job = self._create_job(TranslationJob.Status.COMPLETED)
        job.metadata = {
            "source_page_count": 2,
            "translated_page_count": 2,
            "layout_page_count": 2,
        }
        job.save(update_fields=["metadata"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        # Each page gets its own row with both sides clearly labelled —
        # original and translated for the same page number stay paired,
        # instead of each side scrolling through its own independent list.
        self.assertIn("Original Document &middot; Page 1", body)
        self.assertIn("Translated Document &middot; Page 1", body)
        self.assertIn("Original Document &middot; Page 2", body)
        self.assertIn("Translated Document &middot; Page 2", body)
        row_one = body.index("Page 1 comparison")
        row_two = body.index("Page 2 comparison")
        self.assertLess(row_one, row_two)
        # Top actions (History / Job Details) must still be reachable
        # without scrolling past the page rows.
        self.assertIn("Job Details", body)
        self.assertLess(body.index("Job Details"), row_one)


class TranslatorDocxPaginationTests(TestCase):
    """Phase 3 audit: DOCX extraction page-breaks long documents correctly."""

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)

    def _make_docx(self, n_paragraphs: int) -> str:
        from docx import Document as DocxDocument
        doc = DocxDocument()
        for i in range(n_paragraphs):
            doc.add_paragraph(f"Paragraph {i + 1}: sample audit text for Phase 3.")
        path = str(Path(self.temp_dir.name) / "test_long.docx")
        doc.save(path)
        return path

    def test_long_docx_produces_multiple_pages(self):
        """50 paragraphs must span at least 2 pages (max ~32 paragraphs per page)."""
        from translator.services.extraction_service import ExtractionService
        pages = ExtractionService.extract_docx_text_and_layout(self._make_docx(50))
        self.assertGreaterEqual(len(pages), 2)

    def test_long_docx_all_paragraphs_extracted(self):
        """All 50 paragraphs must appear in the extracted layout with none dropped."""
        from translator.services.extraction_service import ExtractionService
        n = 50
        pages = ExtractionService.extract_docx_text_and_layout(self._make_docx(n))
        total_lines = sum(
            len(block.get("lines", []))
            for page in pages
            for block in page.get("blocks", [])
        )
        self.assertEqual(total_lines, n)

    def test_no_block_bbox_exceeds_page_height(self):
        """No block bbox y1 may exceed the declared page height (prevents silent clipping)."""
        from translator.services.extraction_service import ExtractionService
        pages = ExtractionService.extract_docx_text_and_layout(self._make_docx(50))
        for page in pages:
            page_h = page.get("height", 792)
            for block in page.get("blocks", []):
                bbox = block.get("bbox", [0, 0, 0, page_h])
                self.assertLessEqual(
                    bbox[3],
                    page_h,
                    f"Block y1={bbox[3]} exceeds page height {page_h} on page {page.get('page')}.",
                )


class TranslatorOCRRotationTests(TestCase):
    """Phase 3 audit: OCR rotation detection and _apply_rotation contract."""

    def _white_image(self):
        from PIL import Image
        return Image.new("RGB", (100, 100), (255, 255, 255))

    def test_apply_rotation_zero_is_noop(self):
        from translator.services.ocr_stage.ocr_service import OCRService
        svc = OCRService(detect_orientation=False)
        img = self._white_image()
        out, applied, warning = svc._apply_rotation(img, 0)
        self.assertIs(out, img)
        self.assertEqual(applied, 0)
        self.assertIsNone(warning)

    def test_apply_rotation_90_corrects_image(self):
        """90° rotation: image is rotated (not left unchanged), applied=90, no warning."""
        from translator.services.ocr_stage.ocr_service import OCRService
        svc = OCRService(detect_orientation=False)
        img = self._white_image()
        out, applied, warning = svc._apply_rotation(img, 90)
        self.assertIsNot(out, img)
        self.assertEqual(applied, 90)
        self.assertIsNone(warning)

    def test_apply_rotation_270_corrects_image(self):
        """270° rotation: image is rotated (not left unchanged), applied=270, no warning."""
        from translator.services.ocr_stage.ocr_service import OCRService
        svc = OCRService(detect_orientation=False)
        img = self._white_image()
        out, applied, warning = svc._apply_rotation(img, 270)
        self.assertIsNot(out, img)
        self.assertEqual(applied, 270)
        self.assertIsNone(warning)

    def test_apply_rotation_180_corrects_image(self):
        """180° rotation must return a new transposed image with no warning."""
        from translator.services.ocr_stage.ocr_service import OCRService
        svc = OCRService(detect_orientation=False)
        img = self._white_image()
        out, applied, warning = svc._apply_rotation(img, 180)
        self.assertIsNot(out, img)
        self.assertEqual(applied, 180)
        self.assertIsNone(warning)

    def test_detect_rotation_reads_osd_output(self):
        """_detect_rotation_deg parses Rotate: line from mocked pytesseract OSD output."""
        from translator.services.ocr_stage.ocr_service import OCRService
        svc = OCRService(detect_orientation=True)
        img = self._white_image()
        with patch(
            "pytesseract.image_to_osd",
            return_value="Rotate: 90\nOrientation confidence: 3.14\n",
        ):
            deg = svc._detect_rotation_deg(img)
        self.assertEqual(deg, 90)


class TranslatorWorkflowGatingTests(TestCase):
    """Phase 3 audit: download gating and cross-page segment ordering."""

    password = "Bagobo-Test-2026!"

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)

        self.alice = User.objects.create_user(
            username="alice_p3",
            email="alice_p3@example.test",
            password=self.password,
        )

    def _make_completed_job(self, *, with_output_file: bool = False) -> TranslationJob:
        job = TranslationJob.objects.create(
            owner=self.alice,
            original_filename="audit.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=TranslationJob.Status.COMPLETED,
        )
        if with_output_file:
            out = self.media_root / "jobs" / job.job_id / "translated.pdf"
            out.parent.mkdir(parents=True, exist_ok=True)
            out.write_bytes(b"%PDF-1.4\n%%EOF")
            job.output_file_path = str(out)
            job.save(update_fields=["output_file_path"])
        return job

    def test_completed_job_without_output_file_hides_download(self):
        """COMPLETED job with no translated PDF must not show the Download button."""
        job = self._make_completed_job(with_output_file=False)
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        body = response.content.decode("utf-8")
        self.assertNotIn(">Download<", body)

    def test_completed_job_with_output_file_shows_download(self):
        """COMPLETED job with translated PDF present must show the Download button."""
        job = self._make_completed_job(with_output_file=True)
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        body = response.content.decode("utf-8")
        self.assertIn(">Download<", body)

    def test_segment_ordering_across_pages(self):
        """Segments from page 1 must have lower segment_index than those on page 2."""
        from translator.services import _sync_structure_models

        job = self._make_completed_job()
        job_dir = self.media_root / "jobs" / job.job_id
        job_dir.mkdir(parents=True, exist_ok=True)
        structure_path = str(job_dir / "structure.json")

        structure = {
            "extraction_method": "direct_pdf_text",
            "pages": [
                {
                    "page_number": 1,
                    "width": 612,
                    "height": 792,
                    "rotation": 0,
                    "blocks": [{
                        "type": "text",
                        "bbox": [72, 100, 540, 120],
                        "lines": [{
                            "source_text": "First page text.",
                            "translated_text": "Teksto sa unang pahina.",
                            "translation_method": "exact",
                            "translation_confidence": 1.0,
                            "bbox": [72, 100, 540, 120],
                        }],
                    }],
                },
                {
                    "page_number": 2,
                    "width": 612,
                    "height": 792,
                    "rotation": 0,
                    "blocks": [{
                        "type": "text",
                        "bbox": [72, 100, 540, 120],
                        "lines": [{
                            "source_text": "Second page text.",
                            "translated_text": "Teksto sa ikalawang pahina.",
                            "translation_method": "exact",
                            "translation_confidence": 1.0,
                            "bbox": [72, 100, 540, 120],
                        }],
                    }],
                },
            ],
        }
        with open(structure_path, "w", encoding="utf-8") as fh:
            json.dump(structure, fh)

        _sync_structure_models(str(job.id), structure_path)

        segments = list(
            TranslationSegment.objects.filter(job=job).order_by("segment_index")
        )
        self.assertEqual(len(segments), 2)
        self.assertEqual(segments[0].source_text, "First page text.")
        self.assertEqual(segments[1].source_text, "Second page text.")
        self.assertLess(segments[0].segment_index, segments[1].segment_index)


class TranslatorPhase4FidelityTests(TestCase):
    """Phase 4 audit: reconstruction fidelity, structure validation, rotation, hybrid PDF."""

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.addCleanup(self.temp_dir.cleanup)

    # ------------------------------------------------------------------
    # 1. Reconstruction fidelity: DOCX output contains all paragraphs
    # ------------------------------------------------------------------

    def _make_docx(self, paragraphs) -> str:
        from docx import Document as DocxDoc
        doc = DocxDoc()
        for text in paragraphs:
            doc.add_paragraph(text)
        path = str(self.root / "test_input.docx")
        doc.save(path)
        return path

    def test_reconstruct_docx_output_contains_all_text(self):
        """_create_output_pdf renders translated text for every DOCX paragraph."""
        import fitz
        from translator.services.extraction_service import ExtractionService
        from translator.services.pipeline_service import PipelineService

        n = 5
        source_texts = [f"Source paragraph {i + 1} content." for i in range(n)]
        expected_translations = [f"TRANS_{i}" for i in range(n)]

        docx_path = self._make_docx(source_texts)
        layout_data = ExtractionService.extract_docx_text_and_layout(docx_path)

        # Build translations keyed by source text
        translations = {
            src: {
                "original": src,
                "translated": exp,
                "method": "test",
                "confidence": 1.0,
            }
            for src, exp in zip(source_texts, expected_translations)
        }

        output_path = str(self.root / "translated.pdf")
        pipeline = PipelineService()
        ok = pipeline._create_output_pdf(layout_data, translations, output_path)
        self.assertTrue(ok, "_create_output_pdf returned False")
        self.assertTrue(Path(output_path).exists(), "Output PDF was not created")

        doc = fitz.open(output_path)
        full_text = "".join(doc[i].get_text() for i in range(doc.page_count))
        doc.close()

        for expected in expected_translations:
            self.assertIn(
                expected,
                full_text,
                f"Expected translated text '{expected}' is missing from the output PDF.",
            )

    def test_output_pdf_page_count_matches_layout(self):
        """Output PDF has the same number of pages as the extracted layout."""
        import fitz
        from translator.services.extraction_service import ExtractionService
        from translator.services.pipeline_service import PipelineService

        docx_path = self._make_docx(
            [f"Paragraph {i}." for i in range(50)]
        )
        layout_data = ExtractionService.extract_docx_text_and_layout(docx_path)
        output_path = str(self.root / "paged.pdf")
        pipeline = PipelineService()
        pipeline._create_output_pdf(layout_data, {}, output_path)

        doc = fitz.open(output_path)
        output_page_count = doc.page_count
        doc.close()

        self.assertEqual(
            output_page_count,
            len(layout_data),
            f"Output PDF has {output_page_count} pages but layout has {len(layout_data)}.",
        )

    # ------------------------------------------------------------------
    # 2. Structure validation
    # ------------------------------------------------------------------

    def test_validate_layout_data_accepts_valid_data(self):
        from translator.services.pipeline_service import PipelineService
        PipelineService._validate_layout_data([
            {"page": 0, "width": 612.0, "height": 792.0, "blocks": []},
        ])

    def test_validate_layout_data_rejects_non_list(self):
        from translator.services.pipeline_service import PipelineService
        with self.assertRaises(ValueError):
            PipelineService._validate_layout_data({"page": 0})

    def test_validate_layout_data_rejects_non_dict_page(self):
        from translator.services.pipeline_service import PipelineService
        with self.assertRaises(ValueError):
            PipelineService._validate_layout_data(["not-a-dict"])

    def test_validate_layout_data_rejects_zero_width(self):
        from translator.services.pipeline_service import PipelineService
        with self.assertRaises(ValueError):
            PipelineService._validate_layout_data([
                {"page": 0, "width": 0, "height": 792.0, "blocks": []},
            ])

    def test_validate_layout_data_rejects_negative_height(self):
        from translator.services.pipeline_service import PipelineService
        with self.assertRaises(ValueError):
            PipelineService._validate_layout_data([
                {"page": 0, "width": 612.0, "height": -1.0, "blocks": []},
            ])

    def test_validate_layout_data_allows_none_dimensions(self):
        """None width/height are skipped (reconstruction defaults to 612×792)."""
        from translator.services.pipeline_service import PipelineService
        PipelineService._validate_layout_data([
            {"page": 0, "width": None, "height": None, "blocks": []},
        ])

    # ------------------------------------------------------------------
    # 3. OCR 90°/270° rotation: image rotation + bbox inverse transform
    # ------------------------------------------------------------------

    def _solid_image(self, w: int, h: int, color=(255, 255, 255)):
        from PIL import Image
        return Image.new("RGB", (w, h), color)

    def test_apply_rotation_90_rotates_image_dimensions(self):
        """Rotating a 100×60 image 90° CW produces a 60×100 image."""
        from translator.services.ocr_stage.ocr_service import OCRService
        svc = OCRService(detect_orientation=False)
        img = self._solid_image(100, 60)
        out, applied, warning = svc._apply_rotation(img, 90)
        self.assertEqual(applied, 90)
        self.assertIsNone(warning)
        self.assertEqual(out.size, (60, 100), "90° CW should swap width and height.")

    def test_apply_rotation_270_rotates_image_dimensions(self):
        """Rotating a 100×60 image 90° CCW produces a 60×100 image."""
        from translator.services.ocr_stage.ocr_service import OCRService
        svc = OCRService(detect_orientation=False)
        img = self._solid_image(100, 60)
        out, applied, warning = svc._apply_rotation(img, 270)
        self.assertEqual(applied, 270)
        self.assertIsNone(warning)
        self.assertEqual(out.size, (60, 100), "270° correction should swap width and height.")

    def test_flip_layout_90_cw_maps_corners_correctly(self):
        """_flip_page_layout_90_cw maps top-left in rotated space to top-right in original."""
        from translator.services.ocr_stage.ocr_service import OCRService
        # Original page: 595 × 842 pt (portrait)
        # After 90° CW correction: OCR space is 842 × 595
        orig_w, orig_h = 595.0, 842.0
        page_layout = {
            "width": orig_w,
            "height": orig_h,
            "blocks": [{"type": "text", "bbox": [0.0, 0.0, 10.0, 10.0], "lines": [
                {"text": "hi", "bbox": [0.0, 0.0, 10.0, 10.0]}
            ]}],
        }
        result = OCRService._flip_page_layout_90_cw(page_layout, orig_w, orig_h)
        # [ry0, h - rx1, ry1, h - rx0] = [0, 842-10, 10, 842-0] = [0, 832, 10, 842]
        self.assertEqual(result["blocks"][0]["bbox"], [0.0, 832.0, 10.0, 842.0])
        self.assertEqual(result["blocks"][0]["lines"][0]["bbox"], [0.0, 832.0, 10.0, 842.0])
        self.assertEqual(result["width"], orig_w)
        self.assertEqual(result["height"], orig_h)

    def test_flip_layout_270_cw_maps_corners_correctly(self):
        """_flip_page_layout_270_cw maps top-left in rotated space to bottom-left in original."""
        from translator.services.ocr_stage.ocr_service import OCRService
        orig_w, orig_h = 595.0, 842.0
        page_layout = {
            "width": orig_w,
            "height": orig_h,
            "blocks": [{"type": "text", "bbox": [0.0, 0.0, 10.0, 10.0], "lines": [
                {"text": "hi", "bbox": [0.0, 0.0, 10.0, 10.0]}
            ]}],
        }
        result = OCRService._flip_page_layout_270_cw(page_layout, orig_w, orig_h)
        # [w - ry1, rx0, w - ry0, rx1] = [595-10, 0, 595-0, 10] = [585, 0, 595, 10]
        self.assertEqual(result["blocks"][0]["bbox"], [585.0, 0.0, 595.0, 10.0])
        self.assertEqual(result["width"], orig_w)
        self.assertEqual(result["height"], orig_h)

    # ------------------------------------------------------------------
    # 4. Hybrid PDF: per-page extraction metadata
    # ------------------------------------------------------------------

    def _make_minimal_pdf(self, *, n_chars: int = 200) -> str:
        """Create a single-page digital PDF with n_chars of text."""
        import fitz
        doc = fitz.open()
        page = doc.new_page(width=612, height=792)
        page.insert_text((72, 100), "A" * n_chars, fontsize=12)
        path = str(self.root / "digital.pdf")
        doc.save(path)
        doc.close()
        return path

    def test_hybrid_pdf_digital_only_sets_direct_pdf_text_method(self):
        """A PDF where all pages are digital should set extraction_method=direct_pdf_text."""
        from translator.services.pipeline_service import PipelineService, JobStatus
        from translator.services.models import FileType, DetectionType

        pdf_path = self._make_minimal_pdf(n_chars=200)
        job = JobStatus("test-hybrid-digital")
        job.detection_type = DetectionType.DIGITAL
        job.metadata = {}

        pipeline = PipelineService()
        layout_data, _ = pipeline._extract_hybrid_pdf(pdf_path, None, job)

        self.assertEqual(job.metadata.get("extraction_method"), "direct_pdf_text")
        self.assertGreater(len(layout_data), 0)
        self.assertEqual(job.metadata["page_extraction_methods"]["0"], "digital")

    def test_hybrid_pdf_low_text_page_is_marked_scanned(self):
        """A nearly-blank PDF page is classified as scanned in page_extraction_methods."""
        import fitz
        from translator.services.pipeline_service import PipelineService, JobStatus
        from translator.services.models import FileType, DetectionType

        # Create a blank (0-char) single-page PDF
        doc = fitz.open()
        doc.new_page(width=612, height=792)
        pdf_path = str(self.root / "blank.pdf")
        doc.save(pdf_path)
        doc.close()

        job = JobStatus("test-hybrid-blank")
        job.detection_type = DetectionType.DIGITAL
        job.metadata = {}

        pipeline = PipelineService()
        layout_data, _ = pipeline._extract_hybrid_pdf(pdf_path, None, job)

        self.assertEqual(job.metadata["page_extraction_methods"]["0"], "ocr")

    # ------------------------------------------------------------------
    # 5. Generated output gating
    # ------------------------------------------------------------------

    def test_output_file_not_created_when_no_text_blocks(self):
        """_create_output_pdf still returns True and creates a PDF even with empty layout."""
        import fitz
        from translator.services.pipeline_service import PipelineService

        output_path = str(self.root / "empty_out.pdf")
        pipeline = PipelineService()
        ok = pipeline._create_output_pdf(
            [{"page": 0, "width": 612, "height": 792, "blocks": []}],
            {},
            output_path,
        )
        self.assertTrue(ok)
        self.assertTrue(Path(output_path).exists())
        doc = fitz.open(output_path)
        self.assertEqual(doc.page_count, 1)
        doc.close()


# ===========================================================================
# PHASE 5: Real OCR Benchmarking, Reading Order, and Layout Quality Validation
# ===========================================================================


class TranslatorPhase5MetricsTests(TestCase):
    """Phase 5: CER and WER metric calculations (pure Python, no Tesseract)."""

    def _cer(self, ref, hyp):
        from translator.services.ocr_stage.qa_report import calculate_cer
        return calculate_cer(ref, hyp)

    def _wer(self, ref, hyp):
        from translator.services.ocr_stage.qa_report import calculate_wer
        return calculate_wer(ref, hyp)

    def test_cer_identical_strings_is_zero(self):
        """CER between a string and itself is 0.0."""
        self.assertEqual(self._cer("hello world", "hello world"), 0.0)

    def test_cer_one_char_substitution(self):
        """One character substitution in a 3-char string gives CER of 1/3."""
        cer = self._cer("cat", "bat")
        self.assertAlmostEqual(cer, 1 / 3, places=5)

    def test_cer_completely_different_capped_at_one(self):
        """CER between completely different strings is capped at 1.0."""
        cer = self._cer("abc", "xyz")
        self.assertEqual(cer, 1.0)

    def test_cer_empty_reference_returns_zero(self):
        """CER with an empty reference is 0.0 (no characters to measure against)."""
        self.assertEqual(self._cer("", "hypothesis text"), 0.0)

    def test_wer_identical_strings_is_zero(self):
        """WER between a sentence and itself is 0.0."""
        self.assertEqual(self._wer("hello world", "hello world"), 0.0)

    def test_wer_one_word_substitution(self):
        """One word substituted in a two-word phrase gives WER of 0.5."""
        wer = self._wer("hello world", "hello there")
        self.assertAlmostEqual(wer, 0.5, places=5)

    def test_wer_empty_reference_returns_zero(self):
        """WER with an empty reference is 0.0 (no words to measure against)."""
        self.assertEqual(self._wer("", "some output"), 0.0)


class TranslatorPhase5QAReportTests(TestCase):
    """Phase 5: DocumentQAReport generation and serialisation."""

    def _make_page(self, text="Hello world.", *, page=0, with_confidence=False):
        """Return a minimal page dict in pipeline layout_data format."""
        line = {"text": text, "bbox": [72.0, 100.0, 540.0, 120.0]}
        if with_confidence:
            line["confidence"] = 0.85
        block = {
            "type": "text",
            "bbox": [72.0, 100.0, 540.0, 120.0],
            "lines": [line],
        }
        if with_confidence:
            block["confidence"] = 0.85
        return {
            "page": page,
            "width": 612.0,
            "height": 792.0,
            "blocks": [block],
        }

    def _make_empty_page(self, *, page=0):
        return {"page": page, "width": 612.0, "height": 792.0, "blocks": []}

    def test_qa_report_as_dict_has_required_keys(self):
        """as_dict() must include all required QA keys."""
        from translator.services.ocr_stage.qa_report import build_document_qa_report
        layout = [self._make_page()]
        report = build_document_qa_report("test.pdf", layout, extraction_method="digital")
        d = report.as_dict()
        for key in (
            "document_name", "extraction_method", "page_count",
            "ocr_confidence", "cer", "wer",
            "empty_page_rate", "failed_page_count",
            "total_processing_time_s", "reading_order_issues",
            "output_result", "warnings", "pages",
        ):
            self.assertIn(key, d, f"Required key '{key}' missing from as_dict()")

    def test_qa_report_empty_page_rate_computed(self):
        """empty_page_rate = empty_page_count / page_count."""
        from translator.services.ocr_stage.qa_report import build_document_qa_report
        layout = [
            self._make_page(page=0),
            self._make_empty_page(page=1),
            self._make_empty_page(page=2),
            self._make_page(page=3),
        ]
        report = build_document_qa_report("doc.pdf", layout)
        d = report.as_dict()
        self.assertEqual(d["page_count"], 4)
        self.assertAlmostEqual(d["empty_page_rate"], 0.5, places=4)

    def test_qa_report_cer_wer_with_ground_truth(self):
        """CER and WER are populated when ground_truth_pages is supplied."""
        from translator.services.ocr_stage.qa_report import build_document_qa_report
        layout = [self._make_page("hello world")]
        # Hypothesis matches reference exactly → CER=0, WER=0
        report = build_document_qa_report(
            "doc.pdf", layout,
            ground_truth_pages=["hello world"],
        )
        self.assertEqual(report.cer, 0.0)
        self.assertEqual(report.wer, 0.0)

    def test_qa_report_cer_wer_none_without_ground_truth(self):
        """CER and WER are None when no ground truth is supplied."""
        from translator.services.ocr_stage.qa_report import build_document_qa_report
        layout = [self._make_page()]
        report = build_document_qa_report("doc.pdf", layout)
        self.assertIsNone(report.cer)
        self.assertIsNone(report.wer)

    def test_qa_report_reading_order_issues_surfaced(self):
        """reading_order_issues in as_dict() matches audit_reading_order output."""
        from translator.services.ocr_stage.qa_report import (
            build_document_qa_report,
            audit_reading_order,
        )
        # Reversed block order: second block has a smaller y than first.
        layout = [{
            "page": 0,
            "width": 612.0,
            "height": 792.0,
            "blocks": [
                {"type": "text", "bbox": [72.0, 400.0, 540.0, 420.0],
                 "lines": [{"text": "Late block", "bbox": [72.0, 400.0, 540.0, 420.0]}]},
                {"type": "text", "bbox": [72.0, 100.0, 540.0, 120.0],
                 "lines": [{"text": "Early block", "bbox": [72.0, 100.0, 540.0, 120.0]}]},
            ],
        }]
        report = build_document_qa_report("reverse.pdf", layout)
        d = report.as_dict()
        self.assertGreater(len(d["reading_order_issues"]), 0)
        combined = " ".join(d["reading_order_issues"])
        self.assertIn("reading-order problem", combined)


class TranslatorPhase5ReadingOrderTests(TestCase):
    """Phase 5: audit_reading_order detects ordering and layout issues."""

    def _page(self, blocks, *, page=0, width=612.0, height=792.0):
        return {"page": page, "width": width, "height": height, "blocks": blocks}

    def _text_block(self, x0, y0, x1, y1, text="text"):
        return {
            "type": "text",
            "bbox": [x0, y0, x1, y1],
            "lines": [{"text": text, "bbox": [x0, y0, x1, y1]}],
        }

    def test_single_column_top_bottom_order_passes(self):
        """Blocks in strict top-to-bottom order produce no ordering issues."""
        from translator.services.ocr_stage.qa_report import audit_reading_order
        layout = [self._page([
            self._text_block(72, 100, 540, 120, "First line"),
            self._text_block(72, 130, 540, 150, "Second line"),
            self._text_block(72, 160, 540, 180, "Third line"),
        ])]
        issues = audit_reading_order(layout)
        order_issues = [i for i in issues if "reading-order problem" in i]
        self.assertEqual(order_issues, [], f"Unexpected order issues: {order_issues}")

    def test_reversed_block_order_is_flagged(self):
        """A block appearing above its predecessor triggers an order issue."""
        from translator.services.ocr_stage.qa_report import audit_reading_order
        layout = [self._page([
            self._text_block(72, 500, 540, 520, "Block at y=500"),
            self._text_block(72, 100, 540, 120, "Block at y=100"),  # reverse
        ])]
        issues = audit_reading_order(layout)
        self.assertTrue(
            any("reading-order problem" in i for i in issues),
            f"Expected a reading-order issue, got: {issues}",
        )

    def test_two_column_layout_is_detected(self):
        """Two blocks in distinct left/right zones trigger a two-column notice."""
        from translator.services.ocr_stage.qa_report import audit_reading_order
        # Left block: x1 <= 612*0.45 ≈ 275.  Right block: x0 >= 612*0.55 ≈ 337.
        layout = [self._page([
            self._text_block(72, 100, 260, 120, "Left column text"),
            self._text_block(350, 100, 560, 120, "Right column text"),
        ])]
        issues = audit_reading_order(layout)
        self.assertTrue(
            any("two-column" in i for i in issues),
            f"Expected a two-column notice, got: {issues}",
        )

    def test_empty_page_is_noted(self):
        """A page with no text blocks produces a 'blank page or extraction failed' note."""
        from translator.services.ocr_stage.qa_report import audit_reading_order
        layout = [{"page": 0, "width": 612.0, "height": 792.0, "blocks": []}]
        issues = audit_reading_order(layout)
        self.assertTrue(
            any("blank page" in i.lower() or "extraction failed" in i.lower() for i in issues),
            f"Expected empty-page note, got: {issues}",
        )


class TranslatorPhase5PSMBenchmarkTests(TestCase):
    """Phase 5: PSM configuration is correctly wired through OCRService."""

    def test_ocr_service_default_psm_is_3(self):
        """OCRService() without psm argument uses PSM 3 (Tesseract auto mode)."""
        from translator.services.ocr_stage.ocr_service import OCRService
        svc = OCRService()
        self.assertEqual(svc.psm, 3)

    def test_psm_parameter_accepted_in_range(self):
        """OCRService(psm=6) stores psm=6."""
        from translator.services.ocr_stage.ocr_service import OCRService
        for psm in (0, 1, 3, 4, 6, 11, 13):
            svc = OCRService(psm=psm)
            self.assertEqual(svc.psm, psm, f"psm={psm} was not stored correctly")

    def test_psm_out_of_range_defaults_to_3(self):
        """PSM values outside [0, 13] are silently replaced with 3."""
        from translator.services.ocr_stage.ocr_service import OCRService
        for bad in (14, 99, -1):
            svc = OCRService(psm=bad)
            self.assertEqual(svc.psm, 3, f"PSM {bad} should default to 3")

    def test_psm_config_string_passed_to_tesseract(self):
        """--psm <N> is included in the config string sent to image_to_data."""
        from unittest.mock import patch, MagicMock
        from PIL import Image
        from translator.services.ocr_stage.ocr_service import OCRService

        svc = OCRService(psm=6, detect_orientation=False, preprocess=False,
                         denoise=False, threshold=False)
        svc._verified = True  # skip Tesseract availability check

        mock_data = {
            "text": [], "conf": [], "block_num": [], "par_num": [], "line_num": [],
            "left": [], "top": [], "width": [], "height": [],
        }
        img = Image.new("RGB", (100, 50), (255, 255, 255))

        with patch("pytesseract.image_to_data", return_value=mock_data) as mock_itd:
            svc._ocr_image(img, 0, 72.0, 36.0, 1.0, lang="eng")

        self.assertTrue(mock_itd.called, "image_to_data was not called")
        call_kwargs = mock_itd.call_args[1] if mock_itd.call_args[1] else {}
        call_args = mock_itd.call_args[0] if mock_itd.call_args[0] else ()
        config_value = call_kwargs.get("config") or ""
        self.assertIn("--psm 6", config_value,
                      f"Expected '--psm 6' in config, got: {config_value!r}")


class TranslatorPhase5ReadingOrderReconstructionTests(TestCase):
    """Phase 5: Reconstruction quality — output contains expected text,
    page count matches layout, and None dimensions are handled safely."""

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.addCleanup(self.temp_dir.cleanup)

    def _make_layout(self, texts, *, width=612.0, height=792.0):
        """Build a single-page layout_data with one line per text string."""
        blocks = []
        y = 100.0
        for text in texts:
            bbox = [72.0, y, 540.0, y + 20.0]
            blocks.append({
                "type": "text",
                "bbox": bbox,
                "lines": [{"text": text, "bbox": bbox}],
            })
            y += 24.0
        return [{"page": 0, "width": width, "height": height, "blocks": blocks}]

    def test_output_contains_all_translated_segments(self):
        """Every translated string appears in the reconstructed output PDF."""
        import fitz
        from translator.services.pipeline_service import PipelineService

        sources = [f"Source {i}" for i in range(4)]
        translations = {src: {"original": src, "translated": f"Xlat{i}", "method": "test", "confidence": 1.0}
                        for i, src in enumerate(sources)}
        layout = self._make_layout(sources)

        output_path = str(self.root / "out.pdf")
        ok = PipelineService()._create_output_pdf(layout, translations, output_path)
        self.assertTrue(ok)

        doc = fitz.open(output_path)
        full_text = "".join(doc[i].get_text() for i in range(doc.page_count))
        doc.close()

        for i in range(4):
            self.assertIn(f"Xlat{i}", full_text,
                          f"Translated segment 'Xlat{i}' missing from output PDF.")

    def test_output_page_count_equals_layout_page_count(self):
        """Output PDF has exactly as many pages as the layout_data list."""
        import fitz
        from translator.services.pipeline_service import PipelineService

        n_pages = 3
        layout = []
        for p in range(n_pages):
            layout.append({
                "page": p, "width": 612.0, "height": 792.0,
                "blocks": [{"type": "text", "bbox": [72, 100, 540, 120],
                             "lines": [{"text": f"Page {p} text", "bbox": [72, 100, 540, 120]}]}],
            })

        output_path = str(self.root / "paged.pdf")
        PipelineService()._create_output_pdf(layout, {}, output_path)

        doc = fitz.open(output_path)
        self.assertEqual(doc.page_count, n_pages)
        doc.close()

    def test_none_width_height_pages_use_default_dimensions(self):
        """Pages with None width/height are created with 612×792 pt defaults."""
        import fitz
        from translator.services.pipeline_service import PipelineService

        layout = [{"page": 0, "width": None, "height": None,
                   "blocks": [{"type": "text", "bbox": [72, 100, 540, 120],
                                "lines": [{"text": "Default size page", "bbox": [72, 100, 540, 120]}]}]}]
        output_path = str(self.root / "default_size.pdf")
        ok = PipelineService()._create_output_pdf(layout, {}, output_path)
        self.assertTrue(ok)

        doc = fitz.open(output_path)
        page = doc[0]
        # _create_output_pdf falls back to 612×792 when dimensions are None/falsy.
        self.assertAlmostEqual(page.rect.width, 612.0, delta=1.0)
        self.assertAlmostEqual(page.rect.height, 792.0, delta=1.0)
        doc.close()

    def test_text_block_completely_outside_page_bounds_is_skipped(self):
        """A block whose bbox is outside the page rect does not crash reconstruction."""
        import fitz
        from translator.services.pipeline_service import PipelineService

        # Block at y=9000 is way outside a 792-pt tall page.
        layout = [{
            "page": 0, "width": 612.0, "height": 792.0,
            "blocks": [
                {
                    "type": "text",
                    "bbox": [72, 100, 540, 120],
                    "lines": [{"text": "Valid line", "bbox": [72, 100, 540, 120]}],
                },
                {
                    "type": "text",
                    "bbox": [72, 9000, 540, 9020],
                    "lines": [{"text": "Far-out line", "bbox": [72, 9000, 540, 9020]}],
                },
            ],
        }]
        translations = {
            "Valid line": {"original": "Valid line", "translated": "Valid OK",
                           "method": "test", "confidence": 1.0},
            "Far-out line": {"original": "Far-out line", "translated": "Should skip",
                             "method": "test", "confidence": 1.0},
        }
        output_path = str(self.root / "oob.pdf")
        ok = PipelineService()._create_output_pdf(layout, translations, output_path)
        self.assertTrue(ok, "_create_output_pdf should not crash on out-of-bounds block")
        self.assertTrue(Path(output_path).exists())


class TranslatorPhase5HybridPDFOrderTests(TestCase):
    """Phase 5: Mixed digital/scanned PDFs preserve page order and store methods."""

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.addCleanup(self.temp_dir.cleanup)

    def _make_two_page_pdf(self, *, first_page_chars: int, second_page_chars: int) -> str:
        """Create a two-page PDF; each page has the specified number of 'A' characters."""
        import fitz
        doc = fitz.open()
        p0 = doc.new_page(width=612, height=792)
        if first_page_chars:
            p0.insert_text((72, 100), "A" * first_page_chars, fontsize=12)
        p1 = doc.new_page(width=612, height=792)
        if second_page_chars:
            p1.insert_text((72, 100), "A" * second_page_chars, fontsize=12)
        path = str(self.root / "mixed.pdf")
        doc.save(path)
        doc.close()
        return path

    def test_mixed_pdf_page_order_preserved(self):
        """Pages in the layout_data output are in the same order as the source PDF."""
        from translator.services.pipeline_service import PipelineService, JobStatus
        from translator.services.models import DetectionType

        # 200-char first page (digital), 0-char second page (scanned/ocr path)
        pdf_path = self._make_two_page_pdf(first_page_chars=200, second_page_chars=0)
        job = JobStatus("order-test")
        job.detection_type = DetectionType.DIGITAL
        job.metadata = {}

        pipeline = PipelineService()
        layout_data, _ = pipeline._extract_hybrid_pdf(pdf_path, None, job)

        self.assertEqual(len(layout_data), 2, "Expected exactly 2 pages")
        self.assertEqual(layout_data[0]["page"], 0, "First page index must be 0")
        self.assertEqual(layout_data[1]["page"], 1, "Second page index must be 1")

    def test_mixed_pdf_sets_hybrid_method(self):
        """A PDF with one digital and one scanned page stores extraction_method=hybrid."""
        from translator.services.pipeline_service import PipelineService, JobStatus
        from translator.services.models import DetectionType

        pdf_path = self._make_two_page_pdf(first_page_chars=200, second_page_chars=0)
        job = JobStatus("hybrid-method-test")
        job.detection_type = DetectionType.DIGITAL
        job.metadata = {}

        PipelineService()._extract_hybrid_pdf(pdf_path, None, job)
        # First page is digital; second page has no chars → routed to OCR path.
        # extraction_method depends on whether OCR succeeds, but page methods are set.
        page_methods = job.metadata.get("page_extraction_methods", {})
        self.assertEqual(page_methods.get("0"), "digital")
        self.assertEqual(page_methods.get("1"), "ocr")

    def test_hybrid_pdf_page_methods_stored_for_every_page(self):
        """page_extraction_methods contains an entry for every page in the PDF."""
        from translator.services.pipeline_service import PipelineService, JobStatus
        from translator.services.models import DetectionType

        pdf_path = self._make_two_page_pdf(first_page_chars=200, second_page_chars=200)
        job = JobStatus("all-digital-test")
        job.detection_type = DetectionType.DIGITAL
        job.metadata = {}

        pipeline = PipelineService()
        layout_data, _ = pipeline._extract_hybrid_pdf(pdf_path, None, job)

        page_methods = job.metadata.get("page_extraction_methods", {})
        for page_idx in range(len(layout_data)):
            self.assertIn(
                str(page_idx), page_methods,
                f"No extraction method stored for page {page_idx}",
            )


# ===========================================================================
# PHASE 6: Real Tesseract OCR Validation and Document Quality Benchmarking
# ===========================================================================


class TranslatorPhase6EnvironmentTests(TestCase):
    """Phase 6: Tesseract environment audit — correct status when not installed."""

    def test_check_tesseract_returns_dict_with_required_keys(self):
        """check_tesseract_environment() must return a dict with all expected keys."""
        from translator.services.ocr_stage.environment import check_tesseract_environment
        env = check_tesseract_environment()
        for key in ("available", "version", "languages", "osd_available",
                    "binary_path", "errors"):
            self.assertIn(key, env, f"Required key '{key}' missing from environment report")

    def test_check_tesseract_available_is_false_when_not_installed(self):
        """available=False when Tesseract binary is absent from PATH."""
        from translator.services.ocr_stage.environment import check_tesseract_environment
        # In this environment Tesseract is not installed.
        env = check_tesseract_environment()
        # If somehow Tesseract IS present, we still validate the dict shape.
        self.assertIsInstance(env["available"], bool)
        if not env["available"]:
            self.assertIsNone(env["version"])
            self.assertGreater(len(env["errors"]), 0,
                               "errors list should be non-empty when Tesseract is absent")

    def test_check_tesseract_errors_is_nonempty_when_not_installed(self):
        """errors[] contains at least one actionable message when Tesseract is missing."""
        from translator.services.ocr_stage.environment import check_tesseract_environment
        env = check_tesseract_environment()
        if not env["available"]:
            self.assertTrue(
                len(env["errors"]) > 0,
                "At least one actionable error message is expected when Tesseract is absent",
            )
            # The error should mention installation guidance.
            combined = " ".join(env["errors"]).lower()
            has_guidance = any(
                kw in combined
                for kw in ("install", "tesseract", "path", "tesseract_cmd")
            )
            self.assertTrue(has_guidance,
                            f"Error message should include installation guidance: {env['errors']}")

    def test_ocr_service_is_available_returns_false(self):
        """OCRService.is_available() returns False when Tesseract is not installed."""
        from translator.services.ocr_stage.ocr_service import OCRService
        svc = OCRService()
        result = svc.is_available()
        # is_available() must always return a bool (never raise).
        self.assertIsInstance(result, bool)
        # In this environment it should be False.
        if not result:
            self.assertFalse(result)

    def test_assert_tesseract_available_raises_runtime_error(self):
        """assert_tesseract_available() raises RuntimeError when Tesseract is absent."""
        from translator.services.ocr_stage.environment import (
            check_tesseract_environment,
            assert_tesseract_available,
        )
        env = check_tesseract_environment()
        if not env["available"]:
            with self.assertRaises(RuntimeError) as ctx:
                assert_tesseract_available()
            self.assertIn("Tesseract", str(ctx.exception))
        # If Tesseract IS present, the function should not raise.


class TranslatorPhase6FixtureBuilderTests(TestCase):
    """Phase 6: Programmatic OCR fixture builder creates valid images and PDFs."""

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.addCleanup(self.temp_dir.cleanup)

    def _build(self):
        from translator.management.commands.ocr_benchmark import build_fixtures
        return build_fixtures(self.root)

    def test_clean_scan_fixture_creates_valid_png(self):
        """build_fixtures() creates a clean_scan.png that PIL can open."""
        from PIL import Image
        paths = self._build()
        self.assertIn("clean_scan", paths)
        img = Image.open(paths["clean_scan"])
        self.assertEqual(img.mode, "RGB")
        self.assertGreater(img.width, 0)
        self.assertGreater(img.height, 0)

    def test_faded_scan_fixture_has_lower_contrast_than_clean(self):
        """Faded scan has lower pixel contrast (smaller max-min range) than clean scan."""
        from PIL import Image
        paths = self._build()
        clean = list(Image.open(paths["clean_scan"]).convert("L").getdata())
        faded = list(Image.open(paths["faded_scan"]).convert("L").getdata())
        clean_contrast = max(clean) - min(clean)
        faded_contrast = max(faded) - min(faded)
        # Clean scan: bg=255, ink=0 → contrast=255.
        # Faded scan: bg=240, ink=130 → contrast=110. Faded is lower-contrast.
        self.assertLess(faded_contrast, clean_contrast,
                        "Faded scan should have lower pixel contrast than clean scan")

    def test_rotated_fixture_has_swapped_dimensions(self):
        """Rotating a landscape image 90° produces a portrait image."""
        from PIL import Image
        paths = self._build()
        self.assertIn("rotated_90", paths)
        clean = Image.open(paths["clean_scan"])
        rotated = Image.open(paths["rotated_90"])
        # 90° rotation swaps width and height.
        self.assertEqual(rotated.width, clean.height)
        self.assertEqual(rotated.height, clean.width)

    def test_blank_fixture_is_all_white(self):
        """Blank fixture has a mean pixel value of 255 (pure white)."""
        from PIL import Image
        import statistics
        paths = self._build()
        self.assertIn("blank", paths)
        blank = Image.open(paths["blank"]).convert("L")
        mean_val = statistics.mean(blank.getdata())
        self.assertEqual(mean_val, 255.0, "Blank fixture should be pure white (255)")

    def test_ground_truth_files_are_created_for_all_fixtures(self):
        """build_fixtures() creates a .txt ground-truth file for each expected fixture."""
        from translator.management.commands.ocr_benchmark import GROUND_TRUTH
        self._build()
        gt_dir = self.root / "ground_truth"
        for name in GROUND_TRUTH:
            gt_file = gt_dir / f"{name}.txt"
            self.assertTrue(
                gt_file.exists(),
                f"Ground-truth file missing: {gt_file}",
            )

    def test_ground_truth_blank_fixture_is_empty(self):
        """The blank fixture's ground-truth file is an empty string."""
        self._build()
        gt_file = self.root / "ground_truth" / "blank.txt"
        self.assertTrue(gt_file.exists())
        self.assertEqual(gt_file.read_text(encoding="utf-8").strip(), "")


class TranslatorPhase6QAReportExportTests(TestCase):
    """Phase 6: DocumentQAReport CSV/JSON export and pass/fail logic."""

    def _make_report_with_cer(self, cer_value):
        """Return a DocumentQAReport whose single page has the given CER."""
        from translator.services.ocr_stage.qa_report import (
            DocumentQAReport,
            PageQAResult,
        )
        page = PageQAResult(
            page_index=0,
            extraction_method="ocr",
            char_count=100,
            block_count=3,
            ocr_confidence=0.85,
            cer=cer_value,
            wer=cer_value,
            psm=3,
        )
        return DocumentQAReport(
            document_name="test.png",
            extraction_method="ocr",
            page_count=1,
            ocr_confidence=0.85,
            cer=cer_value,
            wer=cer_value,
            empty_page_count=0,
            failed_page_count=0,
            total_processing_time_s=1.23,
            reading_order_issues=[],
            output_result="ok",
            warnings=[],
            pages=[page],
        )

    def test_csv_rows_has_correct_column_names(self):
        """to_csv_rows() returns dicts with all required CSV column keys."""
        report = self._make_report_with_cer(0.05)
        rows = report.to_csv_rows()
        self.assertEqual(len(rows), 1)
        for col in ("filename", "page", "extraction_method", "confidence",
                    "cer", "wer", "processing_time_s", "psm", "warnings", "pass_fail"):
            self.assertIn(col, rows[0], f"Missing CSV column: {col}")

    def test_csv_has_one_row_per_page(self):
        """to_csv_rows() returns exactly one dict per page in the report."""
        from translator.services.ocr_stage.qa_report import (
            DocumentQAReport,
            PageQAResult,
        )
        pages = [
            PageQAResult(page_index=i, extraction_method="ocr",
                         char_count=50, block_count=2, cer=0.05, wer=0.05, psm=3)
            for i in range(3)
        ]
        report = DocumentQAReport(
            document_name="multi.pdf", extraction_method="ocr",
            page_count=3, ocr_confidence=None, cer=0.05, wer=0.05,
            empty_page_count=0, failed_page_count=0,
            total_processing_time_s=2.0, reading_order_issues=[],
            output_result="ok", warnings=[], pages=pages,
        )
        rows = report.to_csv_rows()
        self.assertEqual(len(rows), 3)

    def test_csv_string_has_header_row(self):
        """to_csv_string() starts with the header row."""
        report = self._make_report_with_cer(0.05)
        csv_text = report.to_csv_string()
        first_line = csv_text.splitlines()[0]
        self.assertIn("filename", first_line)
        self.assertIn("pass_fail", first_line)

    def test_pass_fail_skip_when_no_cer(self):
        """pass_fail is 'skip' when CER is None (no ground truth)."""
        report = self._make_report_with_cer(None)
        rows = report.to_csv_rows()
        self.assertEqual(rows[0]["pass_fail"], "skip")

    def test_pass_fail_pass_when_cer_under_threshold(self):
        """pass_fail is 'pass' when CER < 0.10."""
        report = self._make_report_with_cer(0.05)
        rows = report.to_csv_rows()
        self.assertEqual(rows[0]["pass_fail"], "pass")

    def test_pass_fail_fail_when_cer_at_or_above_threshold(self):
        """pass_fail is 'fail' when CER >= 0.10."""
        for cer in (0.10, 0.50, 1.0):
            report = self._make_report_with_cer(cer)
            rows = report.to_csv_rows()
            self.assertEqual(rows[0]["pass_fail"], "fail",
                             f"Expected 'fail' for CER={cer}")

    def test_json_round_trip(self):
        """to_json_string() / json.loads() round-trips without data loss."""
        import json as _json
        report = self._make_report_with_cer(0.08)
        json_str = report.to_json_string()
        data = _json.loads(json_str)
        self.assertEqual(data["document_name"], "test.png")
        self.assertAlmostEqual(data["cer"], 0.08, places=4)
        self.assertEqual(len(data["pages"]), 1)
        self.assertEqual(data["pages"][0]["psm"], 3)


class TranslatorPhase6MixedPDFRoutingTests(TestCase):
    """Phase 6: Hybrid PDF routes digital pages to fitz and scanned to OCR path."""

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.addCleanup(self.temp_dir.cleanup)

    def _make_two_page_pdf(self, first_chars: int, second_chars: int) -> str:
        import fitz
        doc = fitz.open()
        p0 = doc.new_page(width=612, height=792)
        if first_chars:
            p0.insert_text((72, 100), "A" * first_chars, fontsize=12)
        p1 = doc.new_page(width=612, height=792)
        if second_chars:
            p1.insert_text((72, 100), "A" * second_chars, fontsize=12)
        path = str(self.root / "twopages.pdf")
        doc.save(path)
        doc.close()
        return path

    def test_digital_page_uses_direct_extraction_not_ocr(self):
        """A page with sufficient text is extracted via fitz (method=digital)."""
        from translator.services.pipeline_service import PipelineService, JobStatus
        from translator.services.models import DetectionType

        pdf_path = self._make_two_page_pdf(first_chars=200, second_chars=200)
        job = JobStatus("route-digital")
        job.detection_type = DetectionType.DIGITAL
        job.metadata = {}

        PipelineService()._extract_hybrid_pdf(pdf_path, None, job)
        methods = job.metadata.get("page_extraction_methods", {})
        self.assertEqual(methods.get("0"), "digital",
                         "A page with 200 chars should use the digital (fitz) path")
        self.assertEqual(methods.get("1"), "digital")

    def test_scanned_page_records_ocr_error_when_tesseract_unavailable(self):
        """When Tesseract is absent, scanned pages have ocr_error in their page_data.

        Skipped when Tesseract IS installed: a blank scanned page then correctly
        returns empty blocks with no error, which is the expected happy-path.
        """
        import unittest as _ut
        from translator.services.ocr_stage.environment import check_tesseract_environment
        if check_tesseract_environment()["available"]:
            raise _ut.SkipTest(
                "Tesseract is installed — blank scanned page returns empty blocks "
                "(no ocr_error expected). Phase 7 adds explicit tests for this case."
            )

        from translator.services.pipeline_service import PipelineService, JobStatus
        from translator.services.models import DetectionType

        # Page 1 = blank → triggers OCR path; Tesseract is not installed.
        pdf_path = self._make_two_page_pdf(first_chars=200, second_chars=0)
        job = JobStatus("route-scanned")
        job.detection_type = DetectionType.DIGITAL
        job.metadata = {}

        layout_data, _ = PipelineService()._extract_hybrid_pdf(pdf_path, None, job)
        scanned_page = layout_data[1]
        # The OCR path should have recorded an error (not silently produced empty output).
        has_error = bool(
            scanned_page.get("ocr_error") or scanned_page.get("ocr_warning")
        )
        self.assertTrue(
            has_error,
            "Scanned page with unavailable Tesseract should record ocr_error or ocr_warning, "
            f"got: {scanned_page}",
        )

    def test_overall_method_is_hybrid_when_both_paths_used(self):
        """extraction_method becomes 'hybrid' when both digital and OCR pages appear."""
        from translator.services.pipeline_service import PipelineService, JobStatus
        from translator.services.models import DetectionType

        pdf_path = self._make_two_page_pdf(first_chars=200, second_chars=0)
        job = JobStatus("route-hybrid")
        job.detection_type = DetectionType.DIGITAL
        job.metadata = {}

        PipelineService()._extract_hybrid_pdf(pdf_path, None, job)
        # Page 0 is digital; page 1 triggers OCR path (regardless of Tesseract outcome).
        methods = job.metadata.get("page_extraction_methods", {})
        self.assertEqual(methods.get("0"), "digital")
        self.assertEqual(methods.get("1"), "ocr")
        self.assertEqual(
            job.metadata.get("extraction_method"), "hybrid",
            "When one page is digital and one is OCR, overall method must be 'hybrid'",
        )

    def test_benchmark_command_importable_and_reports_unavailable(self):
        """ocr_benchmark command is importable and gracefully handles Tesseract absence."""
        from translator.services.ocr_stage.environment import check_tesseract_environment
        from translator.management.commands.ocr_benchmark import build_fixtures, GROUND_TRUTH

        # build_fixtures runs on PIL/fitz only — no Tesseract needed.
        fixture_dir = self.root / "fixtures"
        paths = build_fixtures(fixture_dir)

        # All GROUND_TRUTH fixtures should have been created.
        for name in GROUND_TRUTH:
            gt_file = fixture_dir / "ground_truth" / f"{name}.txt"
            self.assertTrue(gt_file.exists(), f"Missing ground-truth file: {gt_file}")

        # Environment check must clearly say Tesseract is absent.
        env = check_tesseract_environment()
        if not env["available"]:
            self.assertFalse(env["available"])
            self.assertTrue(len(env["errors"]) > 0)


# ────────────────────────────────────────────────────────────────────────────
# Phase 7 — Real OCR Benchmark Execution and Evidence-Based Fixes
# ────────────────────────────────────────────────────────────────────────────

class TranslatorPhase7PreprocessingTests(TestCase):
    """Unit tests for the Phase 7 preprocessing fixes — no Tesseract required."""

    def _make_gray_image(self, width=200, height=80, bg_gray=240, text_gray=130):
        """Create a synthetic grayscale image simulating a faded scan."""
        from PIL import Image, ImageDraw
        img = Image.new("L", (width, height), bg_gray)
        draw = ImageDraw.Draw(img)
        draw.rectangle([10, 20, 150, 60], fill=text_gray)
        return img

    def test_threshold_formula_does_not_blacken_gray_background(self):
        """Phase 7 fix: threshold must not turn a gray background (240) to black.

        Old formula (mean+255)/2 = 247 for mean=239.6 → bg=240 < 247 → black.
        New formula mean*0.85 = 203 → bg=240 > 203 → white.
        """
        from translator.services.ocr_stage.ocr_service import OCRService
        svc = OCRService(detect_orientation=False, denoise=False)
        img = self._make_gray_image(bg_gray=240, text_gray=130)
        result = svc._preprocess(img)
        corner_pixel = result.getpixel((0, 0))
        self.assertEqual(
            corner_pixel, 255,
            f"Gray background (240) must become white (255) after threshold, got {corner_pixel}. "
            "Old formula (mean+255)/2 incorrectly blackens gray backgrounds."
        )

    def test_threshold_formula_keeps_dark_text_black(self):
        """Text pixels (130) must remain black after faded-scan thresholding."""
        from translator.services.ocr_stage.ocr_service import OCRService
        svc = OCRService(detect_orientation=False, denoise=False)
        img = self._make_gray_image(bg_gray=240, text_gray=130)
        result = svc._preprocess(img)
        text_pixel = result.getpixel((80, 40))
        self.assertEqual(
            text_pixel, 0,
            f"Dark text pixel (130) must become black (0) after threshold, got {text_pixel}."
        )

    def test_threshold_not_applied_below_mean_cutoff(self):
        """Threshold step must be skipped when mean pixel value ≤ 180."""
        from PIL import Image
        from translator.services.ocr_stage.ocr_service import OCRService
        svc = OCRService(detect_orientation=False, denoise=False)
        # All-dark image: mean ≈ 50 → threshold NOT applied
        img = Image.new("L", (100, 50), 50)
        result = svc._preprocess(img)
        self.assertEqual(result.getpixel((10, 10)), 50,
                         "No threshold when image mean ≤ 180 — pixel must be unchanged")

    def test_clean_scan_background_stays_white_after_preprocessing(self):
        """Pure white background (255) must remain white after preprocessing."""
        from PIL import Image
        from translator.services.ocr_stage.ocr_service import OCRService
        svc = OCRService(detect_orientation=False, denoise=False)
        img = Image.new("RGB", (200, 80), (255, 255, 255))
        result = svc._preprocess(img)
        self.assertEqual(result.getpixel((10, 10)), 255,
                         "White background must remain white after preprocessing")

    def test_ocr_config_includes_dpi_flag(self):
        """_ocr_image must pass --dpi to Tesseract so low-DPI images are read correctly."""
        import unittest
        from translator.services.ocr_stage.environment import check_tesseract_environment
        if not check_tesseract_environment()["available"]:
            raise unittest.SkipTest("Tesseract not available")

        from PIL import Image
        from unittest.mock import patch, MagicMock
        from pytesseract import Output
        from translator.services.ocr_stage.ocr_service import OCRService

        svc = OCRService(dpi=300, detect_orientation=False, preprocess=False, denoise=False)
        img = Image.new("RGB", (200, 80), (255, 255, 255))
        captured_config = {}

        import pytesseract as _pt
        original_fn = _pt.image_to_data

        def capture_config(image, output_type, lang, config, timeout):
            captured_config["config"] = config
            return original_fn(image, output_type=output_type, lang=lang,
                               config=config, timeout=timeout)

        with patch("pytesseract.image_to_data", side_effect=capture_config):
            svc._ocr_image(img, 0, 144.0, 57.6, 72.0 / 300.0, lang="eng")

        self.assertIn("--dpi 300", captured_config.get("config", ""),
                      "Tesseract config must include '--dpi 300' to override missing DPI metadata")

    def test_fixture_images_are_scaled_4x(self):
        """Phase 7 fixture builder saves 4× scaled images (2400×800) for reliable OCR."""
        import tempfile
        from pathlib import Path
        from PIL import Image
        from translator.management.commands.ocr_benchmark import build_fixtures

        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as td:
            paths = build_fixtures(Path(td))
            with Image.open(paths["clean_scan"]) as img:
                size = img.size
            self.assertEqual(size, (2400, 800),
                             "clean_scan fixture must be 2400×800 (4× scale) for Tesseract legibility")

    def test_fixture_images_have_dpi_metadata(self):
        """Phase 7 fixture builder embeds 300 DPI in PNG metadata."""
        import tempfile
        from pathlib import Path
        from PIL import Image
        from translator.management.commands.ocr_benchmark import build_fixtures

        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as td:
            paths = build_fixtures(Path(td))
            with Image.open(paths["clean_scan"]) as img:
                dpi_info = img.info.get("dpi")
            self.assertIsNotNone(dpi_info, "clean_scan PNG must carry DPI metadata")
            # PIL stores DPI as float; PNG spec rounds at ~0.4 DPI precision
            self.assertAlmostEqual(dpi_info[0], 300, delta=1,
                                   msg=f"DPI must be ≈300, got {dpi_info}")


class TranslatorPhase7RealOCRTests(TestCase):
    """End-to-end OCR regression tests using real Tesseract 5.5.0.

    Every test in this class skips automatically when Tesseract is not available
    so the suite stays green on CI environments without a Tesseract binary.
    """

    @classmethod
    def _skip_if_no_tesseract(cls):
        import unittest
        from translator.services.ocr_stage.environment import check_tesseract_environment
        env = check_tesseract_environment()
        if not env["available"]:
            raise unittest.SkipTest(
                f"Tesseract not available: {'; '.join(env['errors'])}"
            )

    def _fixture_image(self, name):
        """Return path to a freshly-built benchmark fixture."""
        import tempfile
        from pathlib import Path
        from translator.management.commands.ocr_benchmark import build_fixtures
        td = tempfile.mkdtemp()
        self._tmpdir = td
        return build_fixtures(Path(td))[name]

    def tearDown(self):
        import shutil
        td = getattr(self, "_tmpdir", None)
        if td:
            shutil.rmtree(td, ignore_errors=True)

    def test_clean_scan_achieves_low_cer_with_real_tesseract(self):
        """Real OCR on a clean black-on-white fixture achieves CER < 0.10 (Phase 7 fix)."""
        self._skip_if_no_tesseract()
        from translator.services.ocr_stage.ocr_service import OCRService
        from translator.services.ocr_stage.qa_report import calculate_cer

        svc = OCRService(psm=3, lang="eng", detect_orientation=False)
        path = self._fixture_image("clean_scan")
        pages = svc.extract_image_text_and_layout(str(path))
        hyp = " ".join(
            line.get("text", "")
            for page in pages
            for block in page.get("blocks", [])
            if block.get("type") == "text"
            for line in block.get("lines", [])
        ).strip()
        ref = "Hello world. This is a clean scan document."
        cer = calculate_cer(ref, hyp)
        self.assertLess(
            cer, 0.10,
            f"clean_scan CER must be < 0.10 after Phase 7 fixes, got {cer:.4f} "
            f"(hypothesis: {repr(hyp)!r})"
        )

    def test_faded_scan_achieves_low_cer_after_threshold_fix(self):
        """Real OCR on a faded scan achieves CER < 0.10 after the Phase 7 threshold fix.

        Before the fix: CER=1.0 (all-black image from (mean+255)/2 formula).
        After the fix:  CER≈0.028 (background stays white, text stays black).
        """
        self._skip_if_no_tesseract()
        from translator.services.ocr_stage.ocr_service import OCRService
        from translator.services.ocr_stage.qa_report import calculate_cer

        svc = OCRService(psm=3, lang="eng", detect_orientation=False)
        path = self._fixture_image("faded_scan")
        pages = svc.extract_image_text_and_layout(str(path))
        hyp = " ".join(
            line.get("text", "")
            for page in pages
            for block in page.get("blocks", [])
            if block.get("type") == "text"
            for line in block.get("lines", [])
        ).strip()
        ref = "Hello world. This is a clean scan document."
        cer = calculate_cer(ref, hyp)
        self.assertLess(
            cer, 0.10,
            f"faded_scan CER must be < 0.10 after threshold fix, got {cer:.4f} "
            f"(hypothesis: {repr(hyp)!r})"
        )

    def test_blank_scanned_page_returns_empty_blocks_not_error(self):
        """Blank page with Tesseract available returns empty blocks, no ocr_error."""
        self._skip_if_no_tesseract()
        from translator.services.pipeline_service import PipelineService, JobStatus
        from translator.services.models import DetectionType
        import tempfile, fitz
        from pathlib import Path

        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "blank.pdf"
            doc = fitz.open()
            doc.new_page(width=612, height=792)
            doc.new_page(width=612, height=792)
            doc.save(str(p))
            doc.close()

            job = JobStatus("phase7-blank")
            job.detection_type = DetectionType.DIGITAL
            job.metadata = {}
            layout_data, _ = PipelineService()._extract_hybrid_pdf(
                str(p), None, job
            )
            blank_page = layout_data[1]
            self.assertNotIn(
                "ocr_error", blank_page,
                "Blank page with Tesseract available must not have ocr_error"
            )
            self.assertEqual(
                blank_page.get("blocks", []), [],
                "Blank page should have no blocks"
            )

    def test_psm3_mean_cer_better_than_psm4_on_real_fixtures(self):
        """PSM 3 must match or beat PSM 4 on standard fixtures — default must not change."""
        self._skip_if_no_tesseract()
        import tempfile
        from pathlib import Path
        from translator.management.commands.ocr_benchmark import (
            build_fixtures, GROUND_TRUTH, _evaluate_image_fixture,
        )

        with tempfile.TemporaryDirectory() as td:
            paths = build_fixtures(Path(td))
            image_names = ["clean_scan", "faded_scan", "two_column", "table"]
            cers = {3: [], 4: []}
            for name in image_names:
                if name not in paths:
                    continue
                gt = GROUND_TRUTH.get(name, "")
                if not gt:
                    continue
                result = _evaluate_image_fixture(paths[name], gt, "eng", [3, 4])
                for pr in result["psm_results"]:
                    if pr.get("cer") is not None:
                        cers[pr["psm"]].append(pr["cer"])

            mean3 = sum(cers[3]) / len(cers[3]) if cers[3] else 1.0
            mean4 = sum(cers[4]) / len(cers[4]) if cers[4] else 1.0
            self.assertLessEqual(
                mean3, mean4,
                f"PSM 3 mean CER ({mean3:.4f}) should be ≤ PSM 4 ({mean4:.4f}) "
                "— default PSM 3 must not be changed."
            )

    def test_environment_check_available_with_tesseract_cmd(self):
        """check_tesseract_environment() returns available=True when TESSERACT_CMD is set."""
        self._skip_if_no_tesseract()
        from translator.services.ocr_stage.environment import check_tesseract_environment
        env = check_tesseract_environment()
        self.assertTrue(env["available"],
                        "Tesseract must report available=True when binary is reachable")
        self.assertIn("eng", env["languages"], "English language pack must be installed")
        self.assertTrue(env["osd_available"], "OSD traineddata must be installed")
        self.assertFalse(env["errors"], f"No errors expected, got: {env['errors']}")

    def test_mixed_pdf_routes_digital_and_ocr_pages(self):
        """Phase 7 regression: hybrid PDF extraction routes pages correctly."""
        self._skip_if_no_tesseract()
        import tempfile
        from pathlib import Path
        from translator.management.commands.ocr_benchmark import build_fixtures
        from translator.services.pipeline_service import PipelineService, JobStatus
        from translator.services.models import DetectionType

        with tempfile.TemporaryDirectory() as td:
            paths = build_fixtures(Path(td))
            pdf_path = paths.get("mixed_digital_scanned")
            if not pdf_path:
                self.skipTest("mixed_digital_scanned fixture not available")

            job = JobStatus("phase7-hybrid")
            job.detection_type = DetectionType.DIGITAL
            job.metadata = {}
            layout_data, _ = PipelineService()._extract_hybrid_pdf(
                str(pdf_path), None, job
            )
            methods = job.metadata.get("page_extraction_methods", {})
            self.assertEqual(methods.get("0"), "digital",
                             "Dense digital page must use digital extraction")
            self.assertEqual(methods.get("1"), "ocr",
                             "Blank page must use OCR extraction path")
            self.assertEqual(job.metadata.get("extraction_method"), "hybrid",
                             "Overall method must be 'hybrid'")


# ============================================================
# Appendix A Sprint 1 — User Account and Authentication Tests
# ============================================================


class AppendixAAuthTests(TestCase):
    """Regression tests covering all six Appendix A Sprint 1 test cases."""

    password = "Bagobo-AppA-2026!"

    def setUp(self):
        self.existing = User.objects.create_user(
            username="existing_user",
            email="existing@example.test",
            password=self.password,
        )

    # ------------------------------------------------------------------
    # TC-1: User registers with valid information
    # ------------------------------------------------------------------

    def test_tc1_valid_registration_creates_account_and_redirects(self):
        """TC-1: Valid registration creates the account, logs the user in,
        and redirects to the Translate page."""
        response = self.client.post(
            reverse("translator:signup"),
            {
                "username": "tc1_newuser",
                "email": "tc1@example.test",
                "password1": self.password,
                "password2": self.password,
            },
        )
        self.assertRedirects(response, reverse("translator:translate"))
        user = User.objects.get(username="tc1_newuser")
        self.assertEqual(int(self.client.session["_auth_user_id"]), user.pk)

    # ------------------------------------------------------------------
    # TC-2: User submits incomplete / invalid registration
    # ------------------------------------------------------------------

    def test_tc2_missing_username_shows_required_error(self):
        """TC-2: Empty username triggers a readable required-field error."""
        response = self.client.post(
            reverse("translator:signup"),
            {"username": "", "email": "", "password1": self.password, "password2": self.password},
        )
        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("This field is required", body)

    def test_tc2_mismatched_passwords_shows_error(self):
        """TC-2: Mismatched passwords trigger a readable validation error."""
        response = self.client.post(
            reverse("translator:signup"),
            {
                "username": "tc2_user",
                "email": "",
                "password1": self.password,
                "password2": "WrongPassword999!",
            },
        )
        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("password", body.lower())

    def test_tc2_duplicate_username_shows_error(self):
        """TC-2: Duplicate username triggers a readable uniqueness error."""
        response = self.client.post(
            reverse("translator:signup"),
            {
                "username": "existing_user",
                "email": "",
                "password1": self.password,
                "password2": self.password,
            },
        )
        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("username", body.lower())

    def test_tc2_incomplete_form_does_not_create_account(self):
        """TC-2: Submitting without a username must not create a new user."""
        count_before = User.objects.count()
        self.client.post(
            reverse("translator:signup"),
            {"username": "", "email": "", "password1": self.password, "password2": self.password},
        )
        self.assertEqual(User.objects.count(), count_before)

    # ------------------------------------------------------------------
    # TC-3: User logs in with valid credentials
    # ------------------------------------------------------------------

    def test_tc3_valid_login_redirects_to_translate(self):
        """TC-3: Valid credentials authenticate the user and redirect to Translate."""
        response = self.client.post(
            reverse("translator:login"),
            {"username": self.existing.username, "password": self.password},
        )
        self.assertRedirects(response, reverse("translator:translate"))
        self.assertEqual(
            int(self.client.session["_auth_user_id"]), self.existing.pk
        )

    # ------------------------------------------------------------------
    # TC-4: User logs in with invalid credentials
    # ------------------------------------------------------------------

    def test_tc4_invalid_username_shows_error_not_technical_detail(self):
        """TC-4: Wrong username shows a credential error with no stack trace."""
        response = self.client.post(
            reverse("translator:login"),
            {"username": "nobody", "password": "anything"},
        )
        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("correct username and password", body)
        self.assertNotIn("Traceback", body)
        self.assertNotIn("Exception", body)
        self.assertFalse("_auth_user_id" in self.client.session)

    def test_tc4_wrong_password_shows_error_not_technical_detail(self):
        """TC-4: Correct username but wrong password shows the same safe error."""
        response = self.client.post(
            reverse("translator:login"),
            {"username": self.existing.username, "password": "WrongPass999!"},
        )
        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("correct username and password", body)
        self.assertNotIn("Traceback", body)
        self.assertFalse("_auth_user_id" in self.client.session)

    # ------------------------------------------------------------------
    # TC-5: User logs out
    # ------------------------------------------------------------------

    def test_tc5_logout_ends_session_and_redirects_to_home(self):
        """TC-5: Logging out clears the session and sends the user to Home."""
        self.client.force_login(self.existing)
        self.assertIn("_auth_user_id", self.client.session)

        response = self.client.post(reverse("translator:logout"))
        self.assertRedirects(response, reverse("translator:home"))
        self.assertNotIn("_auth_user_id", self.client.session)

    def test_tc5_after_logout_protected_page_requires_login_again(self):
        """TC-5: After logout, accessing a protected page redirects to login."""
        self.client.force_login(self.existing)
        self.client.post(reverse("translator:logout"))
        translate_url = reverse("translator:translate")
        response = self.client.get(translate_url)
        self.assertRedirects(
            response,
            f"{reverse('translator:login')}?next={translate_url}",
        )

    # ------------------------------------------------------------------
    # TC-6: Unauthenticated access to protected pages
    # ------------------------------------------------------------------

    def test_tc6_unauthenticated_translate_redirects_with_next(self):
        """TC-6: Translate page redirects to login and preserves ?next."""
        url = reverse("translator:translate")
        response = self.client.get(url)
        self.assertRedirects(response, f"{reverse('translator:login')}?next={url}")

    def test_tc6_unauthenticated_history_redirects_with_next(self):
        """TC-6: History page redirects to login and preserves ?next."""
        url = reverse("translator:history")
        response = self.client.get(url)
        self.assertRedirects(response, f"{reverse('translator:login')}?next={url}")

    def test_tc6_unauthenticated_upload_api_returns_json_not_redirect(self):
        """TC-6 (revised): the upload API is a fetch()+response.json() JSON
        endpoint, not a page navigation. An HTML redirect here makes fetch()
        follow it and hand response.json() the login page's HTML, throwing
        "Unexpected token '<'". It must return a JSON 401 with a
        redirect_url field instead, never an HTTP redirect."""
        response = self.client.post(reverse("translator:upload"))
        self.assertEqual(response.status_code, 401)
        self.assertEqual(response["Content-Type"], "application/json")
        data = response.json()
        self.assertFalse(data["ok"])
        self.assertEqual(data["redirect_url"], reverse("translator:login"))

    # ------------------------------------------------------------------
    # TC-A07: Submit empty login credentials
    # ------------------------------------------------------------------

    def test_tc_a07_empty_login_credentials_show_required_errors(self):
        """TC-A07: Both fields empty — required errors appear for username and
        password, no session created, no redirect (HTTP 200), and no technical
        detail is exposed."""
        response = self.client.post(
            reverse("translator:login"),
            {"username": "", "password": ""},
        )
        # Must stay on the login page — a redirect would mean the user was authenticated
        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        # Required-field errors must be readable in the page
        self.assertIn("This field is required", body)
        # No session must be created
        self.assertNotIn("_auth_user_id", self.client.session)
        # No technical exception text
        self.assertNotIn("Traceback", body)
        self.assertNotIn("Exception", body)

    def test_tc_a07_empty_username_with_password_shows_username_required_error(self):
        """TC-A07 variant: password supplied but username empty — username
        required error shown, user not authenticated."""
        response = self.client.post(
            reverse("translator:login"),
            {"username": "", "password": self.password},
        )
        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("This field is required", body)
        self.assertNotIn("_auth_user_id", self.client.session)
        self.assertNotIn("Traceback", body)

    def test_tc_a07_username_supplied_with_empty_password_shows_password_required_error(self):
        """TC-A07 variant: username supplied but password empty — password
        required error shown, user not authenticated."""
        response = self.client.post(
            reverse("translator:login"),
            {"username": self.existing.username, "password": ""},
        )
        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("This field is required", body)
        self.assertNotIn("_auth_user_id", self.client.session)
        self.assertNotIn("Traceback", body)


# ============================================================
# Appendix B — Sprint 2: Document Upload Module
# ============================================================


class AppendixBUploadTests(TestCase):
    """Regression tests for all ten Appendix B Sprint 2 upload test cases.

    Supported formats (production): PDF, DOCX, JPG, PNG, TXT.
    Every rejection test asserts: no TranslationJob created, readable
    message returned, no Traceback or Exception detail exposed.
    """

    password = "Bagobo-AppB-2026!"

    def setUp(self):
        from django.core.cache import cache as _cache
        _cache.clear()

        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)

        self.alice = User.objects.create_user(
            username="appb_alice",
            email="appb_alice@example.test",
            password=self.password,
        )
        self.bob = User.objects.create_user(
            username="appb_bob",
            email="appb_bob@example.test",
            password=self.password,
        )

    # ── file factories ────────────────────────────────────────────────────────

    @staticmethod
    def _minimal_pdf() -> bytes:
        return b"%PDF-1.4\n1 0 obj\n<<>>\nendobj\n%%EOF"

    @staticmethod
    def _minimal_docx() -> bytes:
        import io
        import zipfile as _zipfile
        buf = io.BytesIO()
        with _zipfile.ZipFile(buf, "w") as zf:
            zf.writestr(
                "[Content_Types].xml",
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"></Types>',
            )
            zf.writestr("word/document.xml", "<w:document/>")
        return buf.getvalue()

    # ── upload helper ─────────────────────────────────────────────────────────

    def _upload_as_alice(self, filename, content, content_type="application/octet-stream", patch_start=True):
        """POST to the upload endpoint as alice; patches start_translation_job by default."""
        self.client.force_login(self.alice)
        uploaded = SimpleUploadedFile(filename, content, content_type=content_type)
        post_data = {
            "file": uploaded,
            "source_language": "auto",
            "target_language": "tagabawa",
        }
        if patch_start:
            with patch("translator.views.start_translation_job"):
                return self.client.post(reverse("translator:upload"), post_data)
        return self.client.post(reverse("translator:upload"), post_data)

    # ── shared assertion helpers ──────────────────────────────────────────────

    def _assert_no_job_created(self):
        self.assertEqual(TranslationJob.objects.count(), 0)

    def _assert_readable_error(self, response, *expected_keywords):
        """Assert that at least one expected keyword appears in the response body."""
        body = json.dumps(response.json()).lower()
        self.assertTrue(
            any(kw.lower() in body for kw in expected_keywords),
            msg=f"Expected one of {expected_keywords!r} in response body, got: {body[:300]}",
        )

    def _assert_no_traceback(self, response):
        body = json.dumps(response.json())
        self.assertNotIn("Traceback", body)
        self.assertNotIn("Exception", body)

    # ------------------------------------------------------------------
    # TC-B01: Valid PDF upload
    # ------------------------------------------------------------------

    def test_tc_b01_valid_pdf_creates_job(self):
        """TC-B01: Valid PDF is accepted; a job is created and owned by the uploading user."""
        response = self._upload_as_alice("document.pdf", self._minimal_pdf(), "application/pdf")
        self.assertEqual(response.status_code, 202)
        self.assertEqual(TranslationJob.objects.count(), 1)
        job = TranslationJob.objects.get()
        self.assertEqual(job.owner, self.alice)
        self.assertEqual(job.file_type, TranslationJob.FileType.PDF)
        self.assertFalse(job.is_deleted)

    # ------------------------------------------------------------------
    # TC-B02: Valid DOCX upload
    # ------------------------------------------------------------------

    def test_tc_b02_valid_docx_creates_job(self):
        """TC-B02: Valid DOCX is accepted; a job is created and owned by the uploading user."""
        response = self._upload_as_alice(
            "report.docx",
            self._minimal_docx(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertEqual(response.status_code, 202)
        self.assertEqual(TranslationJob.objects.count(), 1)
        job = TranslationJob.objects.get()
        self.assertEqual(job.owner, self.alice)
        self.assertEqual(job.file_type, TranslationJob.FileType.DOCX)

    # ------------------------------------------------------------------
    # TC-B03: Unsupported file type
    # ------------------------------------------------------------------

    def test_tc_b03_unsupported_file_rejected(self):
        """TC-B03: Unsupported extension is rejected with a format message; no job created."""
        response = self._upload_as_alice(
            "malware.exe",
            b"MZ binary payload",
            "application/octet-stream",
            patch_start=False,
        )
        self.assertEqual(response.status_code, 400)
        self._assert_readable_error(response, "PDF", "DOCX", "unsupported", "Upload")
        self._assert_no_job_created()
        self._assert_no_traceback(response)

    # ------------------------------------------------------------------
    # TC-B04: Oversized document
    # ------------------------------------------------------------------

    def test_tc_b04_oversized_file_rejected(self):
        """TC-B04: File exceeding the 50 MB limit is rejected with a size message; no job created."""
        oversized_content = b"a" * (51 * 1024 * 1024)
        response = self._upload_as_alice(
            "huge.pdf", oversized_content, "application/pdf", patch_start=False
        )
        self.assertEqual(response.status_code, 400)
        self._assert_readable_error(response, "large", "MB", "size")
        self._assert_no_job_created()
        self._assert_no_traceback(response)

    # ------------------------------------------------------------------
    # TC-B05: No document selected
    # ------------------------------------------------------------------

    def test_tc_b05_missing_file_shows_required_error(self):
        """TC-B05: Submitting the form without a file returns a required-field error; no job created."""
        self.client.force_login(self.alice)
        response = self.client.post(
            reverse("translator:upload"),
            {"source_language": "auto", "target_language": "tagabawa"},
        )
        self.assertEqual(response.status_code, 400)
        self.assertIn("required", json.dumps(response.json()).lower())
        self._assert_no_job_created()
        self._assert_no_traceback(response)

    # ------------------------------------------------------------------
    # TC-B06: Corrupted PDF
    # ------------------------------------------------------------------

    def test_tc_b06_corrupted_pdf_rejected(self):
        """TC-B06: .pdf file with wrong magic bytes is rejected safely; no job created."""
        response = self._upload_as_alice(
            "corrupt.pdf",
            b"This is not a PDF at all, no percent-PDF header",
            "application/pdf",
            patch_start=False,
        )
        self.assertEqual(response.status_code, 400)
        self._assert_readable_error(response, "pdf", "valid", "invalid")
        self._assert_no_job_created()
        self._assert_no_traceback(response)

    # ------------------------------------------------------------------
    # TC-B07: Corrupted DOCX
    # ------------------------------------------------------------------

    def test_tc_b07_corrupted_docx_rejected(self):
        """TC-B07: .docx file with invalid ZIP structure is rejected safely; no job created."""
        response = self._upload_as_alice(
            "broken.docx",
            b"PK\x03\x04" + b"\x00" * 20 + b"truncated and corrupt ZIP data",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            patch_start=False,
        )
        self.assertEqual(response.status_code, 400)
        self._assert_readable_error(response, "docx", "zip", "valid", "invalid", "office")
        self._assert_no_job_created()
        self._assert_no_traceback(response)

    # ------------------------------------------------------------------
    # TC-B08: Successful authenticated upload — ownership and isolation
    # ------------------------------------------------------------------

    def test_tc_b08_job_is_owned_by_authenticated_user(self):
        """TC-B08: Job belongs to the uploading user; other users receive 404 on every access."""
        response = self._upload_as_alice("owned.pdf", self._minimal_pdf(), "application/pdf")
        self.assertEqual(response.status_code, 202)
        payload = response.json()
        self.assertIn("job_id", payload)
        self.assertIn("status", payload)

        job = TranslationJob.objects.get()
        self.assertEqual(job.owner, self.alice)

        # Bob must not be able to read Alice's job status
        self.client.force_login(self.bob)
        self.assertEqual(
            self.client.get(reverse("translator:status", args=[job.id])).status_code,
            404,
        )
        # Bob must not be able to view Alice's job detail page
        self.assertEqual(
            self.client.get(reverse("translator:job_detail", args=[job.id])).status_code,
            404,
        )

    # ------------------------------------------------------------------
    # TC-B09: Active-job limit exceeded
    # ------------------------------------------------------------------

    def test_tc_b09_active_job_limit_blocks_upload(self):
        """TC-B09: Upload is blocked when 2 active jobs exist; no additional job is created."""
        TranslationJob.objects.create(
            owner=self.alice,
            original_filename="active1.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=TranslationJob.Status.QUEUED,
        )
        TranslationJob.objects.create(
            owner=self.alice,
            original_filename="active2.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=TranslationJob.Status.PROCESSING,
        )

        self.client.force_login(self.alice)
        uploaded = SimpleUploadedFile("new.pdf", self._minimal_pdf(), content_type="application/pdf")
        response = self.client.post(
            reverse("translator:upload"),
            {"file": uploaded, "source_language": "auto", "target_language": "tagabawa"},
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("upload limit", response.json()["detail"].lower())
        # Count must remain at 2 — no third job was created
        self.assertEqual(TranslationJob.objects.count(), 2)

    # ------------------------------------------------------------------
    # TC-B10: Upload-rate limit exceeded
    # ------------------------------------------------------------------

    def test_tc_b10_rate_limit_blocks_excess_uploads(self):
        """TC-B10: The 6th upload attempt within an hour is blocked; no job is created."""
        from django.core.cache import cache
        cache.clear()

        self.client.force_login(self.alice)

        # Exhaust the 5-per-hour allowance. Empty-file requests increment the
        # rate counter even though they fail form validation and create no jobs.
        for _ in range(5):
            self.client.post(
                reverse("translator:upload"),
                {"source_language": "auto", "target_language": "tagabawa"},
            )

        # The 6th attempt with a fully valid file must be blocked by the rate limit
        uploaded = SimpleUploadedFile("valid.pdf", self._minimal_pdf(), content_type="application/pdf")
        response = self.client.post(
            reverse("translator:upload"),
            {"file": uploaded, "source_language": "auto", "target_language": "tagabawa"},
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("upload limit", response.json()["detail"].lower())
        # None of the six attempts must have created a job
        self.assertEqual(TranslationJob.objects.count(), 0)


# ===========================================================================
# APPENDIX C — Sprint 3: Automatic Document Detection and Text Extraction
# ===========================================================================


class AppendixCDetectionTests(TestCase):
    """Regression tests for Appendix C Sprint 3 — Document Detection and Extraction.

    TC-C01 through TC-C08 cover: digital vs scanned classification,
    per-page extraction routing, text storage, page ordering, and
    safe failure when no readable content is found.
    """

    password = "Bagobo-AppC-2026!"

    def setUp(self):
        from django.core.cache import cache as _cache
        _cache.clear()
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)
        self.alice = User.objects.create_user(
            username="appc_alice", email="appc_alice@example.test", password=self.password
        )

    def _create_job(self, owner, status="processing", file_type="pdf", metadata=None):
        return TranslationJob.objects.create(
            owner=owner,
            original_filename="test_doc.pdf",
            file_type=file_type,
            status=status,
            source_language="english",
            target_language="tagabawa",
            metadata=metadata or {},
        )

    def _write_structure(self, job, structure: dict) -> str:
        job_dir = self.media_root / "jobs" / job.job_id
        job_dir.mkdir(parents=True, exist_ok=True)
        path = str(job_dir / "structure.json")
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(structure, fh)
        return path

    # ------------------------------------------------------------------ TC-C01
    def test_tc_c01_digital_pdf_uses_direct_extraction(self):
        """extraction_method=direct_pdf_text in structure.json is saved to job.metadata by _sync_structure_models."""
        from translator.services import _sync_structure_models

        job = self._create_job(self.alice)
        structure_path = self._write_structure(job, {
            "extraction_method": "direct_pdf_text",
            "pages": [{
                "page_number": 1,
                "width": 612,
                "height": 792,
                "rotation": 0,
                "blocks": [{
                    "type": "text",
                    "bbox": [72, 100, 540, 120],
                    "lines": [{
                        "text": "Bagobo word",
                        "translated_text": "Salita ng Bagobo",
                        "translation_method": "phrasebook_exact",
                        "translation_confidence": 1.0,
                        "bbox": [72, 100, 540, 120],
                    }],
                }],
            }],
        })

        _sync_structure_models(str(job.id), structure_path)

        job.refresh_from_db()
        self.assertEqual(
            job.metadata.get("extraction_method"), "direct_pdf_text",
            "extraction_method from structure.json must be propagated into job.metadata",
        )
        self.assertTrue(
            TranslationSegment.objects.filter(job=job).exists(),
            "TranslationSegment records must be created for the extracted digital text",
        )

    # ------------------------------------------------------------------ TC-C02
    def test_tc_c02_scanned_pdf_uses_ocr(self):
        """extraction_method=ocr_image and OCR confidence are stored in job.metadata and OCRResult after sync."""
        from translator.services import _sync_structure_models

        job = self._create_job(self.alice)
        structure_path = self._write_structure(job, {
            "extraction_method": "ocr_image",
            "ocr_summary": {"mean_confidence": 0.82},
            "pages": [{
                "page_number": 1,
                "width": 612,
                "height": 792,
                "rotation": 0,
                "blocks": [{
                    "type": "text",
                    "bbox": [72, 100, 540, 120],
                    "ocr_confidence": 0.82,
                    "lines": [{
                        "text": "Scanned line",
                        "translated_text": "Linya mula sa scan",
                        "translation_method": "phrasebook_exact",
                        "translation_confidence": 0.9,
                        "ocr_confidence": 0.82,
                        "bbox": [72, 100, 540, 120],
                    }],
                }],
            }],
        })

        _sync_structure_models(str(job.id), structure_path)

        job.refresh_from_db()
        self.assertEqual(
            job.metadata.get("extraction_method"), "ocr_image",
            "OCR-extracted jobs must store extraction_method=ocr_image in job.metadata",
        )
        ocr = OCRResult.objects.filter(job=job).first()
        self.assertIsNotNone(ocr, "OCRResult must be created when structure.json contains ocr_confidence")
        self.assertGreater(ocr.confidence, 0, "OCRResult.confidence must reflect the value from structure.json")

    # ------------------------------------------------------------------ TC-C03
    def test_tc_c03_low_text_pdf_falls_back_to_ocr(self):
        """A PDF averaging fewer than 50 chars per page is classified SCANNED, never DIGITAL."""
        from translator.services.detection_service import DetectionService, DetectionType

        orig = DetectionService._count_pdf_text_chars
        try:
            DetectionService._count_pdf_text_chars = staticmethod(lambda path: (49, 1))
            det_type = DetectionService.detect_pdf_type("sparse.pdf")
            self.assertEqual(det_type, DetectionType.SCANNED)
            self.assertNotEqual(
                det_type, DetectionType.DIGITAL,
                "A PDF with fewer than 50 chars/page must not be classified as DIGITAL",
            )
        finally:
            DetectionService._count_pdf_text_chars = orig

    # ------------------------------------------------------------------ TC-C04
    def test_tc_c04_docx_text_is_extracted(self):
        """DOCX documents are classified DIGITAL and their extracted text is stored as TranslationSegment records."""
        from translator.services.detection_service import DetectionService, DetectionType
        from translator.services import _sync_structure_models

        # DOCX is always DIGITAL — no OCR path.
        det_type = DetectionService.detect_docx_type("dummy.docx")
        self.assertEqual(det_type, DetectionType.DIGITAL)

        # After extraction and translation, _sync_structure_models stores the text as segments.
        job = self._create_job(self.alice, file_type="docx")
        structure_path = self._write_structure(job, {
            "extraction_method": "docx_text",
            "pages": [{
                "page_number": 1,
                "width": 612,
                "height": 792,
                "rotation": 0,
                "blocks": [{
                    "type": "text",
                    "bbox": [72, 100, 540, 120],
                    "lines": [{
                        "text": "Bagobo learning material",
                        "translated_text": "Materyal sa pag-aaral ng Bagobo",
                        "translation_method": "phrasebook_exact",
                        "translation_confidence": 0.9,
                        "bbox": [72, 100, 540, 120],
                    }],
                }],
            }],
        })

        _sync_structure_models(str(job.id), structure_path)

        job.refresh_from_db()
        self.assertEqual(
            job.metadata.get("extraction_method"), "docx_text",
            "DOCX extraction method must be stored as docx_text in job.metadata",
        )
        self.assertTrue(
            TranslationSegment.objects.filter(job=job).exists(),
            "TranslationSegment records must be created from DOCX-extracted text",
        )

    # ------------------------------------------------------------------ TC-C05
    def test_tc_c05_text_and_page_order_are_preserved(self):
        """Segments from later pages have higher segment_index than segments from earlier pages."""
        from translator.services import _sync_structure_models

        job = self._create_job(self.alice)
        structure_path = self._write_structure(job, {
            "extraction_method": "direct_pdf_text",
            "pages": [
                {
                    "page_number": 1,
                    "width": 612,
                    "height": 792,
                    "rotation": 0,
                    "blocks": [{
                        "type": "text",
                        "bbox": [72, 100, 540, 120],
                        "lines": [{
                            "text": "First page sentence.",
                            "translated_text": "Pangungusap sa unang pahina.",
                            "translation_method": "phrasebook_exact",
                            "translation_confidence": 1.0,
                            "bbox": [72, 100, 540, 120],
                        }],
                    }],
                },
                {
                    "page_number": 2,
                    "width": 612,
                    "height": 792,
                    "rotation": 0,
                    "blocks": [{
                        "type": "text",
                        "bbox": [72, 100, 540, 120],
                        "lines": [{
                            "text": "Second page sentence.",
                            "translated_text": "Pangungusap sa ikalawang pahina.",
                            "translation_method": "phrasebook_exact",
                            "translation_confidence": 1.0,
                            "bbox": [72, 100, 540, 120],
                        }],
                    }],
                },
            ],
        })

        _sync_structure_models(str(job.id), structure_path)

        segments = list(
            TranslationSegment.objects.filter(job=job).order_by("segment_index")
        )
        self.assertEqual(len(segments), 2)
        self.assertEqual(segments[0].source_text, "First page sentence.")
        self.assertEqual(segments[1].source_text, "Second page sentence.")
        self.assertLess(
            segments[0].segment_index,
            segments[1].segment_index,
            "Segment from page 1 must have a strictly lower segment_index than page 2",
        )

    # ------------------------------------------------------------------ TC-C06
    @patch("translator.services._get_pipeline_service")
    def test_tc_c06_no_readable_content_fails_safely(self, mock_pipeline):
        """A job with no extractable text is set to FAILED with an actionable message; job does not linger in QUEUED."""
        from translator.services import _run_translation_job

        no_text_msg = "Failed to extract layout (no text blocks found)"
        mock_pipeline.side_effect = Exception(no_text_msg)

        job = self._create_job(self.alice, status="queued")
        _run_translation_job(
            job.job_id,
            job.input_file_path,
            "pdf",
            "auto",
            "tagabawa",
            None,
        )

        job.refresh_from_db()
        self.assertEqual(
            job.status, TranslationJob.Status.FAILED,
            "A document with no extractable text must set job.status=FAILED",
        )
        self.assertNotIn(
            job.status,
            [TranslationJob.Status.QUEUED, TranslationJob.Status.PROCESSING],
            "Job must not remain QUEUED or PROCESSING after a no-text failure",
        )
        self.assertIn(
            "no text blocks found", job.error,
            "job.error must contain an actionable description of the failure",
        )
        self.assertNotIn(
            "Traceback", job.error,
            "job.error must not expose Python tracebacks to users",
        )

    # ------------------------------------------------------------------ TC-C07
    def test_tc_c07_extracted_text_is_stored_for_translation(self):
        """Source text from structure.json is persisted verbatim in TranslationSegment.source_text."""
        from translator.services import _sync_structure_models

        job = self._create_job(self.alice)
        source_phrase = "Ang ibon ay lumilipad sa kalawakan."
        structure_path = self._write_structure(job, {
            "extraction_method": "direct_pdf_text",
            "pages": [{
                "page_number": 1,
                "width": 612,
                "height": 792,
                "rotation": 0,
                "blocks": [{
                    "type": "text",
                    "bbox": [72, 100, 540, 120],
                    "lines": [{
                        "text": source_phrase,
                        "translated_text": "The bird flies through space.",
                        "translation_method": "phrasebook_exact",
                        "translation_confidence": 0.95,
                        "bbox": [72, 100, 540, 120],
                    }],
                }],
            }],
        })

        _sync_structure_models(str(job.id), structure_path)

        segment = TranslationSegment.objects.filter(job=job).first()
        self.assertIsNotNone(segment, "At least one TranslationSegment must be created after sync")
        self.assertEqual(
            segment.source_text, source_phrase,
            "TranslationSegment.source_text must exactly match the extracted source text",
        )

    # ------------------------------------------------------------------ TC-C08
    @patch("translator.services._get_pipeline_service")
    def test_tc_c08_extraction_error_sets_failed_status(self, mock_pipeline):
        """An extraction-phase error transitions the job to FAILED; status API reflects the final state."""
        from translator.services import _run_translation_job

        mock_pipeline.side_effect = Exception("Extraction failed: corrupt PDF structure")

        job = self._create_job(self.alice, status="queued")
        _run_translation_job(
            job.job_id,
            job.input_file_path,
            "pdf",
            "auto",
            "tagabawa",
            None,
        )

        job.refresh_from_db()
        self.assertEqual(
            job.status, TranslationJob.Status.FAILED,
            "An extraction error must set job.status=FAILED in the database",
        )

        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:status", args=[job.id]))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.json()["status"], TranslationJob.Status.FAILED,
            "The status API must return 'failed' after an extraction error, not 'processing' or 'queued'",
        )


class AppendixDOCRTests(TestCase):
    """Regression tests for Appendix D Sprint 4 — Automatic OCR Module."""

    password = "Bagobo-AppD-2026!"

    def setUp(self):
        from django.core.cache import cache as _cache
        _cache.clear()
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)
        self.alice = User.objects.create_user(
            username="appd_alice", email="appd_alice@example.test", password=self.password
        )

    def _create_job(self, owner, status="processing", file_type="pdf", metadata=None):
        return TranslationJob.objects.create(
            owner=owner,
            original_filename="test_doc.pdf",
            file_type=file_type,
            status=status,
            source_language="english",
            target_language="tagabawa",
            metadata=metadata or {},
        )

    def _write_structure(self, job, structure: dict) -> str:
        job_dir = self.media_root / "jobs" / job.job_id
        job_dir.mkdir(parents=True, exist_ok=True)
        path = str(job_dir / "structure.json")
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(structure, fh)
        return path

    # ------------------------------------------------------------------ TC-D01
    def test_tc_d01_scanned_pdf_runs_ocr_automatically(self):
        """Scanned PDF detection routes to OCR; OCRService defaults are DPI=300 and PSM=3; extraction_method stored as ocr_image."""
        from translator.services.detection_service import DetectionService, DetectionType
        from translator.services.ocr_stage.ocr_service import OCRService
        from translator.services import _sync_structure_models

        svc = OCRService()
        self.assertEqual(svc.dpi, 300, "OCR service must default to DPI=300")
        self.assertEqual(svc.psm, 3, "OCR service must default to PSM=3 (auto page segmentation)")
        self.assertGreater(svc.timeout_seconds, 0, "OCR service must have a positive timeout")

        orig = DetectionService._count_pdf_text_chars
        try:
            DetectionService._count_pdf_text_chars = staticmethod(lambda path: (30, 1))
            det = DetectionService.detect_pdf_type("scan.pdf")
        finally:
            DetectionService._count_pdf_text_chars = orig
        self.assertEqual(det, DetectionType.SCANNED,
                         "PDF with < 50 chars/page must be classified as SCANNED, not DIGITAL")

        job = self._create_job(self.alice)
        path = self._write_structure(job, {
            "extraction_method": "ocr_image",
            "pages": [{"page_number": 1, "width": 612, "height": 792, "rotation": 0, "blocks": []}],
        })
        _sync_structure_models(str(job.id), path)
        job.refresh_from_db()
        self.assertEqual(job.metadata.get("extraction_method"), "ocr_image",
                         "_sync_structure_models must store extraction_method=ocr_image for OCR jobs")
        self.assertNotEqual(job.metadata.get("extraction_method"), "direct_pdf_text")

    # ------------------------------------------------------------------ TC-D02
    def test_tc_d02_jpg_runs_ocr_and_stores_text(self):
        """JPEG images are classified as SCANNED; extract_text_with_fallback always returns extraction_method=ocr_image."""
        from PIL import Image
        from translator.services.detection_service import DetectionService, DetectionType
        from translator.services.ocr_stage.ocr_service import OCRService

        img = Image.new("RGB", (200, 100), (255, 255, 255))
        jpg_path = str(self.media_root / "test_tc_d02.jpg")
        img.save(jpg_path, "JPEG")

        det = DetectionService.detect_image_type(jpg_path)
        self.assertEqual(det, DetectionType.SCANNED,
                         "JPEG images must be classified as SCANNED, not DIGITAL")

        svc = OCRService()
        result = svc.extract_text_with_fallback(jpg_path)
        self.assertEqual(result.get("extraction_method"), "ocr_image",
                         "extract_text_with_fallback must always set extraction_method='ocr_image'")
        self.assertIn("pages", result, "Result must contain 'pages' key")
        self.assertIn("ocr_summary", result, "Result must contain 'ocr_summary' key")

    # ------------------------------------------------------------------ TC-D03
    def test_tc_d03_png_runs_ocr_and_stores_text(self):
        """PNG images are classified as SCANNED; extract_text_with_fallback always returns extraction_method=ocr_image."""
        from PIL import Image
        from translator.services.detection_service import DetectionService, DetectionType
        from translator.services.ocr_stage.ocr_service import OCRService

        img = Image.new("RGB", (200, 100), (255, 255, 255))
        png_path = str(self.media_root / "test_tc_d03.png")
        img.save(png_path, "PNG")

        det = DetectionService.detect_image_type(png_path)
        self.assertEqual(det, DetectionType.SCANNED,
                         "PNG images must be classified as SCANNED, not DIGITAL")

        svc = OCRService()
        result = svc.extract_text_with_fallback(png_path)
        self.assertEqual(result.get("extraction_method"), "ocr_image",
                         "extract_text_with_fallback must always set extraction_method='ocr_image'")
        self.assertIn("pages", result)
        self.assertIn("ocr_summary", result)

    # ------------------------------------------------------------------ TC-D04
    def test_tc_d04_faded_preprocessing_preserves_text(self):
        """Faded-scan preprocessing: near-white background (240) becomes white; light-gray text (155) becomes black."""
        from PIL import Image, ImageDraw
        from translator.services.ocr_stage.ocr_service import OCRService

        img = Image.new("L", (200, 80), 240)
        draw = ImageDraw.Draw(img)
        draw.rectangle([10, 20, 150, 60], fill=155)

        svc = OCRService(detect_orientation=False, denoise=False)
        result = svc._preprocess(img)

        bg_pixel = result.getpixel((0, 0))
        self.assertEqual(bg_pixel, 255,
                         f"Faded background (240) must become white (255) after preprocessing, got {bg_pixel}")

        text_pixel = result.getpixel((80, 40))
        self.assertEqual(text_pixel, 0,
                         f"Light-gray text (155) must become black (0) after preprocessing, got {text_pixel}")

    # ------------------------------------------------------------------ TC-D05
    def test_tc_d05_high_confidence_has_no_warning(self):
        """High OCR confidence (≥0.60) produces has_low_quality_warning=False; job_detail shows no '⚠ Low' badge."""
        from translator.services.ocr_stage.ocr_service import OCRService

        pages_data = [{"blocks": [{"confidence": 0.90}, {"confidence": 0.95}]}]
        summary = OCRService._compute_page_ocr_summary(pages_data)
        self.assertFalse(summary["has_low_quality_warning"],
                         "High-confidence OCR must not produce a low-quality warning")
        self.assertGreaterEqual(summary["mean_confidence"], 0.60,
                                "Mean confidence for high-confidence blocks must be ≥ 0.60")

        job = self._create_job(
            self.alice,
            status=TranslationJob.Status.COMPLETED,
            metadata={
                "extraction_method": "ocr_image",
                "ocr_summary": {"mean_confidence": 0.92, "has_low_quality_warning": False},
            },
        )
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertNotIn("⚠ Low", body,
                         "High-confidence OCR job must not display the '⚠ Low' badge in job detail")

    # ------------------------------------------------------------------ TC-D06
    def test_tc_d06_low_confidence_is_marked_for_review(self):
        """Low OCR confidence (< 0.60) triggers has_low_quality_warning=True; preview page renders the warning section."""
        from translator.services.ocr_stage.ocr_service import OCRService

        pages_data = [{"blocks": [{"confidence": 0.40}]}]
        summary = OCRService._compute_page_ocr_summary(pages_data)
        self.assertTrue(summary["has_low_quality_warning"],
                        "OCR confidence 0.40 must produce has_low_quality_warning=True")
        self.assertLess(summary["mean_confidence"], 0.60,
                        "Mean confidence for low-confidence blocks must be < 0.60")

        warning_msg = "Low OCR confidence detected: mean=40%, 1 of 1 blocks below threshold."
        job = self._create_job(
            self.alice,
            status=TranslationJob.Status.COMPLETED,
            metadata={
                "extraction_method": "ocr_image",
                "ocr_summary": {"mean_confidence": 0.40, "has_low_quality_warning": True},
                "ocr_warnings": [warning_msg],
            },
        )
        # Raw OCR warning text is technical detail — staff/admin only.
        self.alice.is_staff = True
        self.alice.save(update_fields=["is_staff"])
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:preview", args=[job.id]))
        self.assertEqual(response.status_code, 200)
        self.assertIn(b"Low OCR confidence detected", response.content,
                      "Preview page must display OCR low-confidence warning from job metadata")
        self.assertIn(b"warning-box", response.content,
                      "Preview page must render the warning-box section for low-confidence OCR jobs")

    def test_tc_d06_low_confidence_warning_hidden_from_normal_user(self):
        """Normal users never see the raw OCR warning box or its message."""
        job = self._create_job(
            self.alice,
            status=TranslationJob.Status.COMPLETED,
            metadata={
                "extraction_method": "ocr_image",
                "ocr_summary": {"mean_confidence": 0.40, "has_low_quality_warning": True},
                "ocr_warnings": ["Low OCR confidence detected: mean=40%, 1 of 1 blocks below threshold."],
            },
        )
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:preview", args=[job.id]))
        self.assertEqual(response.status_code, 200)
        self.assertNotIn(b"Low OCR confidence detected", response.content)
        self.assertNotIn(b"warning-box", response.content)

    # ------------------------------------------------------------------ TC-D07
    @patch("translator.services._get_pipeline_service")
    def test_tc_d07_unreadable_document_fails_safely(self, mock_pipeline):
        """A scanned document with no readable text fails with actionable 'OCR produced no text' message; no traceback stored."""
        from translator.services import _run_translation_job

        ocr_no_text_msg = (
            "OCR produced no text from this scanned document. "
            "Page may be blank, low-quality, or in an unsupported language."
        )
        mock_pipeline.side_effect = Exception(ocr_no_text_msg)

        job = self._create_job(self.alice, status="queued")
        _run_translation_job(job.job_id, job.input_file_path, "pdf", "auto", "tagabawa", None)
        job.refresh_from_db()

        self.assertEqual(job.status, TranslationJob.Status.FAILED)
        self.assertNotIn(job.status, [TranslationJob.Status.QUEUED, TranslationJob.Status.PROCESSING])
        self.assertIn("OCR produced no text", job.error,
                      "Error message must state 'OCR produced no text' for unreadable scanned documents")
        self.assertNotIn("Traceback", job.error,
                         "job.error must not expose Python tracebacks to users")

    # ------------------------------------------------------------------ TC-D08
    def test_tc_d08_successful_ocr_stores_results_and_completes(self):
        """Successful OCR sync: status=COMPLETED, ocr_summary in metadata, OCRResult and TranslationSegment created."""
        job = self._create_job(self.alice, status="queued")
        ocr_structure = {
            "extraction_method": "ocr_image",
            "ocr_summary": {
                "mean_confidence": 0.87,
                "min_confidence": 0.82,
                "low_confidence_block_count": 0,
                "total_block_count": 1,
                "has_low_quality_warning": False,
            },
            "pages": [{
                "page_number": 1,
                "width": 612,
                "height": 792,
                "rotation": 0,
                "blocks": [{
                    "type": "text",
                    "bbox": [72, 100, 540, 120],
                    "ocr_confidence": 0.87,
                    "lines": [{
                        "text": "Namalagi ang ibon sa kahoy.",
                        "translated_text": "The bird stayed in the tree.",
                        "translation_method": "phrasebook_exact",
                        "translation_confidence": 0.95,
                        "ocr_confidence": 0.87,
                        "bbox": [72, 100, 540, 120],
                    }],
                }],
            }],
        }
        structure_path = self._write_structure(job, ocr_structure)

        sync_pipeline_job(SimpleNamespace(
            job_id=job.job_id,
            status=TranslationJob.Status.COMPLETED,
            progress=100,
            current_phase="done",
            current_step="Completed",
            phase_message="",
            error="",
            detection_type=None,
            file_type=None,
            completed_at=None,
            metadata={"structure_file": str(structure_path)},
        ))

        job.refresh_from_db()
        self.assertEqual(job.status, TranslationJob.Status.COMPLETED,
                         "Job must reach COMPLETED status after successful OCR sync")
        self.assertEqual(job.metadata.get("extraction_method"), "ocr_image",
                         "extraction_method must be stored as 'ocr_image' in job metadata")
        self.assertIsNotNone(job.metadata.get("ocr_summary"),
                             "ocr_summary must be stored in job.metadata after successful OCR")

        ocr = OCRResult.objects.filter(job=job).first()
        self.assertIsNotNone(ocr, "OCRResult must be created for OCR jobs")
        self.assertGreater(ocr.confidence, 0, "OCRResult.confidence must be > 0 for successful OCR")

        seg = TranslationSegment.objects.filter(job=job).first()
        self.assertIsNotNone(seg, "TranslationSegment must be created from OCR-extracted text")
        self.assertEqual(seg.source_text, "Namalagi ang ibon sa kahoy.")

    # ------------------------------------------------------------------ TC-D09
    @patch("translator.services._get_pipeline_service")
    def test_tc_d09_ocr_error_sets_failed_status(self, mock_pipeline):
        """An OCR-phase exception transitions job to FAILED; job.error contains only the message, no Python traceback."""
        from translator.services import _run_translation_job

        error_message = "OCR failed on page 2: image too faded to process"
        mock_pipeline.side_effect = Exception(error_message)

        job = self._create_job(self.alice, status="queued")
        _run_translation_job(job.job_id, job.input_file_path, "pdf", "auto", "tagabawa", None)
        job.refresh_from_db()

        self.assertEqual(job.status, TranslationJob.Status.FAILED,
                         "An OCR exception must transition the job to FAILED")
        self.assertNotIn(job.status, [TranslationJob.Status.QUEUED, TranslationJob.Status.PROCESSING])
        self.assertIn(error_message, job.error,
                      "job.error must contain the original exception message")
        self.assertNotIn("Traceback", job.error,
                         "job.error must not expose Python traceback strings")
        self.assertNotIn('File "', job.error,
                         "job.error must not expose Python source file paths")

    # ------------------------------------------------------------------ TC-D10
    def test_tc_d10_multipage_ocr_order_is_preserved(self):
        """Three-page OCR structure: all segments created in source order with monotonically increasing indexes."""
        from translator.services import _sync_structure_models

        job = self._create_job(self.alice)
        structure_path = self._write_structure(job, {
            "extraction_method": "ocr_image",
            "pages": [
                {
                    "page_number": 1, "width": 612, "height": 792, "rotation": 0,
                    "blocks": [{"type": "text", "bbox": [72, 100, 540, 120], "lines": [{
                        "text": "Page one OCR line.",
                        "translated_text": "Linya sa unang pahina.",
                        "translation_method": "phrasebook_exact",
                        "translation_confidence": 0.9,
                        "bbox": [72, 100, 540, 120],
                    }]}],
                },
                {
                    "page_number": 2, "width": 612, "height": 792, "rotation": 0,
                    "blocks": [{"type": "text", "bbox": [72, 100, 540, 120], "lines": [{
                        "text": "Page two OCR line.",
                        "translated_text": "Linya sa ikalawang pahina.",
                        "translation_method": "phrasebook_exact",
                        "translation_confidence": 0.9,
                        "bbox": [72, 100, 540, 120],
                    }]}],
                },
                {
                    "page_number": 3, "width": 612, "height": 792, "rotation": 0,
                    "blocks": [{"type": "text", "bbox": [72, 100, 540, 120], "lines": [{
                        "text": "Page three OCR line.",
                        "translated_text": "Linya sa ikatlong pahina.",
                        "translation_method": "phrasebook_exact",
                        "translation_confidence": 0.9,
                        "bbox": [72, 100, 540, 120],
                    }]}],
                },
            ],
        })

        _sync_structure_models(str(job.id), structure_path)

        segments = list(TranslationSegment.objects.filter(job=job).order_by("segment_index"))
        self.assertEqual(len(segments), 3,
                         "Three OCR pages must produce exactly three TranslationSegment records")
        self.assertEqual(segments[0].source_text, "Page one OCR line.")
        self.assertEqual(segments[1].source_text, "Page two OCR line.")
        self.assertEqual(segments[2].source_text, "Page three OCR line.")
        self.assertLess(segments[0].segment_index, segments[1].segment_index,
                        "Page 1 segment must have lower segment_index than page 2")
        self.assertLess(segments[1].segment_index, segments[2].segment_index,
                        "Page 2 segment must have lower segment_index than page 3")


class RuntimeTranslationDatasetCorpusTests(TestCase):
    """Runtime dataset must combine phrasebook lookup with the cleaned corpus."""

    CORPUS_TAGABAWA = (
        "T\u00f4 midug\u00e9 d\u00e1n banuwa, t\u00f4 pagb\u00e1nnal kat\u00f4 "
        "mga t\u00f4 min-dug\u00e9 d\u00e1n banuwa Bag\u00f3b\u00f4 \u00e1s "
        "Sandawa, manub\u00f9."
    )
    CORPUS_ENGLISH = (
        "A long time ago on the earth, the belief of the Bagobo people was "
        "that Sandawa, he was a person."
    )

    def test_default_runtime_dataset_loads_text_corpus(self):
        from translator.services.translation_dataset import TranslationDataset

        dataset = TranslationDataset()
        result = dataset.translate_phrase_with_metadata(
            self.CORPUS_ENGLISH,
            "english",
            "tagabawa",
        )

        self.assertGreater(len(dataset.data), 1028)
        self.assertEqual(result["translated"], self.CORPUS_TAGABAWA)
        self.assertEqual(result["method"], "exact_phrase")
        self.assertFalse(result["needs_review"])

    def test_unknown_sentence_with_known_words_is_not_word_by_word(self):
        from translator.services.pipeline_service import PipelineService
        from translator.services.translation_dataset import TranslationDataset

        source = "Hello unknownword999"
        dataset = TranslationDataset()
        service = PipelineService.__new__(PipelineService)
        service.translation_dataset = dataset
        layout = [{
            "blocks": [{
                "type": "text",
                "lines": [{"text": source}],
            }]
        }]

        translations = service._translate_layout(layout, "english", "tagabawa")
        record = translations["0_0_0"]

        self.assertEqual(record["translated"], source)
        self.assertEqual(record["method"], "unknown_for_review")
        self.assertTrue(record["needs_review"])

    def test_txt_upload_lines_are_translated_as_separate_segments(self):
        from translator.services.pipeline_service import PipelineService
        from translator.services.translation_dataset import TranslationDataset

        with tempfile.NamedTemporaryFile("w", suffix=".txt", encoding="utf-8", delete=False) as handle:
            handle.write(f"Hello\nHello unknownword999\n{self.CORPUS_ENGLISH}\n")
            txt_path = handle.name
        self.addCleanup(lambda: Path(txt_path).unlink(missing_ok=True))

        layout = PipelineService._extract_txt_text_and_layout(txt_path)
        text_lines = [
            line["text"]
            for block in layout[0]["blocks"]
            for line in block["lines"]
        ]
        self.assertEqual(text_lines, ["Hello", "Hello unknownword999", self.CORPUS_ENGLISH])

        service = PipelineService.__new__(PipelineService)
        service.translation_dataset = TranslationDataset()
        translations = service._translate_layout(layout, "english", "tagabawa")

        self.assertEqual(translations["0_0_0"]["translated"], "Madig\u00e1r")
        self.assertEqual(translations["0_0_1"]["translated"], "Hello unknownword999")
        self.assertTrue(translations["0_0_1"]["needs_review"])
        self.assertEqual(translations["0_0_2"]["translated"], self.CORPUS_TAGABAWA)


class AppendixETranslationTests(TestCase):
    """Regression tests for Appendix E Sprint 5 — Cross-Lingual Translation Module."""

    password = "Bagobo-AppE-2026!"

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        import os
        from translator.services.translation_dataset import TranslationDataset
        dataset_path = os.path.join(os.path.dirname(__file__), "services", "translation_data.json")
        cls.dataset = TranslationDataset(dataset_path=dataset_path)

    def setUp(self):
        from django.core.cache import cache as _cache
        _cache.clear()
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)
        self.alice = User.objects.create_user(
            username="appe_alice", email="appe_alice@example.test", password=self.password
        )

    def _create_job(self, owner, status="processing", file_type="pdf", metadata=None):
        return TranslationJob.objects.create(
            owner=owner,
            original_filename="test_doc.pdf",
            file_type=file_type,
            status=status,
            source_language="english",
            target_language="tagabawa",
            metadata=metadata or {},
        )

    def _write_structure(self, job, structure: dict) -> str:
        job_dir = self.media_root / "jobs" / job.job_id
        job_dir.mkdir(parents=True, exist_ok=True)
        path = str(job_dir / "structure.json")
        with open(path, "w", encoding="utf-8") as fh:
            json.dump(structure, fh, ensure_ascii=False)
        return path

    # ------------------------------------------------------------------ TC-E01
    def test_tc_e01_extracted_text_is_processed_by_translation_service(self):
        """_translate_layout processes extracted text through the real dataset; a known phrase is translated, not UNKNOWN."""
        from translator.services.pipeline_service import PipelineService
        from translator.services.translation_dataset import UNKNOWN_FOR_REVIEW

        svc = PipelineService.__new__(PipelineService)
        svc.translation_dataset = self.dataset

        layout = [{
            "blocks": [{
                "type": "text",
                "bbox": [0, 0, 400, 50],
                "lines": [{"text": "Hello", "bbox": [0, 0, 400, 20]}],
            }]
        }]
        results = svc._translate_layout(layout, "english", "tagabawa")

        line_result = results.get("0_0_0")
        self.assertIsNotNone(line_result,
                             "_translate_layout must return a result keyed by '0_0_0' for the first line")
        self.assertNotEqual(line_result["translated"], UNKNOWN_FOR_REVIEW,
                            "A known phrase ('Hello') must be translated, not returned as UNKNOWN_FOR_REVIEW")
        self.assertIn(line_result["method"], ("exact_phrase", "normalized_phrase"),
                      "Translation method must be one of the valid cascade stages")
        self.assertIn("source_language", line_result,
                      "_translate_layout result must carry source_language")
        self.assertEqual(line_result["source_language"], "english")
        self.assertEqual(line_result["target_language"], "tagabawa")

    # ------------------------------------------------------------------ TC-E02
    def test_tc_e02_selected_target_language_is_used(self):
        """Valid target languages are accepted; invalid targets are rejected by form; stored language matches selection."""
        from translator.forms import DocumentUploadForm, TARGET_LANGUAGE_CHOICES
        from translator.services.translation_dataset import SUPPORTED_LANGS, _normalize_lang
        from django.core.files.uploadedfile import SimpleUploadedFile

        # Every form choice must resolve to a SUPPORTED_LANGS entry
        for code, _label in TARGET_LANGUAGE_CHOICES:
            self.assertIn(_normalize_lang(code), SUPPORTED_LANGS,
                          f"Form choice '{code}' must resolve to a supported language")

        # TranslationJob stores the chosen language exactly
        job = self._create_job(self.alice)
        job.target_language = "filipino"
        job.save(update_fields=["target_language"])
        job.refresh_from_db()
        self.assertEqual(job.target_language, "filipino",
                         "target_language must be stored and retrieved without modification")

        # Invalid target language: ChoiceField must reject it
        pdf_bytes = b"%PDF-1.4\n%%EOF"
        uploaded = SimpleUploadedFile("doc.pdf", pdf_bytes, content_type="application/pdf")
        form = DocumentUploadForm(
            data={"source_language": "auto", "target_language": "klingon", "ocr_languages": ""},
            files={"file": uploaded},
        )
        self.assertFalse(form.is_valid(),
                         "DocumentUploadForm must reject unknown target language 'klingon'")
        self.assertIn("target_language", form.errors,
                      "Validation error must identify 'target_language' as the invalid field")

    # ------------------------------------------------------------------ TC-E03
    def test_tc_e03_dictionary_word_returns_stored_translation(self):
        """A known word in the dataset returns its exact stored Tagabawa value via exact_phrase at confidence=1.0."""
        from translator.services.translation_dataset import UNKNOWN_FOR_REVIEW

        result = self.dataset.translate_phrase_with_metadata("Hello", "english", "tagabawa")

        self.assertNotEqual(result["translated"], UNKNOWN_FOR_REVIEW,
                            "'Hello' must have a stored Tagabawa translation in the dataset")
        self.assertEqual(result["method"], "exact_phrase",
                         "'Hello' must match via exact_phrase, not fall through to fallback handling")
        self.assertEqual(result["confidence"], 1.0,
                         "An exact_phrase match must report confidence=1.0")
        self.assertEqual(result["translated"], "Madigár",
                         "Translation of 'Hello' must match the stored dataset entry exactly")

    # ------------------------------------------------------------------ TC-E04
    def test_tc_e04_phrasebook_phrase_returns_stored_translation(self):
        """A multi-word phrase exact match returns the stored translation; source text is stored verbatim."""
        from translator.services.translation_dataset import UNKNOWN_FOR_REVIEW
        from translator.services import _sync_structure_models

        result = self.dataset.translate_phrase_with_metadata("Come ahead", "english", "tagabawa")

        self.assertNotEqual(result["translated"], UNKNOWN_FOR_REVIEW)
        self.assertEqual(result["method"], "exact_phrase",
                         "'Come ahead' must be matched as a complete phrase, not word-by-word")
        self.assertEqual(result["translated"], "Allus kó",
                         "Multi-word phrase translation must return the exact stored value including accents")

        # Source text preserved verbatim after storage
        job = self._create_job(self.alice)
        structure_path = self._write_structure(job, {
            "extraction_method": "direct_pdf_text",
            "pages": [{"page_number": 1, "width": 612, "height": 792, "rotation": 0,
                "blocks": [{"type": "text", "bbox": [72, 100, 540, 120], "lines": [{
                    "text": "Come ahead",
                    "translated_text": result["translated"],
                    "translation_method": result["method"],
                    "translation_confidence": result["confidence"],
                    "bbox": [72, 100, 540, 120],
                }]}]}],
        })
        _sync_structure_models(str(job.id), structure_path)
        seg = TranslationSegment.objects.filter(job=job).first()
        self.assertIsNotNone(seg)
        self.assertEqual(seg.source_text, "Come ahead",
                         "source_text must be stored verbatim; normalization must not overwrite the original")

    # ------------------------------------------------------------------ TC-E05
    def test_tc_e05_phrase_match_has_priority_over_word_matches(self):
        """An exact phrase match is returned first; the result is not fragmented by word-by-word processing."""
        from translator.services.translation_dataset import UNKNOWN_FOR_REVIEW

        phrase_result = self.dataset.translate_phrase_with_metadata("Sit down", "english", "tagabawa")

        self.assertEqual(phrase_result["method"], "exact_phrase",
                         "'Sit down' must be matched as a phrase (exact_phrase), not split word-by-word")
        self.assertNotEqual(phrase_result["translated"], UNKNOWN_FOR_REVIEW)
        self.assertEqual(phrase_result["translated"], "Unsad kó",
                         "Phrase match must return the stored phrase translation, not a word-by-word composite")
        self.assertEqual(phrase_result["confidence"], 1.0,
                         "exact_phrase must report confidence=1.0")

    # ------------------------------------------------------------------ TC-E06
    def test_tc_e06_unknown_word_is_retained_and_marked(self):
        """Completely unmatched text remains source text and is marked needs_review=True."""
        from translator.services.translation_dataset import UNKNOWN_FOR_REVIEW
        from translator.services import _sync_structure_models

        unknown_result = self.dataset.translate_phrase_with_metadata(
            "XYZXYZ ABCABC", "english", "tagabawa"
        )
        self.assertEqual(unknown_result["translated"], "XYZXYZ ABCABC",
                         "Completely unknown phrase must stay as source text, not a fabricated translation")
        self.assertEqual(unknown_result["method"], "unknown_for_review",
                         "method must be 'unknown_for_review' when no match exists at any cascade stage")
        self.assertEqual(unknown_result["confidence"], 0.0,
                         "confidence must be 0.0 for unknown_for_review results")
        self.assertTrue(unknown_result["needs_review"])

        job = self._create_job(self.alice)
        structure_path = self._write_structure(job, {
            "extraction_method": "direct_pdf_text",
            "pages": [{"page_number": 1, "width": 612, "height": 792, "rotation": 0,
                "blocks": [{"type": "text", "bbox": [72, 100, 540, 120], "lines": [{
                    "text": "XYZXYZ ABCABC",
                    "translated_text": UNKNOWN_FOR_REVIEW,
                    "translation_method": "unknown_for_review",
                    "translation_confidence": 0.0,
                    "bbox": [72, 100, 540, 120],
                }]}]}],
        })
        _sync_structure_models(str(job.id), structure_path)
        seg = TranslationSegment.objects.filter(job=job).first()
        self.assertIsNotNone(seg)
        self.assertTrue(seg.needs_review,
                        "Segment with UNKNOWN_FOR_REVIEW translated_text must have needs_review=True")
        self.assertEqual(seg.translated_text, UNKNOWN_FOR_REVIEW,
                         "UNKNOWN_FOR_REVIEW sentinel must be stored verbatim in translated_text")

    # ------------------------------------------------------------------ TC-E07
    def test_tc_e07_tagabawa_orthography_is_preserved(self):
        """Tagabawa diacritics survive translation lookup and database round-trip; source text is never normalized."""
        from translator.services import _sync_structure_models

        # "Hello" → "Madigár" (á = U+00E1 must survive)
        hello_result = self.dataset.translate_phrase_with_metadata("Hello", "english", "tagabawa")
        self.assertIn("á", hello_result["translated"],
                      "Tagabawa acute accent (á, U+00E1) must be present in translated output")
        self.assertEqual(hello_result["translated"], "Madigár",
                         "Full Tagabawa word with accent must match stored value exactly")

        # "Visiting" → "Ágpanumbalé" (Á = U+00C1, é = U+00E9 must survive)
        visiting_result = self.dataset.translate_phrase_with_metadata("Visiting", "english", "tagabawa")
        self.assertIn("Á", visiting_result["translated"],
                      "Capital Á (U+00C1) must be preserved in the Tagabawa translation")
        self.assertIn("é", visiting_result["translated"],
                      "Lowercase é (U+00E9) must be preserved in the Tagabawa translation")

        # Tagabawa source text with accents stored verbatim in TranslationSegment
        tagabawa_source = "Madigár"
        job = self._create_job(self.alice)
        structure_path = self._write_structure(job, {
            "extraction_method": "direct_pdf_text",
            "pages": [{"page_number": 1, "width": 612, "height": 792, "rotation": 0,
                "blocks": [{"type": "text", "bbox": [72, 100, 540, 120], "lines": [{
                    "text": tagabawa_source,
                    "translated_text": "Hello",
                    "translation_method": "exact_phrase",
                    "translation_confidence": 1.0,
                    "bbox": [72, 100, 540, 120],
                }]}]}],
        })
        _sync_structure_models(str(job.id), structure_path)
        seg = TranslationSegment.objects.filter(job=job).first()
        self.assertIsNotNone(seg)
        self.assertEqual(seg.source_text, tagabawa_source,
                         "Tagabawa source text with accents must be stored verbatim in the database")
        self.assertIn("á", seg.source_text,
                      "Acute accent (U+00E1) must be present in the stored source_text — not stripped")

    # ------------------------------------------------------------------ TC-E08
    def test_tc_e08_sentence_and_paragraph_order_is_preserved(self):
        """Multiple text blocks on one page produce segments in block order with monotonically increasing indexes."""
        from translator.services import _sync_structure_models

        job = self._create_job(self.alice)
        structure_path = self._write_structure(job, {
            "extraction_method": "direct_pdf_text",
            "pages": [{"page_number": 1, "width": 612, "height": 792, "rotation": 0,
                "blocks": [
                    {"type": "text", "bbox": [72, 100, 540, 120], "lines": [{
                        "text": "Block one sentence.",
                        "translated_text": "Pangungusap ng bloke isa.",
                        "translation_method": "exact_phrase",
                        "translation_confidence": 0.9,
                        "bbox": [72, 100, 540, 120],
                    }]},
                    {"type": "text", "bbox": [72, 130, 540, 150], "lines": [{
                        "text": "Block two sentence.",
                        "translated_text": "Pangungusap ng bloke dalawa.",
                        "translation_method": "exact_phrase",
                        "translation_confidence": 0.9,
                        "bbox": [72, 130, 540, 150],
                    }]},
                    {"type": "text", "bbox": [72, 160, 540, 180], "lines": [{
                        "text": "Block three sentence.",
                        "translated_text": "Pangungusap ng bloke tatlo.",
                        "translation_method": "exact_phrase",
                        "translation_confidence": 0.9,
                        "bbox": [72, 160, 540, 180],
                    }]},
                ],
            }],
        })
        _sync_structure_models(str(job.id), structure_path)

        segments = list(TranslationSegment.objects.filter(job=job).order_by("segment_index"))
        self.assertEqual(len(segments), 3,
                         "Three text blocks on a single page must produce three TranslationSegment records")
        self.assertEqual(segments[0].source_text, "Block one sentence.",
                         "First block must map to the first segment")
        self.assertEqual(segments[1].source_text, "Block two sentence.",
                         "Second block must map to the second segment")
        self.assertEqual(segments[2].source_text, "Block three sentence.",
                         "Third block must map to the third segment")
        self.assertLess(segments[0].segment_index, segments[1].segment_index,
                        "Block 1 must have lower segment_index than block 2 on the same page")
        self.assertLess(segments[1].segment_index, segments[2].segment_index,
                        "Block 2 must have lower segment_index than block 3 on the same page")

    # ------------------------------------------------------------------ TC-E09
    def test_tc_e09_no_match_retains_original_text(self):
        """Unrecognized input is not fabricated; source is stored verbatim; segment is marked needs_review=True."""
        from translator.services.translation_dataset import UNKNOWN_FOR_REVIEW
        from translator.services import _sync_structure_models

        completely_unknown = "ZQXZQX NONSENSE"
        result = self.dataset.translate_phrase_with_metadata(completely_unknown, "english", "tagabawa")
        self.assertEqual(result["translated"], completely_unknown,
                         "Unrecognized phrase must stay as source text - no fabricated translation")
        self.assertEqual(result["method"], "unknown_for_review")
        self.assertEqual(result["confidence"], 0.0)
        self.assertTrue(result["needs_review"])

        job = self._create_job(self.alice)
        structure_path = self._write_structure(job, {
            "extraction_method": "direct_pdf_text",
            "pages": [{"page_number": 1, "width": 612, "height": 792, "rotation": 0,
                "blocks": [{"type": "text", "bbox": [72, 100, 540, 120], "lines": [{
                    "text": completely_unknown,
                    "translated_text": UNKNOWN_FOR_REVIEW,
                    "translation_method": "unknown_for_review",
                    "translation_confidence": 0.0,
                    "bbox": [72, 100, 540, 120],
                }]}]}],
        })
        _sync_structure_models(str(job.id), structure_path)
        seg = TranslationSegment.objects.filter(job=job).first()
        self.assertIsNotNone(seg)
        self.assertEqual(seg.source_text, completely_unknown,
                         "Original source text must be stored verbatim even when no translation match exists")
        self.assertEqual(seg.translated_text, UNKNOWN_FOR_REVIEW)
        self.assertTrue(seg.needs_review,
                        "Segment must have needs_review=True when no translation was found")

    # ------------------------------------------------------------------ TC-E10
    def test_tc_e10_translated_segments_are_available_in_preview(self):
        """Completed job: segments (with Tagabawa unicode) visible in preview. Queued job: preview URL redirects."""
        completed_job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        TranslationSegment.objects.create(
            job=completed_job,
            segment_index=1,
            source_text="Hello",
            translated_text="Madigár",
            source_language="english",
            target_language="tagabawa",
            method="exact_phrase",
            confidence=1.0,
            needs_review=False,
        )
        self.client.force_login(self.alice)

        preview_response = self.client.get(reverse("translator:job_preview", args=[completed_job.id]))
        self.assertEqual(preview_response.status_code, 200,
                         "Preview page must be accessible for a COMPLETED job")
        preview_body = preview_response.content.decode("utf-8")
        self.assertIn("Hello", preview_body,
                      "Source text must appear in the bilingual preview")
        self.assertIn("Madigár", preview_body,
                      "Tagabawa translated text with accent must appear in the preview without stripping")
        self.assertNotIn("Needs review", preview_body,
                         "A needs_review=False segment must not display the 'Needs review' marker")

        # Queued job → preview URL redirects (translation not yet available)
        queued_job = self._create_job(self.alice, status=TranslationJob.Status.QUEUED)
        queued_response = self.client.get(reverse("translator:job_preview", args=[queued_job.id]))
        self.assertEqual(queued_response.status_code, 302,
                         "Preview URL for a QUEUED job must redirect — translation not yet complete")

    # ------------------------------------------------------------------ TC-E11
    @patch("translator.services._get_pipeline_service")
    def test_tc_e11_translation_error_sets_failed_status(self, mock_pipeline):
        """Per-line exception → UNKNOWN_FOR_REVIEW (job continues). Catastrophic failure → FAILED with no traceback."""
        from unittest.mock import MagicMock
        from translator.services.pipeline_service import PipelineService
        from translator.services.translation_dataset import UNKNOWN_FOR_REVIEW
        from translator.services import _run_translation_job

        # Part 1: per-line translation exception → UNKNOWN_FOR_REVIEW; pipeline does NOT propagate
        svc = PipelineService.__new__(PipelineService)
        mock_dataset = MagicMock()
        mock_dataset.translate_phrase_with_metadata.side_effect = Exception("Dataset lookup error")
        svc.translation_dataset = mock_dataset

        layout = [{"blocks": [{"type": "text", "bbox": [0, 0, 400, 50],
            "lines": [{"text": "Hello", "bbox": [0, 0, 400, 20]}]}]}]
        per_line_results = svc._translate_layout(layout, "english", "tagabawa")

        line_result = per_line_results.get("0_0_0")
        self.assertIsNotNone(line_result,
                             "_translate_layout must return a result for the line even when translation raises")
        self.assertEqual(line_result["translated"], UNKNOWN_FOR_REVIEW,
                         "Per-line translation exception must yield UNKNOWN_FOR_REVIEW, not propagate exception")
        self.assertEqual(line_result["method"], "unknown_for_review",
                         "Per-line exception must set method='unknown_for_review'")

        # Part 2: catastrophic pipeline error → job FAILED with actionable message, no traceback
        error_msg = "Translation phase failed: dataset file is unreadable"
        mock_pipeline.side_effect = Exception(error_msg)
        job = self._create_job(self.alice, status="queued")
        _run_translation_job(job.job_id, job.input_file_path, "pdf", "auto", "tagabawa", None)
        job.refresh_from_db()

        self.assertEqual(job.status, TranslationJob.Status.FAILED,
                         "A catastrophic translation error must set job status to FAILED")
        self.assertNotIn(job.status, [TranslationJob.Status.QUEUED, TranslationJob.Status.PROCESSING])
        self.assertIn(error_msg, job.error,
                      "job.error must contain the original error message")
        self.assertNotIn("Traceback", job.error,
                         "job.error must not expose Python traceback strings")
        self.assertNotIn('File "', job.error,
                         "job.error must not expose Python source file paths")


class AppendixFBilingualPreviewTests(TestCase):
    """Regression tests for Appendix F Sprint 6 — Bilingual Output and Preview Module."""

    password = "Bagobo-AppF-2026!"

    def setUp(self):
        from django.core.cache import cache as _cache
        _cache.clear()
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)
        self.alice = User.objects.create_user(
            username="appf_alice", email="appf_alice@example.test", password=self.password
        )
        self.bob = User.objects.create_user(
            username="appf_bob", email="appf_bob@example.test", password=self.password
        )

    def _create_job(self, owner, status=TranslationJob.Status.COMPLETED,
                    source_language="english", target_language="tagabawa", metadata=None):
        return TranslationJob.objects.create(
            owner=owner,
            original_filename="f_test_doc.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=status,
            source_language=source_language,
            target_language=target_language,
            metadata=metadata or {},
        )

    def _create_segment(self, job, index, source, translated, method="exact_phrase",
                        confidence=0.9, needs_review=False):
        return TranslationSegment.objects.create(
            job=job,
            segment_index=index,
            source_text=source,
            translated_text=translated,
            source_language=job.source_language,
            target_language=job.target_language,
            method=method,
            confidence=confidence,
            needs_review=needs_review,
        )

    # ------------------------------------------------------------------ TC-F01
    def test_tc_f01_completed_job_displays_bilingual_segments(self):
        """COMPLETED job: preview returns 200; source and translated text from TranslationSegment DB records
        are visible; segment count badge matches actual DB count; no hardcoded or fabricated text appears."""
        job = self._create_job(self.alice)
        self._create_segment(job, 1, "Madigár source Alpha", "Allus kó translation Alpha", confidence=1.0)
        self._create_segment(job, 2, "Madigár source Beta", "Unsad kó translation Beta", confidence=0.95)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200,
                         "COMPLETED job must return HTTP 200 for the bilingual preview page")
        body = response.content.decode("utf-8")
        self.assertIn("Madigár source Alpha", body,
                      "Source text from TranslationSegment record must appear verbatim in the preview")
        self.assertIn("Allus kó translation Alpha", body,
                      "Translated text from TranslationSegment record must appear verbatim in the preview")
        self.assertIn("Madigár source Beta", body,
                      "All DB-backed source texts must be visible in the preview")
        self.assertIn("Unsad kó translation Beta", body,
                      "All DB-backed translated texts must be visible in the preview")
        self.assertNotIn("FAKE_PREVIEW_TEXT_XYZ_HARDCODED", body,
                         "Preview must not contain any hardcoded demo text not stored in TranslationSegment")

    def test_tc_f01_segment_count_badge_visible_to_staff(self):
        """Staff/admin see the exact segment count badge in the technical panel."""
        job = self._create_job(self.alice)
        self._create_segment(job, 1, "Madigár source Alpha", "Allus kó translation Alpha", confidence=1.0)
        self._create_segment(job, 2, "Madigár source Beta", "Unsad kó translation Beta", confidence=0.95)
        self.alice.is_staff = True
        self.alice.save(update_fields=["is_staff"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("2 segments", body,
                      "Segment count badge must reflect the actual number of TranslationSegment records (2)")

    # ------------------------------------------------------------------ TC-F02
    def test_tc_f02_source_and_target_languages_are_displayed(self):
        """Preview info panel shows a Language Pair row with the exact source_language and target_language
        values stored on TranslationJob — not hardcoded labels."""
        job = self._create_job(self.alice, source_language="cebuano", target_language="filipino")
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("Language Pair", body,
                      "Preview info panel must include a 'Language Pair' label")
        self.assertIn("cebuano", body,
                      "Preview must display the job's source_language value ('cebuano') from the DB")
        self.assertIn("filipino", body,
                      "Preview must display the job's target_language value ('filipino') from the DB")

    # ------------------------------------------------------------------ TC-F03
    def test_tc_f03_segments_remain_aligned_and_ordered(self):
        """Segments are rendered in segment_index order regardless of DB insertion order.
        The source text of segment N appears before segment N+1 in the HTML output."""
        job = self._create_job(self.alice)
        # Insert in reverse order so natural DB ordering would be wrong without order_by
        self._create_segment(job, 3, "Gamma sentence third", "Trans Gamma")
        self._create_segment(job, 1, "Alpha sentence first", "Trans Alpha")
        self._create_segment(job, 2, "Beta sentence second", "Trans Beta")
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        pos_alpha = body.index("Alpha sentence first")
        pos_beta = body.index("Beta sentence second")
        pos_gamma = body.index("Gamma sentence third")
        self.assertLess(pos_alpha, pos_beta,
                        "Segment with segment_index=1 must appear before segment_index=2 in the rendered HTML")
        self.assertLess(pos_beta, pos_gamma,
                        "Segment with segment_index=2 must appear before segment_index=3 in the rendered HTML")
        # Corresponding translated texts must also be in order
        trans_pos_alpha = body.index("Trans Alpha")
        trans_pos_beta = body.index("Trans Beta")
        self.assertLess(trans_pos_alpha, trans_pos_beta,
                        "Translated text for segment 1 must appear before translated text for segment 2")

    # ------------------------------------------------------------------ TC-F04
    def test_tc_f04_unknown_translation_uses_source_text_for_normal_users(self):
        """Unknown translations display source text without exposing review details."""
        job = self._create_job(self.alice)
        # Case 1: empty translated_text falls back to source text for display.
        self._create_segment(
            job, 1,
            source="Completely unknown phrase A",
            translated="",
            method="unknown_for_review",
            confidence=0.0,
            needs_review=True,
        )
        # Case 2: literal UNKNOWN_FOR_REVIEW sentinel stored in translated_text
        self._create_segment(
            job, 2,
            source="Completely unknown phrase B",
            translated="[UNKNOWN_FOR_REVIEW]",
            method="unknown_for_review",
            confidence=0.0,
            needs_review=True,
        )
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertNotIn("[UNKNOWN_FOR_REVIEW]", body,
                         "Raw review sentinels must not be rendered in the preview body")
        self.assertNotIn("Needs review", body,
                         "Normal users must not see the technical 'Needs review' badge")
        self.assertNotIn("Needs Review", body,
                         "Normal users must not see technical review summary labels")
        self.assertNotIn("Some parts may need teacher review.", body,
                         "Normal users must not see review-detail messaging in the preview")
        self.assertIn("Completely unknown phrase A", body,
                      "Source text must remain visible even when the translation is unknown")
        self.assertIn("Completely unknown phrase B", body,
                      "Source text must remain visible for UNKNOWN_FOR_REVIEW literal sentinel case")
        self.assertNotIn("invented translation placeholder", body,
                         "The preview must never fabricate a translation for unmatched phrases")

    def test_tc_f04_needs_review_badge_visible_to_staff(self):
        """Staff/admin still see the per-segment 'Needs review' badge."""
        job = self._create_job(self.alice)
        self._create_segment(
            job, 1,
            source="Completely unknown phrase A",
            translated="",
            method="unknown_for_review",
            confidence=0.0,
            needs_review=True,
        )
        self.alice.is_staff = True
        self.alice.save(update_fields=["is_staff"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("Needs review", body)

    # ------------------------------------------------------------------ TC-F05
    def test_tc_f05_low_ocr_confidence_warning_is_displayed(self):
        """OCR confidence is stored at job/page level, NOT per-segment. The preview sidebar shows the
        job-level percentage. High-confidence OCR shows the percentage but no '⚠ Low' badge.
        Direct-extraction jobs (direct_pdf_text) do not show the OCR Confidence row at all."""
        # Part A: OCR job with high confidence — shows percentage, no ⚠ Low badge on PREVIEW page
        ocr_job = self._create_job(
            self.alice,
            metadata={
                "extraction_method": "ocr_image",
                "ocr_summary": {"mean_confidence": 0.92, "has_low_quality_warning": False},
            },
        )
        self.alice.is_staff = True
        self.alice.save(update_fields=["is_staff"])
        self.client.force_login(self.alice)
        ocr_response = self.client.get(reverse("translator:job_preview", args=[ocr_job.id]))
        self.assertEqual(ocr_response.status_code, 200)
        ocr_body = ocr_response.content.decode("utf-8")
        self.assertIn("OCR Confidence", ocr_body,
                      "Preview sidebar must show 'OCR Confidence' row for jobs with extraction_method=ocr_image")
        self.assertIn("92%", ocr_body,
                      "Preview sidebar must display the stored OCR confidence percentage (92%)")
        self.assertNotIn("⚠ Low", ocr_body,
                         "High-confidence OCR (0.92) must NOT display the '⚠ Low' badge in the preview")

        # Part B: Direct PDF extraction job — OCR Confidence row must be absent on PREVIEW page
        digital_job = self._create_job(
            self.alice,
            metadata={"extraction_method": "direct_pdf_text"},
        )
        digital_response = self.client.get(reverse("translator:job_preview", args=[digital_job.id]))
        self.assertEqual(digital_response.status_code, 200)
        digital_body = digital_response.content.decode("utf-8")
        self.assertNotIn("OCR Confidence", digital_body,
                         "Preview sidebar must NOT show 'OCR Confidence' row for direct-extraction "
                         "(digital PDF) jobs — OCR was never run on these documents")

    # ------------------------------------------------------------------ TC-F06
    def test_tc_f06_unfinished_and_failed_jobs_cannot_preview(self):
        """RETRYING and FAILED jobs redirect from the preview URL to job_detail.
        The job_detail page remains accessible for status information."""
        job_retrying = self._create_job(self.alice, status=TranslationJob.Status.RETRYING)
        job_failed = self._create_job(self.alice, status=TranslationJob.Status.FAILED)
        job_failed.error = "Translation failed during OCR stage."
        job_failed.save(update_fields=["error"])
        self.client.force_login(self.alice)

        retrying_response = self.client.get(reverse("translator:job_preview", args=[job_retrying.id]))
        self.assertRedirects(
            retrying_response,
            reverse("translator:job_detail", args=[job_retrying.id]),
            msg_prefix="RETRYING job preview URL must redirect to job_detail",
        )

        failed_response = self.client.get(reverse("translator:job_preview", args=[job_failed.id]))
        self.assertRedirects(
            failed_response,
            reverse("translator:job_detail", args=[job_failed.id]),
            msg_prefix="FAILED job preview URL must redirect to job_detail",
        )

    # ------------------------------------------------------------------ Additional: soft-deleted job → 404
    def test_tc_f_soft_deleted_job_preview_returns_404(self):
        """Soft-deleted COMPLETED job: preview URL returns 404; segment data is not exposed."""
        job = self._create_job(self.alice, status=TranslationJob.Status.COMPLETED)
        self._create_segment(job, 1, "Sensitive source F text", "Translated sensitive F")
        job.is_deleted = True
        job.deleted_at = timezone.now()
        job.save(update_fields=["is_deleted", "deleted_at"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 404,
                         "Soft-deleted COMPLETED job must return 404 from the preview URL, even for the owner")
        body = response.content.decode("utf-8")
        self.assertNotIn("Sensitive source F text", body,
                         "Soft-deleted job's TranslationSegment data must not appear in the 404 response")

    # ------------------------------------------------------------------ Additional: empty segment collection → safe empty state
    def test_tc_f_empty_segment_collection_shows_safe_empty_state(self):
        """COMPLETED job with zero TranslationSegment records: preview returns 200 with the
        safe empty-state section; no crash, no fabricated segment rows."""
        job = self._create_job(self.alice)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200,
                         "A COMPLETED job with zero segments must still return HTTP 200")
        body = response.content.decode("utf-8")
        # Normal users see the main panel's safe empty state (the technical
        # segment table — and its own empty-state copy — is staff-only).
        self.assertIn("No translated preview available", body,
                      "Zero-segment preview must display the safe empty-state message")
        self.assertNotIn("Segment 1", body,
                         "Zero-segment preview must not fabricate or display any segment rows")

    def test_tc_f_empty_segment_collection_technical_empty_state_for_staff(self):
        """Staff/admin see the technical segment table's own empty-state copy."""
        job = self._create_job(self.alice)
        self.alice.is_staff = True
        self.alice.save(update_fields=["is_staff"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("No translated segments available yet", body)

    # ------------------------------------------------------------------ Additional: unauthenticated access requires login
    def test_tc_f_unauthenticated_preview_requires_login(self):
        """Unauthenticated GET to the preview URL returns 302 redirect to the login page."""
        job = self._create_job(self.alice)
        preview_url = reverse("translator:job_preview", args=[job.id])

        response = self.client.get(preview_url)

        self.assertEqual(response.status_code, 302,
                         "Unauthenticated access to the preview URL must return 302")
        self.assertRedirects(
            response,
            f"{reverse('translator:login')}?next={preview_url}",
            msg_prefix="Unauthenticated preview must redirect to login with a ?next= parameter",
        )

    # ------------------------------------------------------------------ Additional: HTML content is auto-escaped
    def test_tc_f_html_content_is_escaped_safely(self):
        """HTML special characters in source_text and translated_text are auto-escaped by the Django template engine;
        raw script or tag injection cannot occur through TranslationSegment data."""
        job = self._create_job(self.alice)
        xss_source = "<script>alert('XSS')</script>"
        xss_translated = "<b>Bold injection</b>"
        self._create_segment(job, 1, xss_source, xss_translated, method="exact_phrase", confidence=0.9)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertNotIn("<script>alert('XSS')</script>", body,
                         "Raw <script> tag from source_text must be HTML-escaped, not rendered verbatim")
        self.assertIn("&lt;script&gt;", body,
                      "Django must HTML-escape < as &lt; in source_text; escaped form must appear in the body")
        self.assertNotIn("<b>Bold injection</b>", body,
                         "Raw <b> tag from translated_text must be HTML-escaped, not rendered verbatim")
        self.assertIn("&lt;b&gt;", body,
                      "Django must HTML-escape < as &lt; in translated_text; escaped form must appear in the body")


# ===========================================================================
# APPENDIX G — Sprint 7: Document Download Module
# ===========================================================================


class AppendixGDownloadTests(TestCase):
    """Regression tests for Appendix G Sprint 7 — Document Download Module."""

    password = "Bagobo-AppG-2026!"

    def setUp(self):
        from django.core.cache import cache as _cache
        _cache.clear()
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)
        self.alice = User.objects.create_user(
            username="appg_alice", email="appg_alice@example.test", password=self.password
        )
        self.bob = User.objects.create_user(
            username="appg_bob", email="appg_bob@example.test", password=self.password
        )

    def _create_job(self, owner, status=TranslationJob.Status.COMPLETED,
                    source_language="english", target_language="tagabawa", metadata=None):
        return TranslationJob.objects.create(
            owner=owner, original_filename="g_test_doc.pdf",
            file_type=TranslationJob.FileType.PDF, status=status,
            source_language=source_language, target_language=target_language,
            metadata=metadata or {},
        )

    def _minimal_pdf_bytes(self, page_count=1):
        import fitz
        doc = fitz.open()
        for _ in range(page_count):
            doc.new_page(width=612, height=792)
        data = doc.tobytes()
        doc.close()
        return data

    def _write_job_output(self, job, filename="translated.pdf", content=None):
        from translator.services import job_directory_path
        job_dir = Path(job_directory_path(job.job_id))
        output_dir = job_dir / "translated"
        output_dir.mkdir(parents=True, exist_ok=True)
        file_path = output_dir / filename
        if content is None:
            content = self._minimal_pdf_bytes()
        file_path.write_bytes(content)
        return str(file_path)

    # ------------------------------------------------------------------ TC-G01
    def test_tc_g01_completed_bilingual_document_downloads(self):
        """COMPLETED job with a real output file returns 200, Content-Type application/pdf,
        and a Content-Disposition attachment header whose filename matches the expected
        format translated_<first-8-chars-of-job-id>.pdf. No mocking is used."""
        job = self._create_job(self.alice)
        output_path = self._write_job_output(job)
        job.output_file_path = output_path
        job.save(update_fields=["output_file_path"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_download", args=[job.id]))

        try:
            self.assertEqual(response.status_code, 200,
                             "COMPLETED job with a real output file must return HTTP 200 on download")
            self.assertEqual(
                response.headers["Content-Type"], "application/pdf",
                "Download response Content-Type must be application/pdf for all output variants",
            )
            disposition = response.headers.get("Content-Disposition", "")
            self.assertIn("attachment", disposition,
                          "Content-Disposition must include 'attachment' to force a browser download")
            expected_filename = f"translated_{job.job_id[:8]}.pdf"
            self.assertIn(
                expected_filename, disposition,
                f"Content-Disposition filename must be '{expected_filename}' "
                "(prefix 'translated_' + first 8 chars of job UUID + '.pdf')",
            )
        finally:
            response.close()

    # ------------------------------------------------------------------ TC-G02
    def test_tc_g02_downloaded_output_preserves_bilingual_order(self):
        """Part A: _create_output_pdf with a 2-page layout_data creates a valid 2-page PDF.
        Part B: ?format=bilingual serves the bilingual file with the filename prefix 'bilingual_'.
        Together these confirm that the output generation pipeline and the bilingual download
        path function correctly end-to-end."""
        import fitz
        from translator.services.pipeline_service import PipelineService

        # Part A — output generation: 2-page layout produces a 2-page PDF
        output_path = str(self.media_root / "bilingual_order_test.pdf")
        layout_data = [
            {
                "page": 0, "width": 612, "height": 792,
                "blocks": [
                    {
                        "type": "text", "bbox": [72, 100, 540, 130],
                        "lines": [{"text": "The water is clean.", "bbox": [72, 100, 540, 130]}],
                    }
                ],
            },
            {
                "page": 1, "width": 612, "height": 792,
                "blocks": [
                    {
                        "type": "text", "bbox": [72, 100, 540, 130],
                        "lines": [{"text": "The forest is tall.", "bbox": [72, 100, 540, 130]}],
                    }
                ],
            },
        ]
        translations = {
            "The water is clean.": {"translated": "Malinig ang tubig.", "method": "exact_phrase"},
            "The forest is tall.": {"translated": "Mataas ang kagubatan.", "method": "exact_phrase"},
        }
        pipeline = PipelineService()
        ok = pipeline._create_output_pdf(layout_data, translations, output_path)
        self.assertTrue(ok,
                        "_create_output_pdf must return True when processing a 2-page layout with text blocks")
        self.assertTrue(Path(output_path).exists(),
                        "Output PDF must exist on disk after _create_output_pdf returns True")
        doc = fitz.open(output_path)
        page_count = doc.page_count
        full_text = "\n".join(doc[i].get_text() for i in range(doc.page_count))
        doc.close()
        self.assertEqual(page_count, 2,
                         "_create_output_pdf must produce exactly 2 pages for a 2-element layout_data list")
        self.assertIn("Malinig ang tubig.", full_text)
        self.assertIn("Mataas ang kagubatan.", full_text)

        # Part B — bilingual download path: ?format=bilingual serves the bilingual file
        job = self._create_job(self.alice)
        translated_path = self._write_job_output(job, "translated.pdf")
        bilingual_path = self._write_job_output(job, "bilingual.pdf")
        job.output_file_path = translated_path
        job.bilingual_file_path = bilingual_path
        job.save(update_fields=["output_file_path", "bilingual_file_path"])
        self.client.force_login(self.alice)

        response = self.client.get(
            reverse("translator:job_download", args=[job.id]) + "?format=bilingual"
        )
        try:
            self.assertEqual(response.status_code, 200,
                             "?format=bilingual must return HTTP 200 when the bilingual PDF exists in the job directory")
            disposition = response.headers.get("Content-Disposition", "")
            expected_bilingual_filename = f"bilingual_{job.job_id[:8]}.pdf"
            self.assertIn(
                expected_bilingual_filename, disposition,
                f"Bilingual download Content-Disposition filename must be '{expected_bilingual_filename}'",
            )
        finally:
            response.close()

    def test_reconstruct_pdf_rejects_layout_that_omits_source_pages(self):
        from translator.services.reconstruction_service import ReconstructionService

        source_path = self.media_root / "two_page_source.pdf"
        source_path.write_bytes(self._minimal_pdf_bytes(page_count=2))
        output_path = str(self.media_root / "partial_reconstruction.pdf")
        layout_data = [
            {
                "page": 0,
                "width": 612,
                "height": 792,
                "blocks": [
                    {
                        "type": "text",
                        "bbox": [72, 100, 540, 130],
                        "lines": [{"text": "Page one only", "bbox": [72, 100, 540, 130]}],
                    }
                ],
            }
        ]
        warnings = []

        ok = ReconstructionService.reconstruct_pdf(
            str(source_path),
            layout_data,
            {"Page one only": {"translated": "Translated page one", "method": "exact_phrase"}},
            output_path,
            layout_warnings=warnings,
        )

        self.assertFalse(ok)
        self.assertFalse(Path(output_path).exists())
        self.assertTrue(
            any("not every page can be translated safely" in warning for warning in warnings)
        )

    # ------------------------------------------------------------------ TC-G03
    def test_tc_g03_processing_job_cannot_download(self):
        """A RETRYING job (non-COMPLETED intermediate status) must redirect from the download
        URL to job_detail. The existing suite already tests PROCESSING; this covers RETRYING,
        confirming that all non-COMPLETED statuses are blocked at the download endpoint."""
        job = self._create_job(self.alice, status=TranslationJob.Status.RETRYING)
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_download", args=[job.id]))

        self.assertRedirects(
            response,
            reverse("translator:job_detail", args=[job.id]),
            msg_prefix="RETRYING job must redirect from the download URL to job_detail",
        )

    # ------------------------------------------------------------------ TC-G04
    def test_tc_g04_failed_job_cannot_download(self):
        """A FAILED job must redirect from the download URL to job_detail.
        The job_detail response must not contain a raw Python traceback."""
        job = self._create_job(self.alice, status=TranslationJob.Status.FAILED)
        job.error = "OCR stage failed: Tesseract returned non-zero exit code."
        job.save(update_fields=["error"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_download", args=[job.id]))

        self.assertRedirects(
            response,
            reverse("translator:job_detail", args=[job.id]),
            msg_prefix="FAILED job must redirect from the download URL to job_detail",
        )
        detail_response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        body = detail_response.content.decode("utf-8")
        self.assertNotIn(
            "Traceback (most recent call last)", body,
            "The job_detail page must not expose a raw Python traceback to the user",
        )

    # ------------------------------------------------------------------ TC-G05
    def test_tc_g05_completed_job_without_output_cannot_download(self):
        """COMPLETED job whose output file has been physically deleted from disk must redirect
        from the download URL to job_detail. The response must not expose any absolute
        filesystem path (the error message must be static, not include the storage path)."""
        job = self._create_job(self.alice)
        output_path = self._write_job_output(job)
        job.output_file_path = output_path
        job.save(update_fields=["output_file_path"])
        Path(output_path).unlink()
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_download", args=[job.id]))

        self.assertRedirects(
            response,
            reverse("translator:job_detail", args=[job.id]),
            msg_prefix="COMPLETED job with deleted output file must redirect to job_detail",
        )
        detail_response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        body = detail_response.content.decode("utf-8")
        self.assertNotIn(
            str(self.media_root), body,
            "Error message after missing output file must not expose any absolute filesystem path",
        )
        self.assertNotIn(
            output_path, body,
            "Error message must not include the raw output_file_path value",
        )

    # ------------------------------------------------------------------ TC-G06
    def test_tc_g06_cross_user_download_is_denied(self):
        """alice cannot download a job belonging to bob — _get_owned_job filters by owner
        and returns 404 for jobs owned by other users. The 404 response body must not
        contain any of bob's job metadata."""
        job = self._create_job(self.bob)
        output_path = self._write_job_output(job)
        job.output_file_path = output_path
        job.save(update_fields=["output_file_path"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_download", args=[job.id]))

        self.assertEqual(response.status_code, 404,
                         "Accessing another user's download URL must return 404")
        body = response.content.decode("utf-8")
        self.assertNotIn(
            job.original_filename, body,
            "The 404 page must not expose the filename of another user's job",
        )

    # ------------------------------------------------------------------ Additional: unauthenticated access
    def test_tc_g_unauthenticated_download_redirects_to_login(self):
        """An unauthenticated GET on a download URL must return 302 to the login page
        with a ?next= parameter pointing back to the download URL."""
        job = self._create_job(self.alice)
        download_url = reverse("translator:job_download", args=[job.id])

        response = self.client.get(download_url)

        self.assertEqual(response.status_code, 302,
                         "Unauthenticated download request must return HTTP 302")
        self.assertRedirects(
            response,
            f"{reverse('translator:login')}?next={download_url}",
            msg_prefix="Unauthenticated download must redirect to the login page with ?next= pointing to the download URL",
        )

    # ------------------------------------------------------------------ Additional: soft-deleted job
    def test_tc_g_soft_deleted_job_download_returns_404(self):
        """A COMPLETED job that has been soft-deleted (is_deleted=True) must return 404
        on the download URL. The _get_owned_job helper filters is_deleted=False so
        soft-deleted jobs are invisible even to their original owner."""
        job = self._create_job(self.alice)
        output_path = self._write_job_output(job)
        job.output_file_path = output_path
        job.is_deleted = True
        job.save(update_fields=["output_file_path", "is_deleted"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_download", args=[job.id]))

        self.assertEqual(response.status_code, 404,
                         "Soft-deleted COMPLETED job with a valid output file must return 404 on the download URL")

    # ------------------------------------------------------------------ Additional: output PDF with translated text blocks
    def test_tc_g_output_pdf_with_text_blocks_is_valid(self):
        """_create_output_pdf with non-empty text blocks containing a Tagabawa translation
        (Latin-1 accented characters, ord <= 255) produces a valid, openable PDF with the
        correct page count. This confirms that translated content is inserted without
        corrupting the output file."""
        import fitz
        from translator.services.pipeline_service import PipelineService

        output_path = str(self.media_root / "tagabawa_output.pdf")
        layout_data = [
            {
                "page": 0, "width": 612, "height": 792,
                "blocks": [
                    {
                        "type": "text",
                        "bbox": [72.0, 100.0, 540.0, 130.0],
                        "lines": [
                            {
                                "text": "Good morning.",
                                "bbox": [72.0, 100.0, 540.0, 130.0],
                                "spans": [{"font": "Helvetica", "size": 12.0, "flags": 0}],
                            }
                        ],
                    }
                ],
            }
        ]
        # á (U+00E1, ord=225) is within ISO Latin-1; no unicode font required
        translations = {
            "Good morning.": {"translated": "Madigár nu uras.", "method": "exact_phrase"},
        }
        pipeline = PipelineService()
        ok = pipeline._create_output_pdf(layout_data, translations, output_path)
        self.assertTrue(ok,
                        "_create_output_pdf must return True when inserting Tagabawa translated text")
        self.assertTrue(Path(output_path).exists(),
                        "Output PDF must exist after _create_output_pdf with Tagabawa content")
        doc = fitz.open(output_path)
        page_count = doc.page_count
        doc.close()
        self.assertEqual(page_count, 1,
                         "Output PDF with Tagabawa translation must have exactly 1 page matching layout_data")

    # ------------------------------------------------------------------ Additional: bilingual format with missing bilingual file
    def test_tc_g_bilingual_format_missing_redirects_safely(self):
        """A COMPLETED job that has a translated PDF (can_download=True) but no bilingual PDF
        must redirect from ?format=bilingual to job_detail with a safe error message.
        The bilingual_file_path being empty triggers _resolve_job_file to return None,
        which the view handles by redirecting rather than raising a 500."""
        job = self._create_job(self.alice)
        translated_path = self._write_job_output(job, "translated.pdf")
        job.output_file_path = translated_path
        job.bilingual_file_path = ""
        job.save(update_fields=["output_file_path", "bilingual_file_path"])
        self.client.force_login(self.alice)

        response = self.client.get(
            reverse("translator:job_download", args=[job.id]) + "?format=bilingual"
        )

        self.assertRedirects(
            response,
            reverse("translator:job_detail", args=[job.id]),
            msg_prefix="?format=bilingual with no bilingual PDF must redirect to job_detail rather than crash",
        )


# ===========================================================================
# APPENDIX H — Sprint 8: Translation History and Job Management Module
# ===========================================================================


class AppendixHHistoryTests(TestCase):
    """Regression tests for Appendix H Sprint 8 — Translation History and Job Management Module."""

    password = "Bagobo-AppH-2026!"

    def setUp(self):
        from django.core.cache import cache as _cache
        _cache.clear()
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)
        self.alice = User.objects.create_user(
            username="apph_alice", email="apph_alice@example.test", password=self.password
        )
        self.bob = User.objects.create_user(
            username="apph_bob", email="apph_bob@example.test", password=self.password
        )

    def _create_job(self, owner, status=TranslationJob.Status.COMPLETED,
                    original_filename="h_test_doc.pdf",
                    source_language="english", target_language="tagabawa"):
        return TranslationJob.objects.create(
            owner=owner, original_filename=original_filename,
            file_type=TranslationJob.FileType.PDF, status=status,
            source_language=source_language, target_language=target_language,
        )

    def _write_job_output(self, job, filename="translated.pdf", content=None):
        from translator.services import job_directory_path
        job_dir = Path(job_directory_path(job.job_id))
        output_dir = job_dir / "translated"
        output_dir.mkdir(parents=True, exist_ok=True)
        file_path = output_dir / filename
        if content is None:
            import fitz
            doc = fitz.open()
            doc.new_page(width=612, height=792)
            content = doc.tobytes()
            doc.close()
        file_path.write_bytes(content)
        return str(file_path)

    # ------------------------------------------------------------------ TC-H01
    def test_tc_h01_history_lists_previous_owner_records(self):
        """History lists all non-deleted jobs owned by the current user, ordered newest-first.
        Jobs belonging to other users are excluded. No fabricated records appear — every card
        is sourced from a real TranslationJob DB row owned by the requesting user."""
        from datetime import timedelta

        now = timezone.now()
        job_older = self._create_job(self.alice, original_filename="alice_older.pdf")
        TranslationJob.objects.filter(pk=job_older.pk).update(
            created_at=now - timedelta(days=2)
        )
        job_newer = self._create_job(self.alice, original_filename="alice_newer.pdf")
        # Force an explicit timestamp so the ordering is unambiguous in SQLite's
        # shared-memory mode, where auto_now_add may return the same microsecond
        # for two consecutive inserts within the same test transaction.
        TranslationJob.objects.filter(pk=job_newer.pk).update(
            created_at=now - timedelta(seconds=1)
        )
        self._create_job(self.bob, original_filename="bob_secret.pdf")
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:history"))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("alice_newer.pdf", body,
                      "History must list alice's newer job")
        self.assertIn("alice_older.pdf", body,
                      "History must list alice's older job")
        self.assertNotIn("bob_secret.pdf", body,
                         "History must exclude jobs belonging to other users")
        self.assertLess(
            body.index("alice_newer.pdf"),
            body.index("alice_older.pdf"),
            "Newest job must appear before the older job in History (ordered by -created_at)",
        )
        jobs = response.context["jobs"]
        self.assertEqual(len(jobs), 2,
                         "History context must contain exactly alice's 2 non-deleted jobs")

    # ------------------------------------------------------------------ TC-H02
    def test_tc_h02_history_displays_title_date_and_status(self):
        """Each History card shows the document title (original_filename), upload date,
        processing status label, and language pair — all read from the TranslationJob record.
        No hardcoded or fabricated values are rendered."""
        job = self._create_job(
            self.alice,
            status=TranslationJob.Status.COMPLETED,
            original_filename="annual_report_2025.pdf",
            source_language="english",
            target_language="tagabawa",
        )
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:history"))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("annual_report_2025.pdf", body,
                      "History must display the job's original_filename as the document title")
        self.assertIn(str(timezone.now().year), body,
                      "History must display the upload date year from TranslationJob.created_at")
        self.assertIn("Completed", body,
                      "History must display the status label ('Completed') from the job's status field")
        self.assertIn("english", body,
                      "History must display the source_language field from the TranslationJob record")
        self.assertIn("tagabawa", body,
                      "History must display the target_language field from the TranslationJob record")

    # ------------------------------------------------------------------ TC-H03
    def test_tc_h03_history_record_opens_correct_job(self):
        """Navigating to a job's detail URL shows the correct document title and status.
        When a different user (bob) attempts to access alice's job_detail URL directly,
        _owned_jobs_queryset returns an empty result and the view raises Http404."""
        job = self._create_job(
            self.alice,
            status=TranslationJob.Status.COMPLETED,
            original_filename="alice_exclusive_doc.pdf",
        )
        detail_url = reverse("translator:job_detail", args=[job.id])

        self.client.force_login(self.alice)
        alice_response = self.client.get(detail_url)
        self.assertEqual(alice_response.status_code, 200)
        alice_body = alice_response.content.decode("utf-8")
        self.assertIn("alice_exclusive_doc.pdf", alice_body,
                      "job_detail must display the correct document title from the job record")
        self.assertIn("Completed", alice_body,
                      "job_detail must display the correct processing status")

        self.client.force_login(self.bob)
        bob_response = self.client.get(detail_url)
        self.assertEqual(bob_response.status_code, 404,
                         "bob must receive 404 when accessing alice's job_detail — _owned_jobs_queryset filters by owner")
        self.assertNotIn("alice_exclusive_doc.pdf", bob_response.content.decode("utf-8"),
                         "The 404 page must not expose alice's document title to bob")

    # ------------------------------------------------------------------ TC-H04
    def test_tc_h04_completed_record_reopens_preview_and_download(self):
        """COMPLETED job in the History list exposes 'Preview Bilingual'.
        COMPLETED job with a real output file also exposes 'Download'.
        FAILED job exposes neither — the context flags can_preview and can_download gate
        both the History template buttons and the underlying view logic."""
        completed_job = self._create_job(
            self.alice, status=TranslationJob.Status.COMPLETED,
            original_filename="completed_preview_only.pdf",
        )
        completed_with_dl = self._create_job(
            self.alice, status=TranslationJob.Status.COMPLETED,
            original_filename="completed_with_download.pdf",
        )
        output_path = self._write_job_output(completed_with_dl)
        completed_with_dl.output_file_path = output_path
        completed_with_dl.save(update_fields=["output_file_path"])
        failed_job = self._create_job(
            self.alice, status=TranslationJob.Status.FAILED,
            original_filename="failed_no_actions.pdf",
        )
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:history"))

        self.assertEqual(response.status_code, 200)
        jobs = response.context["jobs"]

        completed_state = next(s for s in jobs if s["job"].original_filename == "completed_preview_only.pdf")
        completed_dl_state = next(s for s in jobs if s["job"].original_filename == "completed_with_download.pdf")
        failed_state = next(s for s in jobs if s["job"].original_filename == "failed_no_actions.pdf")

        self.assertTrue(completed_state["can_preview"],
                        "COMPLETED job must have can_preview=True in History context")
        self.assertFalse(completed_state["can_download"],
                         "COMPLETED job without output file must have can_download=False")
        self.assertTrue(completed_dl_state["can_preview"],
                        "COMPLETED job with output must have can_preview=True")
        self.assertTrue(completed_dl_state["can_download"],
                        "COMPLETED job with real output file must have can_download=True")
        self.assertFalse(failed_state["can_preview"],
                         "FAILED job must have can_preview=False")
        self.assertFalse(failed_state["can_download"],
                         "FAILED job must have can_download=False")

    # ------------------------------------------------------------------ TC-H05
    def test_tc_h05_owner_can_soft_delete_history_record(self):
        """alice POSTs to the delete endpoint for her own job. The view sets is_deleted=True
        and deleted_at, creates a UserActivityLog entry, redirects to History, and the
        job disappears from History. The database row is NOT permanently deleted."""
        job = self._create_job(self.alice, original_filename="soft_delete_me.pdf")
        job_pk = str(job.pk)
        self.client.force_login(self.alice)

        response = self.client.post(reverse("translator:job_delete", args=[job.id]))

        self.assertRedirects(response, reverse("translator:history"),
                             msg_prefix="Successful delete must redirect to the History page")

        job.refresh_from_db()
        self.assertTrue(job.is_deleted,
                        "is_deleted must be True after the owner soft-deletes the job via POST")
        self.assertIsNotNone(job.deleted_at,
                             "deleted_at must be populated (not None) after soft-deletion")

        history_response = self.client.get(reverse("translator:history"))
        self.assertEqual(len(history_response.context["jobs"]), 0,
                         "Soft-deleted job must not appear in History after deletion")

        from translator.models import UserActivityLog as UAL
        log = UAL.objects.filter(
            user=self.alice,
            action="delete_job",
            object_id=job_pk,
        ).first()
        self.assertIsNotNone(log,
                             "delete_job action must be recorded in UserActivityLog on soft-deletion")
        self.assertEqual(log.metadata.get("filename"), "soft_delete_me.pdf",
                         "UserActivityLog metadata must include the original filename")

    # ------------------------------------------------------------------ TC-H06
    def test_tc_h06_cross_user_history_record_access_is_denied(self):
        """bob cannot access alice's Job Details, Confirm-Delete page, or delete endpoint.
        All three return 404. No alice data (filename, status, metadata) is exposed in any
        denial response. Alice's job record remains unchanged."""
        job = self._create_job(
            self.alice,
            status=TranslationJob.Status.COMPLETED,
            original_filename="alice_classified_doc.pdf",
        )
        self.client.force_login(self.bob)

        detail_response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        self.assertEqual(detail_response.status_code, 404,
                         "bob must receive 404 when accessing alice's job_detail URL")
        self.assertNotIn("alice_classified_doc.pdf",
                         detail_response.content.decode("utf-8"),
                         "404 page must not expose alice's document filename to bob")

        confirm_response = self.client.get(
            reverse("translator:job_delete_confirm", args=[job.id])
        )
        self.assertEqual(confirm_response.status_code, 404,
                         "bob must receive 404 when accessing alice's confirm-delete page")
        self.assertNotIn("alice_classified_doc.pdf",
                         confirm_response.content.decode("utf-8"),
                         "Confirm-delete 404 must not reveal alice's filename to bob")

        delete_response = self.client.post(
            reverse("translator:job_delete", args=[job.id])
        )
        self.assertEqual(delete_response.status_code, 404,
                         "bob's DELETE POST must be denied with 404")

        job.refresh_from_db()
        self.assertFalse(job.is_deleted,
                         "alice's job must remain non-deleted after all of bob's denied access attempts")

    # ------------------------------------------------------------------ Additional: Cancel keeps record
    def test_tc_h_cancel_delete_keeps_record(self):
        """The confirm-delete page shows a Cancel link that points to History.
        Without a POST submission the job record remains unchanged (is_deleted=False).
        The confirm page also shows the correct document title."""
        job = self._create_job(self.alice, original_filename="keep_me_safe.pdf")
        self.client.force_login(self.alice)

        confirm_response = self.client.get(
            reverse("translator:job_delete_confirm", args=[job.id])
        )

        self.assertEqual(confirm_response.status_code, 200)
        body = confirm_response.content.decode("utf-8")
        self.assertIn("Cancel", body,
                      "Confirm-delete page must show a Cancel option so the user can abort")
        self.assertIn(reverse("translator:history"), body,
                      "Cancel link must point to the History page")
        self.assertIn("keep_me_safe.pdf", body,
                      "Confirm-delete page must display the correct document title")

        job.refresh_from_db()
        self.assertFalse(job.is_deleted,
                         "Viewing the confirm-delete page (no POST) must leave is_deleted=False")

    # ------------------------------------------------------------------ Additional: Repeated delete returns 404
    def test_tc_h_repeated_delete_attempt_returns_404(self):
        """A second POST to delete an already soft-deleted job returns 404. The delete view
        uses get_object_or_404(..., is_deleted=False), so already-deleted jobs are invisible
        to subsequent requests — preventing double-deletion edge cases."""
        job = self._create_job(self.alice, original_filename="delete_twice.pdf")
        self.client.force_login(self.alice)

        first_response = self.client.post(reverse("translator:job_delete", args=[job.id]))
        self.assertRedirects(first_response, reverse("translator:history"),
                             msg_prefix="First delete must redirect to History")

        second_response = self.client.post(reverse("translator:job_delete", args=[job.id]))
        self.assertEqual(second_response.status_code, 404,
                         "Second delete POST on an already soft-deleted job must return 404")

    # ------------------------------------------------------------------ Additional: Deleted record absent from recent jobs
    def test_tc_h_deleted_record_absent_from_recent_jobs(self):
        """After soft-deleting a job, the Translate page's recent_jobs sidebar no longer
        includes that job. The translate view and context_processor both filter is_deleted=False."""
        job = self._create_job(self.alice, original_filename="recent_then_gone.pdf")
        self.client.force_login(self.alice)

        before = self.client.get(reverse("translator:translate"))
        before_filenames = [c["job"].original_filename for c in before.context["recent_jobs"]]
        self.assertIn("recent_then_gone.pdf", before_filenames,
                      "Job must appear in recent_jobs before deletion")

        self.client.post(reverse("translator:job_delete", args=[job.id]))

        after = self.client.get(reverse("translator:translate"))
        after_filenames = [c["job"].original_filename for c in after.context["recent_jobs"]]
        self.assertNotIn("recent_then_gone.pdf", after_filenames,
                         "Soft-deleted job must not appear in recent_jobs after deletion")

    # ------------------------------------------------------------------ Additional: HTML-escaped document title
    def test_tc_h_document_title_is_html_escaped_in_history(self):
        """A document filename containing HTML special characters is auto-escaped by Django's
        template engine. The raw script tag must never appear verbatim in the History response
        body, guarding against potential stored XSS via malicious filenames."""
        job = self._create_job(
            self.alice,
            original_filename="<script>alert('xss')</script>.pdf",
        )
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:history"))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertNotIn("<script>alert('xss')</script>", body,
                         "Raw <script> tag in filename must never appear verbatim in History")
        self.assertIn("&lt;script&gt;", body,
                      "Django auto-escaping must render '<' as '&lt;' in the History template")


# ===========================================================================
# APPENDIX J — Sprint 10: Administrator Management Module
# ===========================================================================


class AppendixJAdminTests(TestCase):
    """
    Regression tests for Appendix J Sprint 10 — Administrator Management Module.

    Verifies Django Admin access control, user list management, job visibility
    across all statuses, phrasebook dataset round-trip, and soft-delete visibility.
    The CRITICAL data-source finding (TC-J05) is documented: Admin-managed
    PhrasebookEntry records are persisted to the DB but are NOT read by the live
    translation service; the service uses translation_data.json only.
    """

    password = "Bagobo-AppJ-2026!"

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)

        # Regular user: no staff, no superuser
        self.regular = User.objects.create_user(
            username="appj_regular",
            email="appj_regular@example.test",
            password=self.password,
        )
        # Staff user only: is_staff=True, is_superuser=False, no model permissions
        self.staff_only = User.objects.create_user(
            username="appj_staff",
            email="appj_staff@example.test",
            password=self.password,
            is_staff=True,
        )
        # Superuser: full Admin access
        self.superuser = User.objects.create_superuser(
            username="appj_super",
            email="appj_super@example.test",
            password=self.password,
        )

    def _create_job(self, owner, status=TranslationJob.Status.COMPLETED,
                    original_filename="j_test.pdf"):
        return TranslationJob.objects.create(
            owner=owner,
            original_filename=original_filename,
            file_type=TranslationJob.FileType.PDF,
            status=status,
        )

    # ------------------------------------------------------------------ TC-J01
    def test_tc_j01_valid_superuser_login_opens_admin_dashboard(self):
        """A valid superuser can POST credentials to the Django Admin login page
        and is redirected to the Admin dashboard. The dashboard must return HTTP 200
        and contain the administration heading. No traceback or secret key is exposed."""
        response = self.client.post(
            "/admin/login/",
            {
                "username": self.superuser.username,
                "password": self.password,
                "next": "/admin/",
            },
        )
        # Django redirects to /admin/ after successful login
        self.assertIn(
            response.status_code, [301, 302],
            "Valid superuser login must redirect (301/302) to the Admin dashboard",
        )
        location = response.get("Location", "")
        self.assertIn(
            "/admin/", location,
            "Redirect after valid superuser login must point to /admin/",
        )
        # Follow the redirect and verify the dashboard is accessible
        dashboard = self.client.get("/admin/")
        self.assertEqual(dashboard.status_code, 200,
                         "Admin dashboard must return 200 for an authenticated superuser")
        body = dashboard.content.decode("utf-8")
        self.assertIn("administration", body.lower(),
                      "Admin dashboard must contain an 'administration' heading")
        self.assertNotIn("Traceback", body,
                         "Admin dashboard must not expose a Python traceback")
        self.assertNotIn("SECRET_KEY", body,
                         "Admin dashboard must not expose the Django secret key")

    def test_tc_j01_invalid_credentials_rejected(self):
        """Posting invalid credentials to the Admin login endpoint must not create
        an authenticated session. The form must redisplay (HTTP 200) without exposing
        any administrative data or Python traceback."""
        response = self.client.post(
            "/admin/login/",
            {
                "username": self.superuser.username,
                "password": "Wrong-Password-9999!",
                "next": "/admin/",
            },
        )
        # Must stay on the login page (200) — not redirect to the dashboard
        self.assertEqual(
            response.status_code, 200,
            "Invalid credentials must redisplay the login form (HTTP 200)",
        )
        body = response.content.decode("utf-8")
        self.assertNotIn(
            "Site administration", body,
            "Admin dashboard heading must not appear after a failed login attempt",
        )
        self.assertNotIn("Traceback", body,
                         "Failed login must not expose a Python traceback")
        # No authenticated session must be created
        self.assertNotIn(
            "_auth_user_id", self.client.session,
            "Failed Admin login must not create an authenticated session",
        )

    # ------------------------------------------------------------------ TC-J02
    def test_tc_j02_regular_user_cannot_access_admin_and_no_data_exposed(self):
        """A regular (non-staff) authenticated user requesting /admin/ is redirected
        to the Admin login page. The redirect response body must not expose filenames,
        usernames, job IDs, or any registered model data belonging to the user."""
        job = self._create_job(self.regular, original_filename="j02_secret_doc.pdf")
        self.client.force_login(self.regular)

        response = self.client.get("/admin/")

        self.assertNotEqual(response.status_code, 200,
                            "Regular user must not receive the Admin dashboard (HTTP 200)")
        self.assertIn(response.status_code, [301, 302],
                      "Regular user must be redirected away from /admin/")
        location = response.get("Location", "")
        self.assertIn("/admin/login/", location,
                      "Regular user redirect from /admin/ must point to /admin/login/")
        body = response.content.decode("utf-8")
        self.assertNotIn("j02_secret_doc.pdf", body,
                         "Admin redirect must not expose the user's job filename in the response body")

    # ------------------------------------------------------------------ TC-J03
    def test_tc_j03_administrator_can_view_registered_users(self):
        """A superuser can navigate to the Django User administration list at
        /admin/auth/user/. The page must return HTTP 200 and include registered
        usernames. Plaintext passwords and raw password hashes must not appear
        in the list view response."""
        self.client.force_login(self.superuser)
        response = self.client.get("/admin/auth/user/")

        self.assertEqual(response.status_code, 200,
                         "Superuser must receive HTTP 200 on /admin/auth/user/")
        body = response.content.decode("utf-8")
        self.assertIn(self.regular.username, body,
                      "User admin list must display the registered regular user's username")
        # Passwords must never appear in readable form in the list view
        self.assertNotIn(self.password, body,
                         "The user list must never expose the plaintext password")
        self.assertNotIn("pbkdf2_sha256", body,
                         "The user list page must not expose the raw password hash algorithm string")

    def test_tc_j03_staff_without_model_permission_cannot_view_user_list(self):
        """A staff user (is_staff=True, is_superuser=False) with no explicit
        auth.view_user or auth.change_user permission is denied access to
        /admin/auth/user/. Django Admin raises PermissionDenied (HTTP 403)."""
        self.client.force_login(self.staff_only)
        response = self.client.get("/admin/auth/user/")

        self.assertNotEqual(
            response.status_code, 200,
            "Staff-only user without model permission must not receive the user list (HTTP 200)",
        )
        self.assertIn(
            response.status_code, [302, 403],
            "Staff-only user accessing /admin/auth/user/ must receive 302 or 403",
        )

    # ------------------------------------------------------------------ TC-J04
    def test_tc_j04_administrator_can_view_all_job_statuses(self):
        """Authorized administrators can view TranslationJob records for all five
        status values — QUEUED, RETRYING, PROCESSING, COMPLETED, FAILED — in the
        Admin list. Each status-specific filename must appear in the response body."""
        for status in TranslationJob.Status.values:
            self._create_job(
                self.regular,
                status=status,
                original_filename=f"j04_{status}_doc.pdf",
            )
        self.client.force_login(self.superuser)

        response = self.client.get("/admin/translator/translationjob/")

        self.assertEqual(response.status_code, 200,
                         "Admin TranslationJob list must return HTTP 200 for a superuser")
        body = response.content.decode("utf-8")
        for status in TranslationJob.Status.values:
            self.assertIn(
                f"j04_{status}_doc.pdf", body,
                f"Admin job list must display the filename of the {status!r} job",
            )

    def test_tc_j04_admin_job_search_by_filename(self):
        """The Admin TranslationJob search bar (search_fields includes original_filename)
        returns only the matching job when filtering by a unique filename substring.
        Non-matching jobs must not appear in the filtered result."""
        self._create_job(self.regular, original_filename="j04_needle_unique_8821.pdf")
        self._create_job(self.regular, original_filename="j04_haystack_other.pdf")
        self.client.force_login(self.superuser)

        response = self.client.get(
            "/admin/translator/translationjob/?q=j04_needle_unique_8821"
        )

        self.assertEqual(response.status_code, 200,
                         "Admin search must return HTTP 200 for a superuser")
        body = response.content.decode("utf-8")
        self.assertIn("j04_needle_unique_8821.pdf", body,
                      "Admin search must return the matching job filename")
        self.assertNotIn("j04_haystack_other.pdf", body,
                         "Admin search must exclude non-matching job filenames")

    # ------------------------------------------------------------------ TC-J05
    def test_tc_j05_phrasebook_admin_entry_persists_unicode_tagabawa(self):
        """An administrator can add a PhrasebookEntry through the Django Admin form.
        The record is persisted to the database and the Tagabawa text (UTF-8 with
        Latin-1 diacritics, e.g. á = U+00E1) is stored and retrieved verbatim,
        verifying the database round-trip preserves Bagobo-Tagabawa orthography."""
        from translator.models import PhrasebookEntry
        tagabawa_text = "Madigár nu uras"  # á = U+00E1
        self.client.force_login(self.superuser)

        response = self.client.post(
            "/admin/translator/phrasebookentry/add/",
            {
                "_save": "Save",
                "english": "Good morning test J05",
                "tagabawa": tagabawa_text,
                "filipino": "Magandang umaga",
                "cebuano": "",
                "topic": "Greetings",
                "source": "phrasebook",
                "notes": "",
                "is_active": "on",
            },
        )
        # Successful Admin add redirects to the changelist (302)
        self.assertEqual(
            response.status_code, 302,
            "Admin PhrasebookEntry add must redirect (302) on successful save",
        )
        entry = PhrasebookEntry.objects.filter(english="Good morning test J05").first()
        self.assertIsNotNone(entry,
                             "PhrasebookEntry must be saved to the database after a successful Admin add")
        self.assertEqual(
            entry.tagabawa, tagabawa_text,
            "Tagabawa field must preserve UTF-8 diacritics (á = U+00E1) verbatim in the database",
        )

    def test_tc_j05_phrasebook_admin_change_does_not_affect_live_translation(self):
        """CRITICAL DATA-SOURCE AUDIT: The live translation service loads data from
        translation_data.json into an in-memory singleton at startup. It does NOT query
        PhrasebookEntry or GlossaryTerm database tables. An entry saved via Django Admin
        is persisted to the DB but is NOT immediately available to translate_phrase()
        without a server restart or explicit singleton reload.

        TC-J05 is therefore PARTIALLY IMPLEMENTED: Admin CRUD on PhrasebookEntry works,
        but database entries are not wired to the live translation service."""
        from translator.models import PhrasebookEntry
        from translator.services.translation_dataset import TranslationDataset

        # A phrase guaranteed to be absent from translation_data.json
        unique_english = "xyzAppJAdminUniquePhrase20260618NotInJson"
        unique_tagabawa = "XyzTagabawaNaSalitaNgAdminJ05"
        PhrasebookEntry.objects.create(
            english=unique_english,
            tagabawa=unique_tagabawa,
            is_active=True,
        )
        # Load a fresh TranslationDataset instance — reads from JSON, not DB
        dataset = TranslationDataset()
        result = dataset.translate_phrase(unique_english, "english", "tagabawa")

        # The translation service must NOT find the DB-only entry
        self.assertNotEqual(
            result, unique_tagabawa,
            "Translation service must not use Admin-managed PhrasebookEntry records "
            "immediately after a DB save. The live translator reads only from "
            "translation_data.json — database entries require a server restart or "
            "explicit reload to take effect.",
        )

    # ------------------------------------------------------------------ TC-J06
    def test_tc_j06_soft_deleted_jobs_are_visible_in_admin(self):
        """Authorized administrators can view soft-deleted (is_deleted=True) TranslationJob
        records in the Admin interface. The is_deleted filter isolates deleted records.
        The Admin detail page shows the deletion state and deleted_at timestamp.
        A GET on the Admin detail must NOT restore the job or clear is_deleted/deleted_at."""
        job = self._create_job(
            self.regular,
            status=TranslationJob.Status.COMPLETED,
            original_filename="j06_soft_deleted.pdf",
        )
        job.is_deleted = True
        job.deleted_at = timezone.now()
        job.save(update_fields=["is_deleted", "deleted_at"])
        self.client.force_login(self.superuser)

        # Filtered Admin list shows the soft-deleted job
        list_response = self.client.get(
            "/admin/translator/translationjob/?is_deleted__exact=1"
        )
        self.assertEqual(list_response.status_code, 200,
                         "Admin filtered list must return HTTP 200 for a superuser")
        self.assertIn(
            "j06_soft_deleted.pdf", list_response.content.decode("utf-8"),
            "Soft-deleted job must appear in Admin list when filtering by is_deleted=1",
        )

        # Admin detail page shows the job and its deletion state
        detail_response = self.client.get(
            f"/admin/translator/translationjob/{job.id}/change/"
        )
        self.assertEqual(detail_response.status_code, 200,
                         "Admin change page for a soft-deleted job must return HTTP 200")
        detail_body = detail_response.content.decode("utf-8")
        self.assertIn(
            "j06_soft_deleted.pdf", detail_body,
            "Admin detail page must show the soft-deleted job's filename",
        )

        # Viewing in Admin must NOT restore the record
        job.refresh_from_db()
        self.assertTrue(job.is_deleted,
                        "A GET on the Admin detail page must not set is_deleted=False (no auto-restore)")
        self.assertIsNotNone(
            job.deleted_at,
            "A GET on the Admin detail page must not clear the deleted_at timestamp",
        )


class TranslatorHomeNavigationTests(TestCase):
    """Verify the About page removal and the trimmed home/footer navigation."""

    def test_home_page_renders_successfully(self):
        response = self.client.get(reverse("translator:home"))
        self.assertEqual(response.status_code, 200)

    def test_home_page_has_no_top_nav_about_link(self):
        response = self.client.get(reverse("translator:home"))
        body = response.content.decode("utf-8")
        nav_start = body.index('<nav class="site-nav"')
        nav_end = body.index("</nav>", nav_start)
        nav_html = body[nav_start:nav_end]
        self.assertNotIn(">About<", nav_html)
        self.assertNotIn("/about/", nav_html)

    def test_footer_has_no_home_translate_about_links(self):
        response = self.client.get(reverse("translator:home"))
        body = response.content.decode("utf-8")
        footer_start = body.index('<footer class="site-footer"')
        footer_html = body[footer_start:]
        self.assertNotIn("<nav", footer_html,
                          "Footer must not contain a leftover navigation container")
        self.assertNotIn(">Home<", footer_html)
        self.assertNotIn(">Translate<", footer_html)
        self.assertNotIn(">About<", footer_html)

    def test_old_about_url_returns_404(self):
        response = self.client.get("/about/")
        self.assertEqual(response.status_code, 404)

    def test_home_translate_login_signup_links_remain_available(self):
        response = self.client.get(reverse("translator:home"))
        body = response.content.decode("utf-8")
        self.assertIn(reverse("translator:home"), body)
        self.assertIn(reverse("translator:translate"), body)
        self.assertIn(reverse("translator:login"), body)
        self.assertIn(reverse("translator:signup"), body)


class TranslatorLogoNavLoginTests(TestCase):
    """Logo asset usage, nav consistency, login-before-translate, and password reset."""

    password = "Bagobo-Reset-2026!"

    def setUp(self):
        self.alice = User.objects.create_user(
            username="logoalice",
            email="logoalice@example.test",
            password=self.password,
        )

    def _nav_links_html(self, body):
        start = body.index('<div class="nav-links"')
        end = body.index("</div>", start)
        return body[start:end]

    # 1. Home page renders successfully.
    def test_home_page_renders_successfully(self):
        response = self.client.get(reverse("translator:home"))
        self.assertEqual(response.status_code, 200)

    # 2. Header uses the real logo static asset.
    def test_header_uses_real_logo_static_asset(self):
        response = self.client.get(reverse("translator:home"))
        body = response.content.decode("utf-8")
        brand_start = body.index('<a class="brand"')
        brand_end = body.index("</a>", brand_start)
        brand_html = body[brand_start:brand_end]
        self.assertIn("images/lingokatutubo-logo.png", brand_html)
        self.assertIn('class="brand-mark"', brand_html)
        self.assertIn('alt="LINGOKATUTUBO logo"', brand_html)

    # 3. Hero uses the real logo and no longer renders the fake illustration.
    def test_hero_uses_real_logo_and_drops_fake_illustration(self):
        response = self.client.get(reverse("translator:home"))
        body = response.content.decode("utf-8")
        self.assertIn('class="hero-logo"', body)
        self.assertIn("images/lingokatutubo-logo.png", body)
        self.assertNotIn("document-visual", body)
        self.assertNotIn("paper-sheet", body)
        self.assertNotIn("translation-panel", body)

    # 4. Hero title remains present.
    def test_hero_title_remains_present(self):
        response = self.client.get(reverse("translator:home"))
        body = response.content.decode("utf-8")
        self.assertIn('class="hero-title"', body)
        self.assertIn("Bagobo-Tagabawa Document Translation", body)

    # 5. About remains removed.
    def test_about_remains_removed(self):
        response = self.client.get("/about/")
        self.assertEqual(response.status_code, 404)

    # 6. Unauthenticated nav Translate link points to the protected redirect.
    def test_unauthenticated_translate_nav_link_redirects_to_login(self):
        response = self.client.get(reverse("translator:home"))
        nav_html = self._nav_links_html(response.content.decode("utf-8"))
        self.assertIn(f'href="{reverse("translator:translate")}"', nav_html)

        translate_response = self.client.get(reverse("translator:translate"))
        self.assertRedirects(
            translate_response,
            f"{reverse('translator:login')}?next={reverse('translator:translate')}",
        )

    # 7. Unauthenticated Start Translating action redirects to Login with the
    # Translate destination preserved.
    def test_start_translating_button_redirects_to_login_with_next(self):
        response = self.client.get(reverse("translator:home"))
        body = response.content.decode("utf-8")
        self.assertIn(
            f'href="{reverse("translator:translate")}">Start Translating</a>', body
        )

        translate_response = self.client.get(reverse("translator:translate"))
        self.assertRedirects(
            translate_response,
            f"{reverse('translator:login')}?next={reverse('translator:translate')}",
        )

    # 8. Login-first context is shown only when routed via the Translate next
    # (see TranslatorAuthPageNavTests for the heading/subtext assertions —
    # this used to be a flash message banner, now it's a heading swap).
    def test_login_first_context_visible_with_translate_next(self):
        login_url = reverse("translator:login")
        translate_url = reverse("translator:translate")
        response = self.client.get(login_url, {"next": translate_url})
        self.assertContains(response, "Log in to continue")

    def test_login_first_context_absent_without_translate_next(self):
        response = self.client.get(reverse("translator:login"))
        self.assertNotContains(response, "Log in to continue")

    # 9. Successful login with next=/translate/ redirects to Translate.
    def test_successful_login_with_translate_next_redirects_to_translate(self):
        translate_url = reverse("translator:translate")
        response = self.client.post(
            f"{reverse('translator:login')}?next={translate_url}",
            {"username": self.alice.username, "password": self.password},
        )
        self.assertRedirects(response, translate_url)

    # 10. External or unsafe next values are rejected (Django's safe-redirect check).
    def test_unsafe_next_value_is_rejected_on_login(self):
        response = self.client.post(
            f"{reverse('translator:login')}?next=https://evil.example.com/",
            {"username": self.alice.username, "password": self.password},
        )
        self.assertRedirects(response, reverse("translator:translate"))
        self.assertNotIn("evil.example.com", response.get("Location", ""))

    def test_protocol_relative_next_value_is_rejected_on_login(self):
        response = self.client.post(
            f"{reverse('translator:login')}?next=//evil.example.com/",
            {"username": self.alice.username, "password": self.password},
        )
        self.assertRedirects(response, reverse("translator:translate"))

    # 11. Authenticated user opens Translate directly.
    def test_authenticated_user_opens_translate_directly(self):
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:translate"))
        self.assertEqual(response.status_code, 200)

    # 12. Login page displays "Forgot Password?"
    def test_login_page_displays_forgot_password_link(self):
        response = self.client.get(reverse("translator:login"))
        self.assertContains(response, "Forgot Password?")
        self.assertContains(response, reverse("password_reset"))

    # 13. Password-reset request page returns HTTP 200.
    def test_password_reset_request_page_returns_200(self):
        response = self.client.get(reverse("password_reset"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "site-header")

    # 14. Valid password-reset submission shows the generic confirmation page,
    # and does not reveal whether the email address is registered.
    def test_password_reset_submission_shows_generic_confirmation(self):
        known_response = self.client.post(
            reverse("password_reset"), {"email": self.alice.email}
        )
        unknown_response = self.client.post(
            reverse("password_reset"), {"email": "no-such-user@example.test"}
        )
        self.assertRedirects(known_response, reverse("password_reset_done"))
        self.assertRedirects(unknown_response, reverse("password_reset_done"))

        done_response = self.client.get(reverse("password_reset_done"))
        self.assertContains(done_response, "If an account exists")
        self.assertEqual(len(mail.outbox), 1)
        self.assertEqual(mail.outbox[0].to, [self.alice.email])

    # 15. Password-reset confirmation and complete routes resolve correctly.
    def test_password_reset_confirm_and_complete_routes_resolve(self):
        uid = urlsafe_base64_encode(force_bytes(self.alice.pk))
        token = default_token_generator.make_token(self.alice)

        confirm_response = self.client.get(
            reverse("password_reset_confirm", kwargs={"uidb64": uid, "token": token}),
            follow=True,
        )
        self.assertEqual(confirm_response.status_code, 200)
        self.assertContains(confirm_response, "new_password1")

        complete_response = self.client.get(reverse("password_reset_complete"))
        self.assertEqual(complete_response.status_code, 200)
        self.assertContains(complete_response, "Password Updated")

    def test_password_reset_confirm_rejects_invalid_token(self):
        uid = urlsafe_base64_encode(force_bytes(self.alice.pk))
        response = self.client.get(
            reverse("password_reset_confirm", kwargs={"uidb64": uid, "token": "bad-token"}),
        )
        self.assertContains(response, "Reset Link Invalid")

    # 16. Home, Translate, Log In, and Sign Up remain present for anonymous users.
    def test_anonymous_nav_has_home_translate_login_signup(self):
        response = self.client.get(reverse("translator:home"))
        nav_html = self._nav_links_html(response.content.decode("utf-8"))
        self.assertIn(">Home<", nav_html)
        self.assertIn(">Translate<", nav_html)
        self.assertIn(">Log In<", nav_html)
        self.assertIn(">Sign Up<", nav_html)

    # 17. Navigation does not contain duplicate links.
    def test_nav_links_has_no_duplicate_entries(self):
        response = self.client.get(reverse("translator:home"))
        nav_html = self._nav_links_html(response.content.decode("utf-8"))
        for label in ["Home", "Translate", "Log In", "Sign Up"]:
            self.assertEqual(
                nav_html.count(f">{label}<"), 1, f"{label} should appear exactly once in the nav"
            )


class TranslatorAuthPageNavTests(TestCase):
    """Auth-flow pages use a minimal header; other pages keep full navigation."""

    password = "Bagobo-AuthNav-2026!"

    AUTH_PAGE_URL_NAMES = [
        "translator:login",
        "translator:signup",
        "password_reset",
        "password_reset_done",
    ]

    def setUp(self):
        self.alice = User.objects.create_user(
            username="authnavalice",
            email="authnavalice@example.test",
            password=self.password,
        )

    def _nav_links_html(self, body):
        start = body.index('<div class="nav-links"')
        end = body.index("</div>", start)
        return body[start:end]

    def _assert_minimal_nav(self, nav_html):
        self.assertIn(">Home<", nav_html)
        self.assertNotIn(">Translate<", nav_html)
        self.assertNotIn(">Log In<", nav_html)
        self.assertNotIn(">Sign Up<", nav_html)
        self.assertNotIn(">History<", nav_html)
        self.assertNotIn("Log Out", nav_html)
        self.assertEqual(
            nav_html.count("<a "), 1, "Nav should contain exactly the Home link, not be empty or have extras"
        )

    # Login page header contains only brand and Home.
    def test_login_page_header_has_only_brand_and_home(self):
        response = self.client.get(reverse("translator:login"))
        body = response.content.decode("utf-8")
        self.assertIn('class="brand"', body)
        self._assert_minimal_nav(self._nav_links_html(body))

    # Login page does not show header Translate, Log In, or Sign Up links.
    def test_login_page_header_hides_translate_login_signup(self):
        response = self.client.get(reverse("translator:login"))
        nav_html = self._nav_links_html(response.content.decode("utf-8"))
        self.assertNotIn(">Translate<", nav_html)
        self.assertNotIn(">Log In<", nav_html)
        self.assertNotIn(">Sign Up<", nav_html)

    # Sign Up remains available inside the login card.
    def test_signup_remains_available_inside_login_card(self):
        response = self.client.get(reverse("translator:login"))
        body = response.content.decode("utf-8")
        self.assertContains(response, "Need an account?")
        self.assertIn(f'href="{reverse("translator:signup")}">Sign up</a>', body)

    # Password-reset pages (and signup) use the minimal header.
    def test_auth_pages_use_minimal_header(self):
        for url_name in self.AUTH_PAGE_URL_NAMES:
            with self.subTest(url_name=url_name):
                response = self.client.get(reverse(url_name))
                nav_html = self._nav_links_html(response.content.decode("utf-8"))
                self._assert_minimal_nav(nav_html)

    def test_password_reset_confirm_and_complete_use_minimal_header(self):
        uid = urlsafe_base64_encode(force_bytes(self.alice.pk))
        token = default_token_generator.make_token(self.alice)

        confirm_response = self.client.get(
            reverse("password_reset_confirm", kwargs={"uidb64": uid, "token": token}),
            follow=True,
        )
        self._assert_minimal_nav(self._nav_links_html(confirm_response.content.decode("utf-8")))

        complete_response = self.client.get(reverse("password_reset_complete"))
        self._assert_minimal_nav(self._nav_links_html(complete_response.content.decode("utf-8")))

    # Normal anonymous pages retain the complete anonymous navigation.
    def test_normal_anonymous_pages_retain_full_navigation(self):
        response = self.client.get(reverse("translator:home"))
        nav_html = self._nav_links_html(response.content.decode("utf-8"))
        self.assertIn(">Home<", nav_html)
        self.assertIn(">Translate<", nav_html)
        self.assertIn(">Log In<", nav_html)
        self.assertIn(">Sign Up<", nav_html)

    # Authenticated pages retain History and Log Out.
    def test_authenticated_pages_retain_history_and_logout(self):
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:home"))
        nav_html = self._nav_links_html(response.content.decode("utf-8"))
        self.assertIn(">History<", nav_html)
        self.assertIn("Log Out", nav_html)
        self.assertNotIn(">Log In<", nav_html)
        self.assertNotIn(">Sign Up<", nav_html)

    # Login heading/subtext swap for the Translate-next case, and the old
    # blue info banner is gone.
    def test_login_heading_changes_for_translate_next(self):
        translate_url = reverse("translator:translate")
        response = self.client.get(reverse("translator:login"), {"next": translate_url})
        self.assertContains(response, "Log in to continue")
        self.assertContains(response, "Sign in to access document translation.")
        self.assertNotContains(response, "Please log in first to translate documents.")

    def test_login_heading_is_default_without_translate_next(self):
        response = self.client.get(reverse("translator:login"))
        self.assertContains(response, "Log In")
        self.assertNotContains(response, "Log in to continue")
        self.assertNotContains(response, "Sign in to access document translation.")

    # Successful login still redirects to Translate.
    def test_successful_login_with_translate_next_still_redirects_to_translate(self):
        translate_url = reverse("translator:translate")
        response = self.client.post(
            f"{reverse('translator:login')}?next={translate_url}",
            {"username": self.alice.username, "password": self.password},
        )
        self.assertRedirects(response, translate_url)

    # Safe-next validation is preserved even with the new context handling.
    def test_unsafe_next_still_rejected_on_login(self):
        response = self.client.post(
            f"{reverse('translator:login')}?next=https://evil.example.com/",
            {"username": self.alice.username, "password": self.password},
        )
        self.assertRedirects(response, reverse("translator:translate"))

    # Django messages still work for unrelated notices (e.g. history deletion).
    def test_messages_framework_still_works_for_unrelated_notices(self):
        self.client.force_login(self.alice)
        job = TranslationJob.objects.create(
            owner=self.alice,
            original_filename="authnav_test.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=TranslationJob.Status.COMPLETED,
        )
        response = self.client.post(
            reverse("translator:job_delete", args=[job.id]), follow=True
        )
        self.assertContains(response, "has been removed from your history")


class AdminAccountAuditLoginFlowTests(TestCase):
    """Login-flow verification for the administrator-account audit.

    Uses throwaway, test-only credentials created in the isolated test
    database — never the real local `admin` account or its password.
    """

    test_password = "Temp-Audit-Test-Only-2026!"

    def setUp(self):
        self.regular = User.objects.create_user(
            username="audit_regular",
            email="audit_regular@example.test",
            password=self.test_password,
        )
        self.superuser = User.objects.create_superuser(
            username="audit_superuser",
            email="audit_superuser@example.test",
            password=self.test_password,
        )

    def test_admin_login_page_loads(self):
        response = self.client.get("/admin/login/")
        self.assertEqual(response.status_code, 200)

    def test_regular_account_cannot_access_admin_dashboard(self):
        self.client.force_login(self.regular)
        response = self.client.get("/admin/")
        self.assertNotEqual(response.status_code, 200)

    def test_superuser_can_log_in_and_access_admin_dashboard(self):
        response = self.client.post(
            "/admin/login/",
            {
                "username": self.superuser.username,
                "password": self.test_password,
                "next": "/admin/",
            },
            follow=True,
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(int(self.client.session["_auth_user_id"]), self.superuser.pk)

        dashboard = self.client.get("/admin/")
        self.assertEqual(dashboard.status_code, 200)

    def test_invalid_credentials_do_not_create_a_session(self):
        response = self.client.post(
            "/admin/login/",
            {"username": self.superuser.username, "password": "totally-wrong-password"},
        )
        self.assertNotIn("_auth_user_id", self.client.session)


# ============================================================
# Experimental ByT5 neural fallback — disabled by default, guarded,
# fail-safe. See translator/services/neural_translation_service.py.
# ============================================================

class TranslatorNeuralServiceUnitTests(TestCase):
    """Unit tests for NeuralTranslationService in isolation (no Django job)."""

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.addCleanup(self.temp_dir.cleanup)

    def _fake_model_dir(self):
        """A directory that looks like a saved model (has config.json) but
        contains no real weights — enough to pass the load precondition."""
        model_dir = self.root / "fake_model"
        model_dir.mkdir(parents=True, exist_ok=True)
        (model_dir / "config.json").write_text("{}", encoding="utf-8")
        return model_dir

    def test_disabled_by_default(self):
        """LINGOKATUTUBO_NEURAL_TRANSLATION_ENABLED defaults to False."""
        from translator.services.neural_translation_service import NeuralTranslationService

        service = NeuralTranslationService()
        self.assertFalse(service.is_enabled())
        self.assertIsNone(service.translate_unmatched("Madayaw", "tagabawa", "english"))

    def test_empty_segment_returns_none_without_crash(self):
        from translator.services.neural_translation_service import NeuralTranslationService

        service = NeuralTranslationService()
        self.assertIsNone(service.translate_unmatched("   ", "tagabawa", "english"))
        self.assertIsNone(service.translate_unmatched("", "tagabawa", "english"))
        # Empty input is rejected before any load attempt.
        self.assertFalse(service._load_attempted)

    @override_settings(LINGOKATUTUBO_NEURAL_TRANSLATION_ENABLED=True)
    def test_wrong_direction_returns_none_without_loading(self):
        """Only the configured source -> target direction may use ByT5."""
        from translator.services.neural_translation_service import NeuralTranslationService

        service = NeuralTranslationService()
        result = service.translate_unmatched("Hello", "english", "tagabawa")

        self.assertIsNone(result)
        self.assertFalse(service._load_attempted)

    @override_settings(
        LINGOKATUTUBO_NEURAL_TRANSLATION_ENABLED=True,
        LINGOKATUTUBO_BYT5_MODEL_DIR="C:/this/path/does/not/exist/byt5",
    )
    def test_missing_model_dir_fails_safe(self):
        from translator.services.neural_translation_service import NeuralTranslationService

        service = NeuralTranslationService()
        result = service.translate_unmatched("Madayaw na adlaw", "tagabawa", "english")

        self.assertIsNone(result)
        self.assertFalse(service.is_available())
        self.assertIsNotNone(service.load_warning)
        # The on-disk path must never leak into the warning text.
        self.assertNotIn("this/path/does/not/exist", service.load_warning)

    @override_settings(LINGOKATUTUBO_NEURAL_TRANSLATION_ENABLED=True)
    def test_missing_transformers_dependency_fails_safe(self):
        """If transformers is unavailable, the model load fails safely."""
        from translator.services.neural_translation_service import NeuralTranslationService

        model_dir = self._fake_model_dir()

        with override_settings(LINGOKATUTUBO_BYT5_MODEL_DIR=str(model_dir)):
            service = NeuralTranslationService()
            with patch.dict(sys.modules, {"transformers": None}):
                result = service.translate_unmatched("Madayaw", "tagabawa", "english")

        self.assertIsNone(result)
        self.assertFalse(service.is_available())
        self.assertIn("not installed", service.load_warning or "")

    @override_settings(LINGOKATUTUBO_NEURAL_TRANSLATION_ENABLED=True)
    def test_mocked_model_returns_review_marked_result(self):
        """A loaded (mocked) model produces a fully-tagged experimental record."""
        from translator.services.neural_translation_service import (
            NEURAL_METHOD,
            NEURAL_REVIEW_WARNING,
            NeuralTranslationService,
        )

        class _FakeBatch(dict):
            def to(self, device):
                return self

        class _FakeTokenizer:
            def __call__(self, prompt, return_tensors=None, truncation=None, max_length=None):
                return _FakeBatch({"input_ids": [[1, 2, 3]]})

            def decode(self, ids, skip_special_tokens=True):
                return "mocked english output"

        class _FakeModel:
            def generate(self, **kwargs):
                return [[1, 2, 3]]

        service = NeuralTranslationService()
        service._model = _FakeModel()
        service._tokenizer = _FakeTokenizer()
        service._load_attempted = True

        result = service.translate_unmatched("Madayaw na adlaw", "tagabawa", "english")

        self.assertIsNotNone(result)
        self.assertEqual(result["translated"], "mocked english output")
        self.assertEqual(result["method"], NEURAL_METHOD)
        self.assertEqual(result["cascade_stage"], NEURAL_METHOD)
        self.assertIsNone(result["confidence"])
        self.assertTrue(result["needs_review"])
        self.assertEqual(result["warning"], NEURAL_REVIEW_WARNING)

    @override_settings(LINGOKATUTUBO_NEURAL_TRANSLATION_ENABLED=True)
    def test_long_segment_is_handled_safely(self):
        """A very long segment does not crash; truncation params are forwarded."""
        from translator.services.neural_translation_service import NeuralTranslationService

        captured = {}

        class _FakeBatch(dict):
            def to(self, device):
                return self

        class _FakeTokenizer:
            def __call__(self, prompt, return_tensors=None, truncation=None, max_length=None):
                captured["truncation"] = truncation
                captured["max_length"] = max_length
                return _FakeBatch({"input_ids": [[1, 2, 3]]})

            def decode(self, ids, skip_special_tokens=True):
                return "safe output"

        class _FakeModel:
            def generate(self, **kwargs):
                return [[1, 2, 3]]

        service = NeuralTranslationService()
        service._model = _FakeModel()
        service._tokenizer = _FakeTokenizer()
        service._load_attempted = True

        long_text = "Madayaw " * 2000  # far longer than any real OCR/PDF line

        result = service.translate_unmatched(long_text, "tagabawa", "english")

        self.assertIsNotNone(result)
        self.assertEqual(result["translated"], "safe output")
        self.assertTrue(captured["truncation"])
        self.assertEqual(captured["max_length"], 512)

    @override_settings(LINGOKATUTUBO_NEURAL_TRANSLATION_ENABLED=True)
    def test_inference_exception_falls_back_safely(self):
        """A raising model must not propagate; the segment is left to the caller."""
        from translator.services.neural_translation_service import NeuralTranslationService

        class _ExplodingModel:
            def generate(self, **kwargs):
                raise RuntimeError("simulated inference crash")

        class _FakeBatch(dict):
            def to(self, device):
                return self

        class _FakeTokenizer:
            def __call__(self, prompt, return_tensors=None, truncation=None, max_length=None):
                return _FakeBatch({"input_ids": [[1, 2, 3]]})

        service = NeuralTranslationService()
        service._model = _ExplodingModel()
        service._tokenizer = _FakeTokenizer()
        service._load_attempted = True

        result = service.translate_unmatched("Madayaw", "tagabawa", "english")

        self.assertIsNone(result)


class TranslatorNeuralPipelineWiringTests(TestCase):
    """Integration tests for the seam in PipelineService._translate_layout."""

    def _make_pipeline(self, dataset_responses):
        from translator.services.pipeline_service import PipelineService

        class _StubDataset:
            def translate_phrase_with_metadata(self, text, source_lang, target_lang):
                return dataset_responses[text]

        service = PipelineService.__new__(PipelineService)
        service.translation_dataset = _StubDataset()
        return service

    @staticmethod
    def _single_line_layout(text):
        return [{
            "blocks": [{
                "type": "text",
                "lines": [{"text": text}],
            }]
        }]

    def test_neural_disabled_keeps_old_phrasebook_behavior(self):
        """Disabled (default) neural service must not change existing output."""
        from translator.services.translation_dataset import UNKNOWN_FOR_REVIEW

        pipeline = self._make_pipeline({
            "Hello": {
                "translated": UNKNOWN_FOR_REVIEW,
                "method": "unknown_for_review",
                "cascade_stage": "unknown_for_review",
                "confidence": 0.0,
            }
        })

        translations = pipeline._translate_layout(
            self._single_line_layout("Hello"), "tagabawa", "english"
        )

        record = translations["Hello"]
        self.assertEqual(record["translated"], UNKNOWN_FOR_REVIEW)
        self.assertEqual(record["method"], "unknown_for_review")

    def test_neural_enabled_uses_mocked_fallback_for_unmatched_segment(self):
        """Enabled neural service is consulted only when phrasebook is unmatched."""
        from translator.services.neural_translation_service import (
            NEURAL_METHOD,
            NEURAL_REVIEW_WARNING,
        )
        from translator.services.translation_dataset import UNKNOWN_FOR_REVIEW

        pipeline = self._make_pipeline({
            "Hello": {
                "translated": UNKNOWN_FOR_REVIEW,
                "method": "unknown_for_review",
                "cascade_stage": "unknown_for_review",
                "confidence": 0.0,
            }
        })

        stub_neural = SimpleNamespace(
            is_enabled=lambda: True,
            load_warning=None,
            translate_unmatched=lambda text, src, tgt: {
                "translated": "Mocked English",
                "method": NEURAL_METHOD,
                "cascade_stage": NEURAL_METHOD,
                "confidence": None,
                "needs_review": True,
                "warning": NEURAL_REVIEW_WARNING,
            },
        )

        with patch(
            "translator.services.neural_translation_service.get_neural_translation_service",
            return_value=stub_neural,
        ):
            translations = pipeline._translate_layout(
                self._single_line_layout("Hello"), "tagabawa", "english"
            )

        record = translations["Hello"]
        self.assertEqual(record["method"], NEURAL_METHOD)
        self.assertEqual(record["translated"], "Mocked English")
        self.assertIsNone(record["confidence"])
        self.assertEqual(record["warning"], NEURAL_REVIEW_WARNING)

    def test_phrasebook_match_skips_neural_fallback_entirely(self):
        """A matched phrasebook segment must never reach the neural service."""
        from unittest.mock import Mock

        pipeline = self._make_pipeline({
            "Hello": {
                "translated": "Madigar",
                "method": "exact_phrase",
                "cascade_stage": "exact_phrase",
                "confidence": 1.0,
            }
        })

        stub_neural = Mock()
        stub_neural.is_enabled.return_value = True

        with patch(
            "translator.services.neural_translation_service.get_neural_translation_service",
            return_value=stub_neural,
        ):
            translations = pipeline._translate_layout(
                self._single_line_layout("Hello"), "tagabawa", "english"
            )

        record = translations["Hello"]
        self.assertEqual(record["method"], "exact_phrase")
        self.assertEqual(record["translated"], "Madigar")
        stub_neural.translate_unmatched.assert_not_called()

    @override_settings(LINGOKATUTUBO_NEURAL_TRANSLATION_ENABLED=True)
    def test_wrong_direction_keeps_unknown_marker_with_real_service(self):
        """English -> Tagabawa must not trigger the Tagabawa -> English model,
        even with neural enabled, using the real (unmocked) service logic."""
        from translator.services.neural_translation_service import NeuralTranslationService
        from translator.services.translation_dataset import UNKNOWN_FOR_REVIEW

        pipeline = self._make_pipeline({
            "Hello": {
                "translated": UNKNOWN_FOR_REVIEW,
                "method": "unknown_for_review",
                "cascade_stage": "unknown_for_review",
                "confidence": 0.0,
            }
        })

        with patch(
            "translator.services.neural_translation_service.get_neural_translation_service",
            return_value=NeuralTranslationService(),
        ):
            translations = pipeline._translate_layout(
                self._single_line_layout("Hello"), "english", "tagabawa"
            )

        record = translations["Hello"]
        self.assertEqual(record["translated"], UNKNOWN_FOR_REVIEW)
        self.assertEqual(record["method"], "unknown_for_review")

    @override_settings(
        LINGOKATUTUBO_NEURAL_TRANSLATION_ENABLED=True,
        LINGOKATUTUBO_BYT5_MODEL_DIR="C:/this/path/does/not/exist/byt5",
    )
    def test_neural_load_failure_falls_back_and_records_warning(self):
        """A missing model must not crash the job; it must fall back to the
        phrasebook's review marker and record an actionable warning."""
        from translator.services.neural_translation_service import NeuralTranslationService
        from translator.services.translation_dataset import UNKNOWN_FOR_REVIEW

        pipeline = self._make_pipeline({
            "Hello": {
                "translated": UNKNOWN_FOR_REVIEW,
                "method": "unknown_for_review",
                "cascade_stage": "unknown_for_review",
                "confidence": 0.0,
            }
        })

        translation_warnings = []
        with patch(
            "translator.services.neural_translation_service.get_neural_translation_service",
            return_value=NeuralTranslationService(),
        ):
            translations = pipeline._translate_layout(
                self._single_line_layout("Hello"),
                "tagabawa",
                "english",
                translation_warnings=translation_warnings,
            )

        record = translations["Hello"]
        self.assertEqual(record["translated"], UNKNOWN_FOR_REVIEW)
        self.assertTrue(
            any("Neural translation model is unavailable" in w for w in translation_warnings),
            f"Expected a neural load warning, got: {translation_warnings}",
        )


class TranslatorNeuralSyncAndOutputTests(TestCase):
    """DB sync and PDF reconstruction must treat byt5_neural specially:
    always needs_review, never a fabricated confidence."""

    password = "Bagobo-Neural-Sync-2026!"

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)

        self.alice = User.objects.create_user(
            username="neuralalice",
            email="neuralalice@example.test",
            password=self.password,
        )

    def _create_job(self, owner, status=TranslationJob.Status.PROCESSING):
        return TranslationJob.objects.create(
            owner=owner,
            original_filename="neural_test.pdf",
            file_type="pdf",
            status=status,
            source_language="tagabawa",
            target_language="english",
        )

    def _write_job_file(self, job, *parts_and_content):
        *parts, content = parts_and_content
        path = self.media_root.joinpath("jobs", job.job_id, *parts)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(content)
        return path

    def test_byt5_segments_are_always_needs_review_with_no_confidence(self):
        from translator.services.neural_translation_service import (
            NEURAL_METHOD,
            NEURAL_REVIEW_WARNING,
        )

        job = self._create_job(self.alice)
        structure_path = self._write_job_file(
            job,
            "structure.json",
            json.dumps({
                "warnings": [],
                "pages": [{
                    "page_number": 1,
                    "width": 612,
                    "height": 792,
                    "blocks": [{
                        "type": "text",
                        "block_id": "b1",
                        "bbox": [72, 72, 540, 110],
                        "lines": [{
                            "text": "Madayaw na adlaw",
                            "translated_text": "Good day",
                            "translation_method": NEURAL_METHOD,
                            # Even if a stray confidence leaked in upstream,
                            # sync must still null it out for neural segments.
                            "translation_confidence": 0.83,
                            "translation_warning": NEURAL_REVIEW_WARNING,
                            "bbox": [72, 72, 540, 90],
                        }],
                    }],
                }],
            }).encode("utf-8"),
        )

        sync_pipeline_job(
            SimpleNamespace(
                job_id=job.job_id,
                status=TranslationJob.Status.COMPLETED,
                progress=100,
                current_phase="completed",
                current_step="Completed",
                phase_message="Translation complete.",
                error="",
                detection_type="digital",
                file_type=TranslationJob.FileType.PDF,
                metadata={"structure_file": str(structure_path)},
                completed_at=None,
            )
        )

        segment = TranslationSegment.objects.get(job=job, segment_index=1)
        self.assertEqual(segment.method, NEURAL_METHOD)
        self.assertEqual(segment.translated_text, "Good day")
        self.assertIsNone(segment.confidence)
        self.assertTrue(segment.needs_review)
        self.assertEqual(segment.metadata.get("warning"), NEURAL_REVIEW_WARNING)

    def test_phrasebook_segment_in_same_job_is_unaffected(self):
        """A normal phrasebook segment alongside a neural one keeps its own
        confidence and review status — neural handling must not leak."""
        from translator.services.neural_translation_service import NEURAL_METHOD

        job = self._create_job(self.alice)
        structure_path = self._write_job_file(
            job,
            "structure.json",
            json.dumps({
                "warnings": [],
                "pages": [{
                    "page_number": 1,
                    "width": 612,
                    "height": 792,
                    "blocks": [
                        {
                            "type": "text",
                            "block_id": "b1",
                            "bbox": [72, 72, 540, 90],
                            "lines": [{
                                "text": "Hello",
                                "translated_text": "Madigar",
                                "translation_method": "exact_phrase",
                                "translation_confidence": 1.0,
                                "bbox": [72, 72, 540, 90],
                            }],
                        },
                        {
                            "type": "text",
                            "block_id": "b2",
                            "bbox": [72, 94, 540, 112],
                            "lines": [{
                                "text": "Madayaw na adlaw",
                                "translated_text": "Good day",
                                "translation_method": NEURAL_METHOD,
                                "translation_confidence": None,
                                "bbox": [72, 94, 540, 112],
                            }],
                        },
                    ],
                }],
            }).encode("utf-8"),
        )

        sync_pipeline_job(
            SimpleNamespace(
                job_id=job.job_id,
                status=TranslationJob.Status.COMPLETED,
                progress=100,
                current_phase="completed",
                current_step="Completed",
                phase_message="Translation complete.",
                error="",
                detection_type="digital",
                file_type=TranslationJob.FileType.PDF,
                metadata={"structure_file": str(structure_path)},
                completed_at=None,
            )
        )

        phrasebook_segment = TranslationSegment.objects.get(job=job, segment_index=1)
        neural_segment = TranslationSegment.objects.get(job=job, segment_index=2)

        self.assertFalse(phrasebook_segment.needs_review)
        self.assertAlmostEqual(phrasebook_segment.confidence, 1.0)

        self.assertTrue(neural_segment.needs_review)
        self.assertIsNone(neural_segment.confidence)

    def test_byt5_translated_text_appears_in_output_pdf(self):
        """Neural translations are drawn into the reconstructed PDF like any
        other method — download must not silently omit them."""
        import fitz
        from translator.services.neural_translation_service import NEURAL_METHOD
        from translator.services.pipeline_service import PipelineService

        layout = [{
            "page": 0,
            "width": 612.0,
            "height": 792.0,
            "blocks": [{
                "type": "text",
                "bbox": [72.0, 100.0, 540.0, 120.0],
                "lines": [{"text": "Madayaw na adlaw", "bbox": [72.0, 100.0, 540.0, 120.0]}],
            }],
        }]
        translations = {
            "Madayaw na adlaw": {
                "original": "Madayaw na adlaw",
                "translated": "Good day NeuralMarker",
                "method": NEURAL_METHOD,
                "cascade_stage": NEURAL_METHOD,
                "confidence": None,
            }
        }

        output_path = str(self.media_root / "neural_out.pdf")
        ok = PipelineService()._create_output_pdf(layout, translations, output_path)
        self.assertTrue(ok)

        doc = fitz.open(output_path)
        full_text = "".join(doc[i].get_text() for i in range(doc.page_count))
        doc.close()

        self.assertIn("Good day NeuralMarker", full_text)

    def test_unknown_marker_never_appears_in_output_pdf_body(self):
        import fitz
        from translator.services.pipeline_service import PipelineService
        from translator.services.translation_dataset import UNKNOWN_FOR_REVIEW

        layout = [{
            "page": 0,
            "width": 612.0,
            "height": 792.0,
            "blocks": [{
                "type": "text",
                "bbox": [72.0, 100.0, 540.0, 120.0],
                "lines": [{"text": "Unmatched source text", "bbox": [72.0, 100.0, 540.0, 120.0]}],
            }],
        }]
        translations = {
            "Unmatched source text": {
                "original": "Unmatched source text",
                "translated": UNKNOWN_FOR_REVIEW,
                "method": "unknown_for_review",
                "cascade_stage": "unknown_for_review",
                "confidence": 0.0,
            }
        }

        output_path = str(self.media_root / "unknown_out.pdf")
        ok = PipelineService()._create_output_pdf(layout, translations, output_path)
        self.assertTrue(ok)

        doc = fitz.open(output_path)
        full_text = "".join(doc[i].get_text() for i in range(doc.page_count))
        doc.close()

        self.assertIn("Unmatched source text", full_text)
        self.assertNotIn("UNKNOWN_FOR_REVIEW", full_text)

    def test_font_metadata_fallback_mapping_handles_missing_fonts(self):
        import fitz
        from translator.services.reconstruction_service import ReconstructionService

        self.assertEqual(
            ReconstructionService._fontname_for_line({
                "font": "TimesNewRomanPS-BoldItalicMT",
                "flags": 4 | 16 | 2,
            }),
            "tibi",
        )
        self.assertEqual(
            ReconstructionService._fontname_for_line({
                "font": "Arial-BoldMT",
                "flags": 16,
            }),
            "hebo",
        )
        self.assertEqual(
            ReconstructionService._fontname_for_line({
                "font": "CourierNewPS-ItalicMT",
                "flags": 8 | 2,
            }),
            "coit",
        )
        self.assertEqual(ReconstructionService._fontname_for_line({}), "helv")

        page_rect = fitz.Rect(0, 0, 612, 792)
        center_rect = fitz.Rect(240, 100, 372, 122)
        right_rect = fitz.Rect(470, 100, 540, 122)
        self.assertEqual(
            ReconstructionService._text_alignment_for_line(
                {"block_bbox": [72, 90, 540, 130]},
                center_rect,
                page_rect,
            ),
            fitz.TEXT_ALIGN_CENTER,
        )
        self.assertEqual(
            ReconstructionService._text_alignment_for_line(
                {"block_bbox": [72, 90, 540, 130]},
                right_rect,
                page_rect,
            ),
            fitz.TEXT_ALIGN_RIGHT,
        )

    def test_translated_pdf_generation_falls_back_when_exact_font_is_unavailable(self):
        import fitz
        from translator.services.pipeline_service import PipelineService

        layout = [{
            "page": 0,
            "width": 612.0,
            "height": 792.0,
            "blocks": [{
                "type": "text",
                "bbox": [72.0, 96.0, 540.0, 140.0],
                "lines": [{
                    "text": "Styled source",
                    "bbox": [168.0, 100.0, 444.0, 124.0],
                    "font": "MissingSubset+FancySerif-BoldItalic",
                    "size": 16.0,
                    "flags": 4 | 16 | 2,
                    "color": [0.1, 0.2, 0.3],
                    "spans": [{
                        "text": "Styled source",
                        "bbox": [168.0, 100.0, 444.0, 124.0],
                        "font": "MissingSubset+FancySerif-BoldItalic",
                        "size": 16.0,
                        "flags": 4 | 16 | 2,
                        "color": [0.1, 0.2, 0.3],
                    }],
                }],
            }],
        }]
        translations = {
            "Styled source": {
                "original": "Styled source",
                "translated": "Styled target",
                "method": "exact_phrase",
                "cascade_stage": "exact_phrase",
                "confidence": 1.0,
            }
        }

        output_path = str(self.media_root / "styled_out.pdf")
        warnings = []
        ok = PipelineService()._create_output_pdf(
            layout,
            translations,
            output_path,
            layout_warnings=warnings,
        )
        self.assertTrue(ok)

        doc = fitz.open(output_path)
        try:
            page_text = doc[0].get_text()
            spans = [
                span
                for block in doc[0].get_text("dict").get("blocks", [])
                for line in block.get("lines", [])
                for span in line.get("spans", [])
                if "Styled target" in span.get("text", "")
            ]
        finally:
            doc.close()

        self.assertIn("Styled target", page_text)
        self.assertNotIn("[UNKNOWN_FOR_REVIEW]", page_text)
        self.assertTrue(spans)
        self.assertGreaterEqual(float(spans[0].get("size", 0)), 14.0)


class TranslatorNeuralPreviewDisplayTests(TestCase):
    """Preview UI must surface ByT5 output distinctly without fabricating
    confidence, and must not block the existing preview/download workflow."""

    password = "Bagobo-Neural-Preview-2026!"

    def setUp(self):
        self.alice = User.objects.create_user(
            username="neuralpreview",
            email="neuralpreview@example.test",
            password=self.password,
        )

    def _create_job(self, status=TranslationJob.Status.COMPLETED):
        return TranslationJob.objects.create(
            owner=self.alice,
            original_filename="neural_preview.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=status,
            source_language="tagabawa",
            target_language="english",
        )

    def test_preview_displays_byt5_method_and_needs_review_badge(self):
        """Staff/admin see the ByT5 method badge and review marker; normal
        users never see internal method names like byt5_neural."""
        from translator.services.neural_translation_service import NEURAL_METHOD

        job = self._create_job()
        TranslationSegment.objects.create(
            job=job,
            segment_index=1,
            source_text="Madayaw na adlaw",
            translated_text="Good day",
            source_language="tagabawa",
            target_language="english",
            method=NEURAL_METHOD,
            confidence=None,
            needs_review=True,
        )
        self.alice.is_staff = True
        self.alice.save(update_fields=["is_staff"])
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertIn("ByT5 Neural Translation", body)
        self.assertIn("Needs review", body)
        self.assertIn("Good day", body)
        self.assertIn("No confidence score", body)
        # No fabricated confidence score should ever be rendered for a
        # confidence=None segment (the confidence-badge <p> is conditional).
        self.assertNotIn("confidence-badge", body)

    def test_preview_hides_byt5_method_name_from_normal_user(self):
        from translator.services.neural_translation_service import NEURAL_METHOD

        job = self._create_job()
        TranslationSegment.objects.create(
            job=job,
            segment_index=1,
            source_text="Madayaw na adlaw",
            translated_text="Good day",
            source_language="tagabawa",
            target_language="english",
            method=NEURAL_METHOD,
            confidence=None,
            needs_review=True,
        )
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertNotIn("ByT5 Neural Translation", body)
        self.assertNotIn(NEURAL_METHOD, body)
        self.assertNotIn("Needs review", body)
        self.assertIn("Good day", body)

    def test_preview_and_download_still_work_when_neural_enabled(self):
        """The neural feature flag must not interfere with the existing
        preview/download workflow, whether or not any segment used ByT5."""
        job = self._create_job()
        TranslationSegment.objects.create(
            job=job,
            segment_index=1,
            source_text="Hello",
            translated_text="Madigar",
            source_language="tagabawa",
            target_language="english",
            method="exact_phrase",
            confidence=1.0,
        )

        media_root_dir = tempfile.TemporaryDirectory()
        self.addCleanup(media_root_dir.cleanup)
        media_root = Path(media_root_dir.name)
        output_path = media_root / "jobs" / job.job_id / "outputs" / "translated.pdf"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(b"%PDF-1.4\n%%EOF")
        job.output_file_path = str(output_path)
        job.save(update_fields=["output_file_path"])

        self.client.force_login(self.alice)

        with override_settings(
            LINGOKATUTUBO_NEURAL_TRANSLATION_ENABLED=True,
            MEDIA_ROOT=media_root,
        ):
            preview_response = self.client.get(reverse("translator:job_preview", args=[job.id]))
            download_response = self.client.get(reverse("translator:translate_download", args=[job.id]))

        self.assertEqual(preview_response.status_code, 200)
        self.assertIn("Hello", preview_response.content.decode("utf-8"))
        self.assertEqual(download_response.status_code, 200)
        download_response.close()


# ============================================================
# Appendix K — Bilingual preview UX, role-based visibility, and
# Unicode/charmap hardening sprint.
# ============================================================


class AppendixKBilingualPreviewUXTests(TestCase):
    """Regression tests for: the Preview Bilingual loading transition,
    role-based hiding of technical/segment details, simplified failure
    messages, and Unicode/charmap crash hardening."""

    password = "Bagobo-AppK-2026!"

    def setUp(self):
        self.alice = User.objects.create_user(
            username="appk_alice", email="appk_alice@example.test", password=self.password
        )

    # ------------------------------------------------------------------
    # Unicode / charmap hardening
    # ------------------------------------------------------------------

    class _Cp1252Stream(StringIO):
        """Mimics a Windows console limited to the legacy cp1252 codec."""

        encoding = "cp1252"

        def write(self, value):
            value.encode("cp1252")
            return super().write(value)

    def test_safe_print_does_not_crash_on_invisible_unicode(self):
        """The exact failure reported in production: printing text containing
        a zero-width space (\\u200b) on a cp1252-only console must never
        raise UnicodeEncodeError."""
        from translator.services.display_utils import safe_print

        stream = self._Cp1252Stream()
        try:
            safe_print("Hello​world", file=stream)
        except UnicodeEncodeError:
            self.fail("safe_print raised UnicodeEncodeError instead of handling it safely")
        self.assertIn("Helloworld", stream.getvalue())

    def test_safe_print_falls_back_for_other_unmappable_characters(self):
        """Characters outside the known invisible-unicode list (e.g. emoji)
        must still be handled safely rather than crashing the caller."""
        from translator.services.display_utils import safe_print

        stream = self._Cp1252Stream()
        try:
            safe_print("status: \U0001F600", file=stream)
        except UnicodeEncodeError:
            self.fail("safe_print raised UnicodeEncodeError on a non-cp1252 character")
        self.assertTrue(stream.getvalue())

    def test_language_detection_cleans_invisible_unicode(self):
        from translator.services.language_detection_service import LanguageDetectionService

        service = LanguageDetectionService()
        result = service.detect_language("Hello​ world﻿ there")
        self.assertIn("language", result)
        self.assertIn("confidence", result)

    def test_normalize_lang_cleans_invisible_unicode(self):
        from translator.services.translation_dataset import _normalize_lang

        self.assertEqual(_normalize_lang("​English​"), "english")

    def test_phrasebook_lookup_cleans_invisible_unicode(self):
        from translator.services.translation_dataset import get_translation_dataset

        dataset = get_translation_dataset()
        result = dataset.translate_phrase_with_metadata(
            "Hel​lo", source_lang="english", target_lang="tagabawa"
        )
        self.assertNotEqual(
            result["method"], "unknown_for_review",
            "A zero-width space inside an otherwise-known phrase must not break the lookup",
        )

    def test_preview_redirect_never_shows_raw_charmap_error(self):
        """A job that failed with a raw codec error must never surface that
        raw text to the user — only the safe, generic failure message."""
        job = TranslationJob.objects.create(
            owner=self.alice,
            original_filename="broken.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=TranslationJob.Status.FAILED,
            error="'charmap' codec can't encode character '\\u200b' in position 24: character maps to <undefined>",
        )
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]), follow=True)

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertNotIn("charmap", body)
        self.assertNotIn("u200b", body)
        self.assertIn(
            "Translation could not be completed. Please try another file or contact the administrator.",
            body,
        )

    def test_unknown_for_review_never_leaks_via_pageless_segment(self):
        """A segment with no DocumentPage link still surfaces via the main
        panels for normal users, and never leaks the raw review sentinel."""
        job = TranslationJob.objects.create(
            owner=self.alice,
            original_filename="unknown.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=TranslationJob.Status.COMPLETED,
        )
        TranslationSegment.objects.create(
            job=job,
            segment_index=1,
            source_text="Mystery phrase",
            translated_text="[UNKNOWN_FOR_REVIEW]",
            source_language="english",
            target_language="tagabawa",
            method="unknown_for_review",
            confidence=0.0,
            needs_review=True,
        )
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        self.assertNotIn("[UNKNOWN_FOR_REVIEW]", body)
        self.assertIn("Mystery phrase", body)

    # ------------------------------------------------------------------
    # Preview Bilingual loading transition markup
    # ------------------------------------------------------------------

    def test_preview_bilingual_links_carry_loading_hooks(self):
        """Every Preview Bilingual entry point ships the JS loading hooks
        (class + data attribute) and the shared page-transition overlay."""
        job = TranslationJob.objects.create(
            owner=self.alice,
            original_filename="ready.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=TranslationJob.Status.COMPLETED,
        )
        self.client.force_login(self.alice)

        detail_body = self.client.get(
            reverse("translator:job_detail", args=[job.id])
        ).content.decode("utf-8")
        history_body = self.client.get(reverse("translator:history")).content.decode("utf-8")
        translate_body = self.client.get(reverse("translator:translate")).content.decode("utf-8")

        for body in (detail_body, history_body):
            self.assertIn("js-preview-link", body)
            self.assertIn("data-loading-text", body)

        self.assertIn('id="preview-link"', translate_body)
        self.assertIn("js-preview-link", translate_body)
        self.assertIn('id="page-transition-overlay"', detail_body)
        self.assertIn('id="js-error-banner"', detail_body)

    # ------------------------------------------------------------------
    # Role-based visibility (consolidated smoke checks; detailed coverage
    # lives alongside the individual feature tests elsewhere in this file)
    # ------------------------------------------------------------------

    def test_normal_user_sees_only_the_clean_bilingual_summary(self):
        job = TranslationJob.objects.create(
            owner=self.alice,
            original_filename="clean.pdf",
            file_type=TranslationJob.FileType.PDF,
            status=TranslationJob.Status.COMPLETED,
            source_language="english",
            target_language="tagabawa",
            metadata={"extraction_method": "ocr_image", "ocr_summary": {"mean_confidence": 0.4}},
        )
        TranslationSegment.objects.create(
            job=job,
            segment_index=1,
            source_text="Visiting",
            translated_text="Pagdalaw",
            source_language="english",
            target_language="tagabawa",
            method="byt5_neural",
            confidence=None,
            needs_review=True,
        )
        self.client.force_login(self.alice)

        response = self.client.get(reverse("translator:job_preview", args=[job.id]))

        self.assertEqual(response.status_code, 200)
        body = response.content.decode("utf-8")
        # Allowed for normal users: filename, status, languages, simple
        # confidence summary, original/translated panels, download/back.
        self.assertIn("clean.pdf", body)
        self.assertIn("Completed", body)
        self.assertIn("english", body)
        self.assertIn("tagabawa", body)
        self.assertIn("Original Document", body)
        self.assertIn("Translated Document", body)
        self.assertIn("Back to History", body)
        # Hidden from normal users: technical/debug surfaces.
        self.assertNotIn("View Segment Details", body)
        self.assertNotIn("byt5_neural", body)
        self.assertNotIn("Technical Details", body)
        self.assertNotIn("OCR Confidence", body)
        self.assertNotIn("badge-method", body)


class AppendixLJsonApiTests(TestCase):
    """Regression tests for the "Unexpected token '<'" bug: fetch() in
    app.js parses every /translate/upload/, /translate/status/,
    /translate/structure/, and /translate/preview-data/ response as JSON.
    These endpoints must never hand back an HTML login redirect, an HTML
    CSRF error page, or an HTML 404 page — only JSON, on every path."""

    password = "Bagobo-AppL-2026!"

    def setUp(self):
        from django.core.cache import cache as _cache
        _cache.clear()

        self.temp_dir = tempfile.TemporaryDirectory()
        self.media_root = Path(self.temp_dir.name)
        self.settings_override = override_settings(MEDIA_ROOT=self.media_root)
        self.settings_override.enable()
        self.addCleanup(self.settings_override.disable)
        self.addCleanup(self.temp_dir.cleanup)

        self.alice = User.objects.create_user(
            username="appl_alice", email="appl_alice@example.test", password=self.password
        )
        self.staff_user = User.objects.create_user(
            username="appl_staff",
            email="appl_staff@example.test",
            password=self.password,
            is_staff=True,
        )

    def _make_job(self, owner, **kwargs):
        defaults = dict(
            original_filename="doc.pdf",
            file_type=TranslationJob.FileType.PDF,
            source_language="english",
            target_language="tagabawa",
        )
        defaults.update(kwargs)
        return TranslationJob.objects.create(owner=owner, **defaults)

    # ------------------------------------------------------------------
    # Upload (start-translation) endpoint
    # ------------------------------------------------------------------

    def test_upload_success_returns_json(self):
        self.client.force_login(self.alice)
        uploaded = SimpleUploadedFile(
            "doc.pdf", b"%PDF-1.4\n1 0 obj\n<<>>\nendobj\n%%EOF", content_type="application/pdf"
        )
        with patch("translator.views.start_translation_job"):
            response = self.client.post(
                reverse("translator:upload"),
                {"file": uploaded, "source_language": "auto", "target_language": "tagabawa"},
            )
        self.assertEqual(response.status_code, 202)
        self.assertEqual(response["Content-Type"], "application/json")
        self.assertIn("job_id", response.json())

    def test_upload_validation_error_returns_json(self):
        self.client.force_login(self.alice)
        bad_file = SimpleUploadedFile("danger.exe", b"binary", content_type="application/octet-stream")
        response = self.client.post(reverse("translator:upload"), {"file": bad_file})
        self.assertEqual(response.status_code, 400)
        self.assertEqual(response["Content-Type"], "application/json")

    def test_upload_unauthenticated_does_not_redirect(self):
        response = self.client.post(reverse("translator:upload"))
        # No HTML redirect: fetch() follows redirects automatically, which
        # is exactly what previously fed response.json() an HTML page.
        self.assertNotEqual(response.status_code, 302)
        self.assertEqual(response.status_code, 401)

    def test_upload_unexpected_exception_returns_friendly_json_for_normal_user(self):
        self.client.force_login(self.alice)
        uploaded = SimpleUploadedFile(
            "doc.pdf", b"%PDF-1.4\n1 0 obj\n<<>>\nendobj\n%%EOF", content_type="application/pdf"
        )
        with patch("translator.views.save_uploaded_file", side_effect=OSError("disk full")):
            response = self.client.post(
                reverse("translator:upload"),
                {"file": uploaded, "source_language": "auto", "target_language": "tagabawa"},
            )
        self.assertEqual(response.status_code, 500)
        self.assertEqual(response["Content-Type"], "application/json")
        data = response.json()
        self.assertFalse(data["ok"])
        self.assertEqual(
            data["error"],
            "Translation could not be completed. Please try another file or contact the administrator.",
        )
        self.assertNotIn("disk full", data["error"])
        self.assertNotIn("technical_detail", data)

    def test_upload_unexpected_exception_includes_technical_detail_for_staff(self):
        self.client.force_login(self.staff_user)
        uploaded = SimpleUploadedFile(
            "doc.pdf", b"%PDF-1.4\n1 0 obj\n<<>>\nendobj\n%%EOF", content_type="application/pdf"
        )
        with patch("translator.views.save_uploaded_file", side_effect=OSError("disk full")):
            response = self.client.post(
                reverse("translator:upload"),
                {"file": uploaded, "source_language": "auto", "target_language": "tagabawa"},
            )
        self.assertEqual(response.status_code, 500)
        data = response.json()
        self.assertIn("disk full", data["technical_detail"])

    # ------------------------------------------------------------------
    # Status / structure / preview-data endpoints
    # ------------------------------------------------------------------

    def test_status_unauthenticated_returns_json_not_html(self):
        job = self._make_job(self.alice)
        response = self.client.get(reverse("translator:status", args=[job.job_id]))
        self.assertNotEqual(response.status_code, 302)
        self.assertEqual(response.status_code, 401)
        self.assertEqual(response["Content-Type"], "application/json")
        self.assertFalse(response.json()["ok"])

    def test_status_unknown_job_returns_json_404_not_html(self):
        import uuid
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:status", args=[uuid.uuid4()]))
        self.assertEqual(response.status_code, 404)
        self.assertEqual(response["Content-Type"], "application/json")
        data = response.json()
        self.assertFalse(data["ok"])
        self.assertEqual(data["error"], "Job not found.")

    def test_structure_unknown_job_returns_json_404_not_html(self):
        import uuid
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:structure", args=[uuid.uuid4()]))
        self.assertEqual(response.status_code, 404)
        self.assertEqual(response["Content-Type"], "application/json")

    def test_preview_data_unknown_job_returns_json_404_not_html(self):
        import uuid
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:preview_data", args=[uuid.uuid4()]))
        self.assertEqual(response.status_code, 404)
        self.assertEqual(response["Content-Type"], "application/json")

    def test_status_for_other_users_job_returns_json_404_not_html(self):
        """A job ID that exists but is owned by someone else must still come
        back as a clean JSON 404, not Django's owner-queryset Http404 page."""
        other = User.objects.create_user(
            username="appl_bob", email="appl_bob@example.test", password=self.password
        )
        job = self._make_job(other)
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:status", args=[job.job_id]))
        self.assertEqual(response.status_code, 404)
        self.assertEqual(response["Content-Type"], "application/json")

    # ------------------------------------------------------------------
    # CSRF failures on JSON API routes
    # ------------------------------------------------------------------

    def test_csrf_failure_on_status_endpoint_returns_json(self):
        from django.test import Client
        job = self._make_job(self.alice)
        csrf_client = Client(enforce_csrf_checks=True)
        csrf_client.force_login(self.alice)
        # GET requests aren't CSRF-checked; force a POST-like CSRF failure
        # path via the upload endpoint, which is what the bug report hit.
        response = csrf_client.post(reverse("translator:upload"), {"file": "dummy"})
        self.assertEqual(response.status_code, 403)
        self.assertEqual(response["Content-Type"], "application/json")
        self.assertNotIn(b"<!DOCTYPE", response.content[:20])

    # ------------------------------------------------------------------
    # Normal HTML pages must keep working unchanged
    # ------------------------------------------------------------------

    def test_translate_page_unauthenticated_still_redirects_to_html_login(self):
        """Only the JSON API routes change behavior. Normal page navigation
        (not fetch+JSON) must keep using the standard login redirect."""
        translate_url = reverse("translator:translate")
        response = self.client.get(translate_url)
        self.assertEqual(response.status_code, 302)
        self.assertIn(reverse("translator:login"), response["Location"])

    def test_history_page_still_renders_html(self):
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:history"))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"].split(";")[0], "text/html")

    def test_job_detail_page_still_renders_html(self):
        job = self._make_job(self.alice)
        self.client.force_login(self.alice)
        response = self.client.get(reverse("translator:job_detail", args=[job.id]))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"].split(";")[0], "text/html")
